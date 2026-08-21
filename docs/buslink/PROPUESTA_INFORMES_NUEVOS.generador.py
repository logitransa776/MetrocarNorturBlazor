# -*- coding: utf-8 -*-
"""
Genera PROPUESTA_INFORMES_NUEVOS.docx a partir de PROPUESTA_INFORMES_NUEVOS.md.
Mismo patrón visual que INFORME_ESTADO_MIGRACION.generador.py (colores NORTUR).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

AZUL   = RGBColor(0x00, 0x3A, 0xA0)
NARANJ = RGBColor(0xF9, 0x94, 0x10)
GRIS   = RGBColor(0x5A, 0x60, 0x6A)
ROJO   = RGBColor(0xB3, 0x26, 0x1E)
VERDE  = RGBColor(0x1B, 0x6E, 0x3C)
NEGRO  = RGBColor(0x1F, 0x24, 0x2C)

HEX_AZUL   = "003AA0"
HEX_NARANJ = "F99410"
HEX_CLARO  = "EEF2F9"
HEX_SUAVE  = "F7F8FA"
HEX_AVISO  = "FFF4E3"
HEX_ROJO   = "FDECEA"
HEX_VERDE  = "EAF5EE"

BODY = "Segoe UI"
MONO = "Consolas"

doc = Document()

st = doc.styles["Normal"]
st.font.name = BODY
st.font.size = Pt(10.5)
st.font.color.rgb = NEGRO
st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15

for sec in doc.sections:
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.9)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)


def shade(el, hexcolor):
    el.append(parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls("w"), hexcolor)))


def border(par, edge="bottom", size=12, color=HEX_NARANJ, space=4):
    pPr = par._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    e = OxmlElement("w:" + edge)
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(size))
    e.set(qn("w:space"), str(space))
    e.set(qn("w:color"), color)
    pBdr.append(e)


def run(par, text, bold=False, italic=False, size=10.5, color=None, font=BODY, hl=None):
    r = par.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color is not None:
        r.font.color.rgb = color
    if hl:
        shade(r._element.get_or_add_rPr(), hl)
    return r


def p(text="", size=10.5, bold=False, italic=False, color=None, space_after=6,
      space_before=0, align=None, left=0):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    if left:
        par.paragraph_format.left_indent = Cm(left)
    if align:
        par.alignment = align
    if text:
        run(par, text, bold=bold, italic=italic, size=size, color=color)
    return par


def rich(parts, space_after=6, space_before=0, left=0, size=10.5):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    if left:
        par.paragraph_format.left_indent = Cm(left)
    for txt, fmt in parts:
        f = dict(fmt)
        f.setdefault("size", size)
        run(par, txt, **f)
    return par


def h1(text, nro=None):
    par = doc.add_paragraph()
    par.paragraph_format.page_break_before = True
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.keep_with_next = True
    if nro:
        run(par, nro + "  ", bold=True, size=17, color=NARANJ)
    run(par, text, bold=True, size=17, color=AZUL)
    border(par, "bottom", size=10, color=HEX_NARANJ, space=3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def h2(text, color=AZUL):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.keep_with_next = True
    run(par, text, bold=True, size=13, color=color)
    return par


def h3(text, color=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(9)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.keep_with_next = True
    run(par, text, bold=True, size=11, color=color or GRIS)
    return par


def bullet(text_parts, level=0):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(0.75 + 0.6 * level)
    if isinstance(text_parts, str):
        run(par, text_parts)
    else:
        for txt, fmt in text_parts:
            f = dict(fmt)
            f.setdefault("size", 10.5)
            run(par, txt, **f)
    return par


def _cell_margins(cell, lr=110, tb=70):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", tb), ("left", lr), ("bottom", tb), ("right", lr)):
        e = OxmlElement("w:" + side)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def _tbl_borders(tbl, color, left_only=False, size=4):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    sides = ["left"] if left_only else ["top", "left", "bottom", "right", "insideH", "insideV"]
    todos = ["top", "left", "bottom", "right", "insideH", "insideV"]
    for s in todos:
        e = OxmlElement("w:" + s)
        if s in sides:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(24 if left_only else size))
            e.set(qn("w:color"), color)
        else:
            e.set(qn("w:val"), "none")
        e.set(qn("w:space"), "0")
        borders.append(e)
    tblPr.append(borders)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:tblHeader")
    e.set(qn("w:val"), "true")
    trPr.append(e)


def callout(titulo, texto, fill=HEX_AVISO, color_titulo=NARANJ, borde=HEX_NARANJ):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), fill)
    cell.text = ""
    par = cell.paragraphs[0]
    par.paragraph_format.space_after = Pt(2)
    run(par, titulo, bold=True, size=10.5, color=color_titulo)
    par2 = cell.add_paragraph()
    par2.paragraph_format.space_after = Pt(0)
    if isinstance(texto, str):
        run(par2, texto, size=10)
    else:
        for txt, fmt in texto:
            f = dict(fmt)
            f.setdefault("size", 10)
            run(par2, txt, **f)
    _cell_margins(cell, 140, 110)
    _tbl_borders(t, borde, left_only=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def tabla(headers, filas, widths=None, size=9.5, header_fill=HEX_AZUL, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, htxt in enumerate(headers):
        c = hdr.cells[i]
        if header_fill:
            shade(c._tc.get_or_add_tcPr(), header_fill)
        c.text = ""
        par = c.paragraphs[0]
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        run(par, htxt, bold=True, size=size, color=RGBColor(0xFF, 0xFF, 0xFF))
        _cell_margins(c, 90, 55)
    _repeat_header(hdr)
    for j, fila in enumerate(filas):
        row = t.add_row()
        for i, celda in enumerate(fila):
            c = row.cells[i]
            if zebra and j % 2 == 1:
                shade(c._tc.get_or_add_tcPr(), HEX_SUAVE)
            c.text = ""
            par = c.paragraphs[0]
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.space_before = Pt(0)
            if isinstance(celda, str):
                run(par, celda, size=size)
            else:
                for txt, fmt in celda:
                    f = dict(fmt)
                    f.setdefault("size", size)
                    run(par, txt, **f)
            _cell_margins(c, 90, 55)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    if header_fill is None:
        t._tbl.remove(t.rows[0]._tr)
    _tbl_borders(t, "D5DAE3", size=4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def code_block(lineas, size=9):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), "F4F6F9")
    cell.text = ""
    for i, ln in enumerate(lineas):
        par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        run(par, ln, font=MONO, size=size, color=RGBColor(0x33, 0x39, 0x42))
    _cell_margins(cell, 140, 90)
    _tbl_borders(t, "C9D2E0", left_only=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


B = {"bold": True}
N = {}
R = {"color": ROJO, "bold": True}
V = {"color": VERDE, "bold": True}
G = {"color": GRIS}


def mono(txt):
    return (txt, {"font": MONO, "size": 9.5, "color": AZUL})


# ═══════════════════════════ PORTADA ═══════════════════════════
for _ in range(3):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(0)
run(par, "BUSLINK  ·  METROCAR NORTUR", bold=True, size=11, color=NARANJ)
par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(0)
run(par, "Propuesta de informes nuevos", size=11, color=GRIS)

par = doc.add_paragraph()
par.paragraph_format.space_before = Pt(18)
par.paragraph_format.space_after = Pt(0)
run(par, "Seis informes que", bold=True, size=32, color=AZUL)
par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(4)
run(par, "hoy nadie puede ver", bold=True, size=32, color=NARANJ)
border(par, "bottom", size=18, color=HEX_NARANJ, space=8)

par = doc.add_paragraph()
par.paragraph_format.space_before = Pt(14)
run(par, "Ideas de informes nuevos para Buslink — que no existen ni en Buslink ni en el "
         "FoxPro — medidas contra la base replicaVPF, no estimadas. Incluye qué se descartó "
         "y por qué, y las queries para reproducir cada número.",
    size=12, color=GRIS)

for _ in range(2):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

tabla(
    ["", ""],
    [
        [[("Fecha", B)], "10 de agosto de 2026"],
        [[("Estado", B)], "Propuesta — ninguno construido todavía"],
        [[("Alcance", B)], "Informes que NO tienen gemelo en el FoxPro: preguntas nuevas, no migraciones"],
        [[("Método", B)], "Cada afirmación se verificó con COUNT/GROUP BY directo contra replicaVPF (año 2026, "
                          "_deleted = 0), no se estimó"],
    ],
    widths=[3.4, 12.4], size=10, header_fill=None, zebra=True,
)


# ═══════════════════════ 1. DIAGNÓSTICO ═══════════════════════
h1("Diagnóstico: qué contesta Buslink hoy y qué no", "1")

p("El catálogo (Services/InformesCatalogo.cs) tiene 17 informes disponibles y 4 marcados "
  "“Próximamente”, repartidos así:")

tabla(
    ["Módulo", "Disponibles"],
    [["Reservas", "3"], ["Tráfico", "2"], ["Flota", "4"],
     ["Facturación", "4"], ["Combustible", "3"], ["Sistema", "1"]],
    widths=[10, 5.8],
)

p("Vistos juntos, los 17 contestan variantes de dos preguntas: cuánto volumen (reservas, "
  "pax, km) y cuánta plata, agregados por entidad (cliente, servicio, chofer, unidad). Es "
  "una cobertura buena de el qué, y nula de el cómo.")

h2("Las cinco dimensiones que ningún informe toca")

tabla(
    ["Dimensión", "Evidencia en la base (2026)", "Informes que la usan"],
    [
        ["Tiempo real de operación", "53.036 viajes con hs_inicio y hs_fin. Duración real promedio: 91 min", "0"],
        ["Tercerización", "viaje.fletero poblado al 90%: NORTUR 43.295 vs terceros 9.767", "0"],
        ["Quién opera el sistema", "u_create poblado al 100%. LEONARDO: 47.578 altas = 77% del total", "0"],
        ["Calidad del despacho", "chequeo = 1 en 115.110 de 172.301 (67%). EZEIZA al 48%, CABECERA_KM al 73%", "0"],
        ["Plata a nivel viaje", "liquidacion_detalle.id_viaje permite bajar el ingreso a la unidad y al servicio",
         "0 (solo por cliente)"],
    ],
    widths=[3.6, 10.4, 2.2], size=9,
)


# ═══════════════════════ 2. TRAMPAS ═══════════════════════
h1("Trampas de datos verificadas", "2")

p("Condicionan el diseño de cualquier informe nuevo. Confirmadas con COUNT, no supuestas.")

tabla(
    ["Campo", "Realidad medida", "Consecuencia"],
    [
        ["viaje.total, .importe, .descuento,\n.hext, .imp_liq, .adi_*",
         "0 filas con valor ≠ 0 en 2025-2026",
         "La plata vive solo en liquidacion_detalle. No hay costo de chofer en la base"],
        ["viaje.km_recorri", "3 filas pobladas", "Inservible. Usar viaje.km (82% poblado)"],
        ["viaje.hs_present", "13.571 de 172.301 (8%)",
         "No se puede medir puntualidad de presentación. Idea descartada"],
        ["viaje.hs_fin_apr",
         "100% poblado, pero es inicio + 2h fijo en ~30% de los casos",
         "Es un presupuesto, no un hecho. Sirve como línea base contra la cual medir "
         "desvío — nunca como duración real"],
        ["viaje.id_operado", "3 filas pobladas", "El operador comercial no se puede analizar desde viaje"],
        ["viaje.f_modify en cancelados",
         "1.967 de 2.304 cancelados de 2026 sin f_modify",
         "La anticipación de la cancelación solo se mide en el 15% de los casos. Idea descartada"],
        ["CABECERA_KM / CABECERA_SERV", "~74% del volumen",
         "Son modos de facturación, no servicios. Excluir o segmentar siempre"],
    ],
    widths=[3.6, 5.4, 7.2], size=9,
)


# ═══════════════════════ 3. LAS SEIS PROPUESTAS ═══════════════════════
h1("Las seis propuestas", "3")

callout(
    "Patrón de construcción",
    "Cada informe sigue la regla vigente de CLAUDE.md § 7: patrón dashboard completo "
    "(filtros compactos, KPIs, ApexCharts con color por entidad, pivote con drill-down, "
    "Excel multi-hoja y cross-filter). El alta en el hub es una línea en InformesCatalogo.cs.",
    fill=HEX_CLARO, color_titulo=AZUL, borde=HEX_AZUL,
)


def informe(nro, titulo, modulo, pregunta, fuente_parts, contenido_items,
            porque_parts, esfuerzo, riesgo, advertencia=None, cuidado=None, extra_parts=None):
    h2("{}  {}   —  {}".format(nro, titulo, modulo), color=AZUL)
    rich([("Pregunta: ", B), (pregunta, {})], space_after=4)
    rich([("Fuente: ", B)] + fuente_parts, space_after=4)
    h3("Contenido", color=GRIS)
    for it in contenido_items:
        bullet(it)
    if extra_parts:
        rich(extra_parts, space_after=4, space_before=4)
    h3("Por qué importa", color=GRIS)
    rich(porque_parts, space_after=6)
    if advertencia:
        callout("Advertencia", advertencia, fill=HEX_ROJO, color_titulo=ROJO, borde="B3261E")
    if cuidado:
        callout("Cuidado de diseño", cuidado, fill=HEX_AVISO, color_titulo=NARANJ, borde=HEX_NARANJ)
    rich([("Esfuerzo: ", B), (esfuerzo, {}), ("     Riesgo: ", B), (riesgo, {})], space_after=10)


callout(
    "✅ CONSTRUIDO el 11/08/2026",
    [("Vive en ", N), ("/panel-operador", {"font": MONO, "size": 9.5, "color": AZUL}),
     (", bajo Informes → En desarrollo (permiso S) hasta que el cliente lo apruebe. "
      "El alcance de abajo es el REAL, ya corregido contra lo que los datos permiten.", N)],
    fill=HEX_VERDE, color_titulo=VERDE, borde="1B6E3C",
)

informe(
    "3.1", "Panel del Operador (auditoría de carga)", "Sistema",
    "¿Quién carga el trabajo, cuándo, y quién modifica lo que cargó otro?",
    [("viaje.u_create, f_create, u_modify, f_modify — los cuatro poblados. Una sola tabla, "
      "sin joins. Es el más barato de los seis.", {})],
    [
        "Ranking de altas por operador y evolución diaria/mensual de la carga",
        "Antelación: días entre que se cargó la reserva y la fecha del viaje, y las cargas "
        "retroactivas (cargadas después de que el viaje ocurrió)",
        "Matriz quién modificó lo de quién — el indicador de fricción entre operadores",
        "Concentración: % de la carga en el operador top, con el mix de clientes al lado — es lo "
        "que distingue “una persona hace todo” de “a una persona le tocó el contrato grande”",
        "Calidad de lo cargado: % cancelado y % sin asignar por operador",
        "Control: operadores con cargas que no figuran en el padrón de Usuarios",
    ],
    [("1. Hoy, como gestión. ", B),
     ("LEONARDO cargó el 77% de las reservas de 2026 (47.578 de ~62.000). OSVALDO el 20%. "
      "Los otros nueve usuarios se reparten el 3%. Es un riesgo de continuidad que el dueño "
      "no tiene medido.\n", {}),
     ("2. Después del día D, como control. ", B),
     ("Cuando Buslink pase a escribir el circuito viaje, este informe es la caja negra que "
      "contesta “¿quién tocó esto?”. Conviene tenerlo antes del corte, para tener "
      "línea base de cómo se cargaba en FoxPro.", {})],
    "Bajo", "Ninguno (solo lectura, una tabla)",
    cuidado=[("Dos ideas de la propuesta original NO se pudieron hacer (medido 11/08/2026): ", N),
             ("la curva horaria de carga", B),
             (" — f_create es date, sin hora, y _created_at es el timestamp de la réplica (todas "
              "las filas comparten el mismo instante de importación) — y ", N),
             ("la latencia de asignación", B),
             (", porque no hay historial de estados: no existe registro de cuándo una reserva pasó "
              "a ASIGNADO. Se reemplazaron por la antelación y la calidad de lo cargado, que son "
              "mejores y sí tienen respaldo en el dato.", N)],
)

informe(
    "3.2", "Panel de Tercerización — Nortur vs Fleteros", "Flota",
    "¿Cuánto de la operación se está dando a terceros, y cuánto podría absorber la flota propia?",
    [("viaje.fletero (90% poblado). Valores 2026: NORTUR 43.295 · vacío 9.523 · VANSQ 3.066 · "
      "MVTRAVEL 2.368 · NEUQUEN 2.163 · TEB 1.221 · TEDESCHI 505 · MASIMIGLIA 381 · y una cola menor.",
      {})],
    [
        "% tercerizado por mes, cliente, servicio y tipo de unidad; km y pax tercerizados",
        "Ranking de fleteros con tendencia mes a mes",
        "El cruce que lo justifica: días y franjas donde se contrató afuera teniendo unidades "
        "propias libres — reutilizando el modelo de Services/OcupacionFlota.cs, sin query nueva",
    ],
    [("Es el informe de más impacto económico directo de los seis, y engancha con lo que ya "
      "destapó el Panel de Flota (33% de la demanda sin cubrir con flota propia).", {})],
    "Medio", "Bajo",
    advertencia=[("Este es el campo correcto para separar flota propia de contratada. La regla "
                  "FoxPro “interno ≥ 1000 = contratado” está rota.", {})],
    cuidado=[("Limitación a declarar en pantalla: ", B),
             ("se mide volumen, no costo. La liquidación a fleteros está muerta desde el "
              "21/12/2023, así que no hay contra qué valorizar lo tercerizado.", {})],
)

informe(
    "3.3", "Panel de Puntualidad y Duración Real", "Tráfico",
    "¿Qué servicios se van sistemáticamente del tiempo presupuestado?",
    [("hs_inicio → hs_fin (duración real, 53.036 casos en 2026, promedio 91 min) contra "
      "hs_fin_apr (duración presupuestada).", {})],
    [
        "Desvío promedio y mediano (real − presupuestado) por servicio, cliente, chofer y franja",
        "Distribución de duraciones por servicio (dónde está la cola larga, no solo el promedio)",
        "Los servicios que se pasan siempre vs los que se pasan a veces",
        "Filtro que aísle los casos con hs_fin_apr = default de 2 h, que son los que más ruido meten",
    ],
    [("Es munición comercial concreta — “el CITY que te facturo como 3 h me ocupa "
      "4”. Y es la primera métrica de tiempo real que va a tener la empresa.", {})],
    "Medio", "Medio — el valor depende de comunicar bien la limitación",
    cuidado=[("El informe debe explicar en pantalla que hs_fin_apr es un presupuesto con "
              "default, o el usuario va a leer el desvío como error de los choferes.", {})],
)

informe(
    "3.4", "Tablero de Cumplimiento Operativo (chequeo y avisos)", "Tráfico",
    "¿Qué servicios están saliendo sin chequear?",
    [("viaje.chequeo, chequeo_ag, hs_aviso (el motor de avisos F4 ya está vivo en Buslink).", {})],
    [
        "% chequeado por servicio, franja horaria, día de semana y operador",
        "Listado de los que se escaparon, exportable, para revisión del jefe de tráfico",
        "Evolución mensual del cumplimiento",
    ],
    [("Que los servicios de aeropuerto — los más sensibles, con vuelo de por medio — "
      "chequeen por debajo de las cabeceras es exactamente el tipo de hallazgo por el que "
      "después piden el informe formal.", {})],
    "Bajo-medio", "Bajo",
    extra_parts=None,
)

h3("El hallazgo que lo dispara (2026, hasta 10/07)", color=GRIS)
tabla(
    ["Servicio", "Viajes", "Chequeados", "%"],
    [
        ["CABECERA_KM", "36.058", "26.464", "73%"],
        ["CABECERA_SERV", "9.280", "5.654", "61%"],
        ["TRASLADO", "2.296", "1.351", "59%"],
        [[("EZEIZA", B)], [("1.896", B)], [("914", B)], [("48%", R)]],
        ["GUARDIA8", "1.579", "934", "59%"],
        ["AEROPARQUE", "1.353", "812", "60%"],
        ["CITY", "1.060", "428", "40%"],
        ["CENA SHOW", "635", "229", "36%"],
    ],
    widths=[4.6, 3, 3.4, 2.4], size=9.5,
)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

informe(
    "3.5", "Rentabilidad por unidad y por servicio", "Facturación",
    "¿Qué unidad y qué servicio dejan plata?",
    [("liquidacion_detalle (tiene id_viaje → se puede bajar el ingreso al viaje, y de ahí a la "
      "unidad por viaje.id_vehicu2) cruzado con combustible por unidad (vehiculo_sobre) y km "
      "(viaje.km).", {})],
    [
        "Ingreso, litros, km y contribución por unidad, por servicio y por mes",
        "$/km e ingreso por unidad disponible",
        "Ranking de unidades y servicios por contribución",
    ],
    [("Además: la facturación hay que calcularla del detalle (importe + incremento − "
      "descuento, moneda por línea), no de liquidacion.total, que tiene cargas corruptas.", {})],
    "Alto",
    "Alto — es el que más se puede equivocar y el que peor se perdona si da un número mal",
    advertencia=[("viaje.imp_liq (costo de chofer) está en 0 en toda la tabla. El margen que se "
                  "puede calcular es contribución antes de mano de obra, y hay que rotularlo "
                  "así. Vender esto como “rentabilidad” a secas sería mentirle al "
                  "dueño.", {})],
)

informe(
    "3.6", "Curva de demanda y dimensionamiento de flota", "Flota",
    "¿Cuántas unidades necesito de verdad un martes a las 7 de la mañana?",
    [("hs_inicio / hs_fin_apr de todo un rango, no de un día. Extiende el tablero de ocupación "
      "(Services/OcupacionFlota.cs), que hoy es de una sola fecha.", {})],
    [
        "Simultaneidad de servicios en bandas de 15 min, promediada por día de semana y mes",
        "Pico de demanda contra flota disponible → dónde estructuralmente no se llega",
        "Heatmap día de semana × hora",
        "Cruce con 3.2: los picos no cubiertos son los que se están tercerizando",
    ],
    [("Es el informe que justifica comprar (o no comprar) una unidad, y el más vistoso para "
      "mostrarle al dueño.", {})],
    "Medio-alto", "Medio",
    cuidado=[("Usa hs_fin_apr para el fin (es lo único disponible a futuro), con el sesgo del "
              "default de 2 h → infla la ocupación ~55%. Hay que mostrar la banda "
              "optimista/pesimista, no un número solo.", {})],
)


# ═══════════════════════ 4. ORDEN RECOMENDADO ═══════════════════════
h1("Orden recomendado", "4")

tabla(
    ["#", "Informe", "Esfuerzo", "Impacto", "Por qué en ese lugar"],
    [
        ["1", [("Panel del Operador  ✅ hecho", V)], "Bajo", "Alto",
         "Barato, sin riesgo, y conviene tener la línea base antes del día D"],
        ["2", [("Panel de Tercerización", B)], "Medio", "Muy alto",
         "El de más impacto económico directo"],
        ["3", [("Puntualidad y Duración Real", B)], "Medio", "Alto",
         "Primera métrica de tiempo real de la empresa; argumento comercial"],
        ["4", [("Curva de demanda", B)], "Medio-alto", "Alto",
         "El más vistoso; se apoya en 3.2"],
        ["5", [("Cumplimiento Operativo", B)], "Bajo-medio", "Medio",
         "Barato, pero es control interno: menos vendible hacia afuera"],
        ["6", [("Rentabilidad por unidad", B)], "Alto", "Alto",
         "Último a propósito: el que peor se perdona si da mal"],
    ],
    widths=[0.8, 3.6, 2.2, 2.2, 6.4], size=9,
)


# ═══════════════════════ 5. DESCARTADOS ═══════════════════════
h1("Descartados, y por qué", "5")

tabla(
    ["Idea", "Motivo"],
    [
        ["Anticipación de cancelaciones",
         "1.967 de 2.304 cancelados de 2026 no tienen f_modify. Solo se podría medir el 15%"],
        ["Puntualidad de presentación", "hs_present poblado al 8%"],
        ["Lead time de reserva",
         "El 78% cae en “2-7 días” porque lo generan las plantillas: mide el proceso "
         "interno, no al cliente"],
        ["Apercibimientos por chofer", "chofer_sancion con 0 filas (auditoría 09/08/2026)"],
        ["Liquidación a fleteros", "Muerta desde el 21/12/2023"],
    ],
    widths=[4.6, 10.6], size=9.5,
)

h2("Bonus baratos (fuera del top 6)")
bullet([("Panel de Choferes — ", B),
        ("gemelo del Panel de Flota: quién trabaja y quién no, francos, antigüedad, registro "
         "por vencer, siniestros. Todas las tablas ya están migradas, así que es casi todo "
         "ensamblado.", {})])
bullet([("Informe de vuelos — ", B),
        ("viaje.vuelo está poblado al 98% con códigos reales (AR1775, LA426…), aunque "
         "57.997 de 2026 son SIN VUELO. Nicho, para la operación de aeropuerto.", {})])


# ═══════════════════════ 6. QUERIES ═══════════════════════
h1("Queries de verificación", "6")

p("Todo lo afirmado acá se reproduce con estas consultas (server local "
  "DESKTOP-CV6LF0O\\SQLEXPRESS). Recordar: literales de fecha en formato "
  "‘yyyyMMdd’ sin guiones, y el server nuevo es SQL Server 2012 (sin STRING_AGG, "
  "TRIM, CONCAT_WS).")

code_block([
    "-- Población de campos (cambiar el campo en el COUNT)",
    "SELECT COUNT(*) total, COUNT(NULLIF(LTRIM(u_create),'')) poblado",
    "FROM viaje WHERE _deleted = 0 AND f_reserva >= '20250101' AND f_reserva < '20270101';",
    "",
    "-- Concentración de carga por operador",
    "SELECT TOP 12 u_create, COUNT(*) n, MIN(f_create) desde, MAX(f_create) hasta",
    "FROM viaje WHERE _deleted = 0 AND f_reserva >= '20260101'",
    "GROUP BY u_create ORDER BY 2 DESC;",
    "",
    "-- Tercerización",
    "SELECT TOP 10 fletero, COUNT(*) n",
    "FROM viaje WHERE _deleted = 0 AND f_reserva >= '20260101'",
    "GROUP BY fletero ORDER BY 2 DESC;",
    "",
    "-- Duración real",
    "SELECT COUNT(*) n, AVG(DATEDIFF(minute, hs_inicio, hs_fin) * 1.0) prom",
    "FROM viaje WHERE _deleted = 0 AND f_reserva >= '20260101'",
    "  AND hs_fin IS NOT NULL AND hs_inicio IS NOT NULL;",
    "",
    "-- Cumplimiento de chequeo por servicio",
    "SELECT TOP 8 id_servici, COUNT(*) n, SUM(CASE WHEN chequeo = 1 THEN 1 ELSE 0 END) chequeados",
    "FROM viaje WHERE _deleted = 0 AND f_reserva >= '20260101' AND f_reserva < '20260710'",
    "GROUP BY id_servici ORDER BY 2 DESC;",
])


# ── pie ──
doc.add_paragraph()
par = doc.add_paragraph()
border(par, "top", size=8, color="D5DAE3", space=6)
par.paragraph_format.space_before = Pt(14)
run(par, "Buslink · Metrocar Nortur — Propuesta de informes nuevos · 10 de agosto de 2026",
    size=9, color=GRIS)
par2 = doc.add_paragraph()
run(par2, "Todos los números medidos con COUNT/GROUP BY directo contra replicaVPF "
          "(DESKTOP-CV6LF0O\\SQLEXPRESS), año 2026, _deleted = 0.", size=9, color=GRIS)

for sec in doc.sections:
    footer_par = sec.footer.paragraphs[0]
    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_par.add_run()
    r.font.size = Pt(8.5)
    r.font.name = BODY
    r.font.color.rgb = GRIS
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    r._r.append(fld1); r._r.append(instr); r._r.append(fld2)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "PROPUESTA_INFORMES_NUEVOS.docx")
doc.save(out)
print("OK ->", out)
