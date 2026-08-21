# -*- coding: utf-8 -*-
"""
Genera la guia de teclas F1-F9 de la Planilla de Trafico (Buslink) en .docx.
Fuentes: docs/PlanoFoxPro/trafico/TRAFICO_F4_AVISO.md, TRAFICO_CRONOGRAMA.md,
TRAFICO_F2_NOVEDADES.md + la implementacion en AyudaTeclasDialog.razor.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

AZUL   = RGBColor(0x00, 0x3A, 0xA0)
NARANJ = RGBColor(0xF9, 0x94, 0x10)
GRIS   = RGBColor(0x5A, 0x60, 0x6A)
ROJO   = RGBColor(0xB3, 0x26, 0x1E)
VERDE  = RGBColor(0x1B, 0x6E, 0x3C)
NEGRO  = RGBColor(0x1F, 0x24, 0x2C)

HEX_AZUL   = "003AA0"
HEX_NARANJ = "F99410"
HEX_CLARO  = "EEF2F9"
HEX_SUAVE  = "F7F8FA"
HEX_AVISO  = "FFF4E3"

BODY = "Segoe UI"
MONO = "Consolas"

doc = Document()

# ───────────────────────── estilos base ─────────────────────────
st = doc.styles["Normal"]
st.font.name = BODY
st.font.size = Pt(10.5)
st.font.color.rgb = NEGRO
st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15

for sec in doc.sections:
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)


def shade(el, hexcolor):
    el.append(parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls("w"), hexcolor)))


def border(par, edge="bottom", size=12, color=HEX_NARANJ, space=4):
    pPr = par._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    e = OxmlElement("w:" + edge)
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(size))
    e.set(qn("w:space"), str(space))
    e.set(qn("w:color"), color)
    pBdr.append(e)


def par_shade(par, hexcolor):
    shade(par._p.get_or_add_pPr(), hexcolor)


def run(par, text, bold=False, italic=False, size=10.5, color=None, font=BODY, hl=None):
    r = par.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color is not None:
        r.font.color.rgb = color
    if hl:
        shade(r._element.get_or_add_rPr(), hl)
    return r


def p(text="", size=10.5, bold=False, italic=False, color=None, space_after=6,
      space_before=0, align=None, left=0):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    if left:
        par.paragraph_format.left_indent = Cm(left)
    if align:
        par.alignment = align
    if text:
        run(par, text, bold=bold, italic=italic, size=size, color=color)
    return par


def rich(parts, space_after=6, space_before=0, left=0, size=10.5):
    """parts = lista de (texto, dict-de-formato)"""
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    if left:
        par.paragraph_format.left_indent = Cm(left)
    for txt, fmt in parts:
        f = dict(fmt)
        f.setdefault("size", size)
        run(par, txt, **f)
    return par


_primer_h1 = [True]


def h1(text, nro=None):
    par = doc.add_paragraph()
    # Cada capítulo abre página propia: evita encabezados huérfanos al pie y
    # páginas casi vacías por un renglón que se desborda.
    par.paragraph_format.page_break_before = True
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.keep_with_next = True
    if nro:
        run(par, nro + "  ", bold=True, size=17, color=NARANJ)
    run(par, text, bold=True, size=17, color=AZUL)
    border(par, "bottom", size=10, color=HEX_NARANJ, space=3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def h2(text, color=AZUL):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.keep_with_next = True
    run(par, text, bold=True, size=13, color=color)
    return par


def h3(text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(9)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.keep_with_next = True
    run(par, text, bold=True, size=11, color=GRIS)
    return par


def bullet(text_parts, level=0):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(0.75 + 0.6 * level)
    if isinstance(text_parts, str):
        run(par, text_parts)
    else:
        for txt, fmt in text_parts:
            f = dict(fmt)
            f.setdefault("size", 10.5)
            run(par, txt, **f)
    return par


def pasos(items):
    """Lista numerada manual. El estilo 'List Number' de Word comparte una sola
    numeración para todo el documento: el segundo 'Paso a paso' arrancaba en 5.
    Numerando a mano, cada lista arranca en 1."""
    for i, text_parts in enumerate(items, start=1):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.left_indent = Cm(1.05)
        par.paragraph_format.first_line_indent = Cm(-0.6)
        run(par, "{}.	".format(i), bold=True, color=NARANJ)
        if isinstance(text_parts, str):
            run(par, text_parts)
        else:
            for txt, fmt in text_parts:
                f = dict(fmt)
                f.setdefault("size", 10.5)
                run(par, txt, **f)


def callout(titulo, texto, fill=HEX_AVISO, color_titulo=NARANJ, borde=HEX_NARANJ):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), fill)
    cell.text = ""
    par = cell.paragraphs[0]
    par.paragraph_format.space_after = Pt(2)
    run(par, titulo, bold=True, size=10.5, color=color_titulo)
    par2 = cell.add_paragraph()
    par2.paragraph_format.space_after = Pt(0)
    if isinstance(texto, str):
        run(par2, texto, size=10)
    else:
        for txt, fmt in texto:
            f = dict(fmt)
            f.setdefault("size", 10)
            run(par2, txt, **f)
    _cell_margins(cell, 140, 110)
    _tbl_borders(t, borde, left_only=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _cell_margins(cell, lr=110, tb=70):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", tb), ("left", lr), ("bottom", tb), ("right", lr)):
        e = OxmlElement("w:" + side)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def _tbl_borders(tbl, color, left_only=False, size=4):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    sides = ["left"] if left_only else ["top", "left", "bottom", "right", "insideH", "insideV"]
    todos = ["top", "left", "bottom", "right", "insideH", "insideV"]
    for s in todos:
        e = OxmlElement("w:" + s)
        if s in sides:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(24 if left_only else size))
            e.set(qn("w:color"), color)
        else:
            e.set(qn("w:val"), "none")
        e.set(qn("w:space"), "0")
        borders.append(e)
    tblPr.append(borders)


def tabla(headers, filas, widths=None, size=9.5, header_fill=HEX_AZUL, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, htxt in enumerate(headers):
        c = hdr.cells[i]
        if header_fill:
            shade(c._tc.get_or_add_tcPr(), header_fill)
        c.text = ""
        par = c.paragraphs[0]
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        run(par, htxt, bold=True, size=size, color=RGBColor(0xFF, 0xFF, 0xFF))
        _cell_margins(c, 90, 55)
    _repeat_header(hdr)
    for j, fila in enumerate(filas):
        row = t.add_row()
        for i, celda in enumerate(fila):
            c = row.cells[i]
            if zebra and j % 2 == 1:
                shade(c._tc.get_or_add_tcPr(), HEX_SUAVE)
            c.text = ""
            par = c.paragraphs[0]
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.space_before = Pt(0)
            if isinstance(celda, str):
                run(par, celda, size=size)
            else:
                for txt, fmt in celda:
                    f = dict(fmt)
                    f.setdefault("size", size)
                    run(par, txt, **f)
            _cell_margins(c, 90, 55)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    if header_fill is None:          # tabla sin encabezado (ficha de portada)
        t._tbl.remove(t.rows[0]._tr)
    _tbl_borders(t, "D5DAE3", size=4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:tblHeader")
    e.set(qn("w:val"), "true")
    trPr.append(e)


def tecla(txt):
    return (txt, {"font": MONO, "bold": True, "size": 10, "color": AZUL, "hl": HEX_CLARO})


def code_block(lineas, size=9):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), "F4F6F9")
    cell.text = ""
    for i, ln in enumerate(lineas):
        par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        run(par, ln, font=MONO, size=size, color=RGBColor(0x33, 0x39, 0x42))
    _cell_margins(cell, 140, 90)
    _tbl_borders(t, "C9D2E0", left_only=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def ficha(key, titulo, escribe, permiso=None):
    """Cabecera de ficha de una tecla."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), HEX_CLARO)
    cell.text = ""
    par = cell.paragraphs[0]
    par.paragraph_format.space_after = Pt(0)
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.keep_with_next = True
    run(par, key, bold=True, size=15, font=MONO, color=AZUL)
    run(par, "   " + titulo + "     ", bold=True, size=14, color=NEGRO)
    if escribe:
        run(par, " MODIFICA DATOS ", bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF), hl="C9741A")
    if permiso:
        run(par, "   permiso " + permiso, bold=True, size=8.5, color=GRIS)
    _cell_margins(cell, 140, 110)
    _tbl_borders(t, HEX_AZUL, left_only=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def campo(etiqueta, texto):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(4)
    par.paragraph_format.left_indent = Cm(0.3)
    run(par, etiqueta + "   ", bold=True, size=9, color=NARANJ)
    if isinstance(texto, str):
        run(par, texto)
    else:
        for txt, fmt in texto:
            f = dict(fmt)
            f.setdefault("size", 10.5)
            run(par, txt, **f)
    return par


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ═══════════════════════════ PORTADA ═══════════════════════════
for _ in range(4):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(0)
run(par, "BUSLINK", bold=True, size=11, color=NARANJ)
par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(0)
run(par, "Módulo Tráfico · Planilla del día", size=11, color=GRIS)

par = doc.add_paragraph()
par.paragraph_format.space_before = Pt(18)
par.paragraph_format.space_after = Pt(0)
run(par, "Guía de teclas", bold=True, size=34, color=AZUL)
par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(4)
run(par, "F1 a F9", bold=True, size=34, color=NARANJ)
border(par, "bottom", size=18, color=HEX_NARANJ, space=8)

par = doc.add_paragraph()
par.paragraph_format.space_before = Pt(14)
run(par, "Todo lo que hacen los atajos de la planilla de tráfico: qué resuelve cada uno, "
         "cuándo se usa, qué queda registrado y en qué se diferencia del Metrocar.",
    size=12, color=GRIS)

for _ in range(2):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

tabla(
    ["", ""],
    [
        [[("Para quién", {"bold": True})], "Operadores de la mesa de tráfico, diagramadores y supervisión"],
        [[("Alcance", {"bold": True})], "Las 9 teclas de función + Ctrl+F8 + el atajo de tipeo directo"],
        [[("Nivel", {"bold": True})], "Básico–intermedio. No hace falta saber nada de la base de datos"],
        [[("Sistema", {"bold": True})], "Buslink (Blazor) — con las diferencias marcadas contra Metrocar (FoxPro)"],
        [[("Versión", {"bold": True})], "Agosto 2026"],
    ],
    widths=[3.6, 12.4], size=10, header_fill=None, zebra=True,
)


# ═══════════════════════════ CONTENIDO ═══════════════════════════
h1("Contenido")

contenido = [
    ("1.", "Antes de las teclas: cómo leer la planilla", "Las tres unidades de un viaje, los códigos y los permisos"),
    ("2.", "Mapa rápido", "Las once teclas en una tabla"),
    ("3.", "Familia 1 — Ver la planilla", "F1 · F3 · F5"),
    ("4.", "Familia 2 — Anotar sobre el servicio", "F2 · F4"),
    ("5.", "Familia 3 — Cambiar la unidad prevista", "F6 · F7 · F8 · F9 · Ctrl+F8 · tipeo directo"),
    ("6.", "Las cinco teclas de unidad, comparadas", "Cuál usar en cada situación"),
    ("7.", "Modo diagramador y modo operador", "Por qué la misma tecla hace cosas distintas"),
    ("8.", "Lo que las teclas F NO hacen", "Para no buscarlas donde no están"),
    ("9.", "Preguntas frecuentes", ""),
    ("A.", "Anexo — Qué cambia en la base con cada tecla", "Para perfil técnico"),
    ("B.", "Anexo — Diferencias entre Metrocar y Buslink", ""),
    ("C.", "Anexo — Qué está activo hoy y qué espera al día D", ""),
]
for nro, tit, sub in contenido:
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(5)
    par.paragraph_format.left_indent = Cm(0.2)
    run(par, nro + "  ", bold=True, size=11, color=NARANJ)
    run(par, tit, bold=True, size=11, color=AZUL)
    if sub:
        run(par, "   " + sub, size=10, color=GRIS)

callout(
    "Cómo usar esta guía",
    [("Si ya sabés lo que hacen las teclas y solo querés acordarte de cuál era, andá directo al ",
      {}), ("Mapa rápido", {"bold": True}), (" (capítulo 2). Si estás aprendiendo, leé el capítulo 1 antes que nada: "
      "las teclas de unidad no se entienden sin las ", {}), ("tres unidades de un viaje", {"bold": True}), (".", {})],
    fill=HEX_CLARO, color_titulo=AZUL, borde=HEX_AZUL,
)


# ═══════════════════════ 1. CONCEPTOS ═══════════════════════
h1("Antes de las teclas: cómo leer la planilla", "1")

p("Cinco de las once teclas hacen lo mismo: cambiar la unidad prevista de un servicio. "
  "Si no está clara la diferencia entre «prevista» y «asignada», esas cinco teclas parecen "
  "arbitrarias. Este capítulo es corto y evita casi todos los errores.")

h2("1.1  Un viaje tiene tres unidades, no una")

p("En la grilla hay tres columnas que hablan de unidades. No son lo mismo ni las escribe la "
  "misma persona:")

tabla(
    ["Columna", "Se llama", "Qué es", "Quién la carga"],
    [
        [[("U/Pr", {"bold": True, "font": MONO, "color": AZUL})],
         "Unidad Programada",
         "La unidad que el diagramador planificó para ese servicio, normalmente la noche anterior.",
         "El diagramador"],
        [[("U/Cb", {"bold": True, "font": MONO, "color": AZUL})],
         "Unidad del Cronograma",
         "La unidad prevista vigente ahora. Arranca igual a U/Pr y va cambiando durante el día "
         "según lo que pase.",
         "Diagramador y operador"],
        [[("U/As", {"bold": True, "font": MONO, "color": AZUL})],
         "Unidad Asignada",
         "La unidad real que salió a la calle, con su chofer. Es la que factura y la que se libera "
         "al terminar.",
         "El operador, con Asig U/P"],
    ],
    widths=[1.8, 3.3, 7.2, 3.7], size=9.5,
)

callout(
    "La frase que hay que retener",
    [("Las teclas ", {}), ("F6 a F9 y Ctrl+F8 solo tocan U/Cb", {"bold": True}),
     (" (y U/Pr si sos diagramador). ", {}),
     ("Ninguna asigna la unidad ni el chofer.", {"bold": True}),
     (" La asignación real es el botón ", {}), ("Asig U/P", {"bold": True}),
     (" de la barra, y es otra cosa completamente distinta: ahí sí se compromete un vehículo, "
      "se toma el odómetro y el servicio pasa a EN CURSO.", {})],
)

h3("El ciclo completo de un servicio")

code_block([
    " Nace la reserva ─► Diagramación  ─► Ajuste del día  ─► Asignación   ─► Servicio",
    "   (Reservas)       (arma el plan)   (algo cambió)      (Asig U/P)      en la calle",
    "",
    " U/Cb = S/C         U/Pr = NT0049    U/Cb = NT0051      U/As = NT0051   CURSO ─► FIN",
    "                    U/Cb = NT0049    (F6-F9, Ctrl+F8)   + chofer",
], size=8.5)

p("Las teclas de esta guía viven en las dos etapas del medio. Lo de la derecha —asignar, "
  "liberar, finalizar— se hace con botones, no con teclas.")

h2("1.2  Los códigos que vas a ver en U/Pr y U/Cb")

tabla(
    ["Código", "Significa", "Cómo se carga"],
    [
        [[("S/C", {"font": MONO, "bold": True})], "Sin cronograma. El servicio no tiene todavía ninguna unidad prevista.",
         [tecla("F6")]],
        [[("NORTUR", {"font": MONO, "bold": True})], "Lo cubre la empresa, pero todavía no se definió con qué interno.",
         [tecla("F7")]],
        [[("NT0049", {"font": MONO, "bold": True})], "Interno 49 de la flota propia. Dos letras de prefijo + cuatro dígitos.",
         [tecla("F9")]],
        [[("TT0012", {"font": MONO, "bold": True})], "Interno 12 de un transportista contratado (fletero) con prefijo TT.",
         [tecla("F8")]],
        [[("TEDESCHI", {"font": MONO, "bold": True})], "Una empresa contratada entera, sin especificar interno. Solo la "
         "usan los fleteros que no diagraman unidad por unidad.",
         [tecla("F8"), (" en modo diagramador", {})]],
    ],
    widths=[2.6, 9.4, 4.0], size=9.5,
)

p("El número siempre va con ceros a la izquierda hasta cuatro dígitos: el interno 49 es "
  "NT0049, no NT49. Buslink lo completa solo; no hace falta tipear los ceros.", size=10)

h2("1.3  Quién puede hacer qué: las letras de permiso")

p("El sistema no muestra las mismas teclas a todo el mundo. Tres letras del permiso del "
  "usuario deciden qué ves:")

tabla(
    ["Letra", "Qué habilita", "Efecto sobre las teclas"],
    [
        [[("T", {"bold": True, "font": MONO, "size": 11, "color": AZUL})],
         "Tráfico — operador de la mesa",
         "Puede cambiar la unidad prevista, pero solo la vigente (U/Cb) y siempre con un motivo."],
        [[("D", {"bold": True, "font": MONO, "size": 11, "color": AZUL})],
         "Diagramador",
         "Trabaja el plan: cambia U/Pr y U/Cb juntas, sin motivo, y puede hacer cambios masivos. "
         "Además habilita F5."],
        [[("C", {"bold": True, "font": MONO, "size": 11, "color": AZUL})],
         "Avisos de chequeo",
         "Habilita F4 y hace que suenen las alarmas de los servicios que están por salir."],
    ],
    widths=[1.6, 4.6, 9.8], size=9.5,
)

bullet([("Si tenés ", {}), ("D", {"bold": True, "font": MONO}), (" y ", {}), ("T", {"bold": True, "font": MONO}),
        (" a la vez, mandan las reglas de ", {}), ("D", {"bold": True, "font": MONO}), (": entrás como diagramador.", {})])
bullet([("La ventana de ayuda del ", {}), ("F1", {"font": MONO, "bold": True}),
        (" oculta las teclas que no podés usar. Si un compañero ve una tecla que vos no, "
         "es una diferencia de permisos, no un error.", {})])
bullet([("Los permisos los administra Supervisión desde ", {}), ("Usuarios y Permisos", {"bold": True}),
        ("; no se cambian desde la planilla.", {})])


# ═══════════════════════ 2. MAPA RAPIDO ═══════════════════════
h1("Mapa rápido", "2")

p("Las once teclas de la planilla. La columna «Modifica» marca las que dejan un cambio en la "
  "reserva; el resto solo cambia lo que ves en pantalla.")

tabla(
    ["Tecla", "Qué hace", "Familia", "Modifica", "Permiso"],
    [
        [[tecla("F1")], "Abre la ayuda de atajos", "Ver la planilla", "—", "todos"],
        [[tecla("F2")], "Carga una novedad en el libro de guardia", "Anotar", "sí", "todos"],
        [[tecla("F3")], "Actualiza la planilla ahora mismo", "Ver la planilla", "—", "todos"],
        [[tecla("F4")], "Fija la hora del aviso de ese servicio", "Anotar", "sí", "C"],
        [[tecla("F5")], "Cambia el rango de fechas de trabajo", "Ver la planilla", "—", "D"],
        [[tecla("F6")], "Deja el servicio sin unidad prevista (S/C)", "Cambiar unidad", "sí", "T o D"],
        [[tecla("F7")], "Pasa el servicio a la empresa (NORTUR)", "Cambiar unidad", "sí", "T o D"],
        [[tecla("F8")], "Pasa el servicio a un fletero", "Cambiar unidad", "sí", "T o D"],
        [[tecla("F9")], "Pasa el servicio a un interno propio", "Cambiar unidad", "sí", "T o D"],
        [[tecla("Ctrl+F8")], "Vuelve a la unidad que puso el diagramador", "Cambiar unidad", "sí", "T o D"],
        [[tecla("A-Z 0-9")], "Tipear abre el cambio de unidad ya filtrado", "Cambiar unidad", "sí", "T o D"],
    ],
    widths=[2.2, 6.6, 3.2, 1.9, 2.1], size=9.5,
)

callout(
    "Dos teclas que el navegador también usa",
    [("F1", {"font": MONO, "bold": True}), (" abre normalmente la ayuda del navegador y ", {}),
     ("F3", {"font": MONO, "bold": True}), (" el «buscar siguiente». Buslink se las queda para sí "
     "mientras el cursor está sobre la grilla. En cambio ", {}),
     ("F5", {"font": MONO, "bold": True}), (" (recargar) y ", {}), ("F11", {"font": MONO, "bold": True}),
     (" (pantalla completa) se le siguen dejando al navegador cuando no estás parado en la planilla.", {})],
)

h2("Cómo se dispara cada una")

p("Todas las teclas actúan sobre la fila donde está el cursor y solo cuando el foco está en la "
  "grilla. Si acabás de tipear en un filtro, hacé un clic en cualquier fila primero.")

bullet([("Con la tecla", {"bold": True}), (" — el camino rápido, el mismo del Metrocar.", {})])
bullet([("Con el clic derecho sobre la fila", {"bold": True}),
        (" — el menú contextual tiene los mismos ítems, con el atajo escrito al costado. "
         "Es el camino recomendable mientras estás aprendiendo.", {})])
bullet([("Desde los botones de la barra", {"bold": True}),
        (" — Avisos, Novedad y F1 están también arriba, a la vista.", {})])


# ═══════════════════════ 3. FAMILIA VER ═══════════════════════
h1("Familia 1 — Ver la planilla", "3")

p("Tres teclas que no tocan ningún dato: solo cambian lo que estás mirando. Se pueden apretar "
  "sin miedo.")

# ── F1 ──
ficha("F1", "Esta ayuda", escribe=False)
campo("QUÉ HACE", "Abre la lista de atajos de la planilla, agrupada por familia. Cada tecla "
      "aparece en una línea y se despliega al tocarla, con qué hace, cuándo conviene usarla y un "
      "distintivo naranja si modifica la reserva.")
campo("CUÁNDO", "Cuando no te acordás cuál era la tecla, o cuando entra alguien nuevo a la mesa.")
campo("OJO CON", [("La ayuda muestra solo las teclas que ", {}), ("vos", {"italic": True}),
                  (" podés usar. Si no tenés el permiso D no vas a ver el F5, y si no tenés el "
                   "C no vas a ver el F4. Es a propósito.", {})])
campo("DIFERENCIA", "En el Metrocar la ventana «Zoon» listaba las teclas pero no decía qué hacía "
      "ninguna: había que saberlo de antes. Acá cada una está explicada.")

# ── F3 ──
ficha("F3", "Actualizar ahora", escribe=False)
campo("QUÉ HACE", "Vuelve a traer los servicios del día desde la base, en el momento.")
campo("CUÁNDO", [("Casi nunca hace falta. ", {"bold": True}),
                 ("La planilla se actualiza sola cada 60 segundos, resalta las filas que "
                  "cambiaron y muestra la hora de la última actualización arriba. F3 está para "
                  "forzarla cuando querés confirmar algo al instante: por ejemplo, "
                  "después de que un compañero te avisó por teléfono que acaba de asignar.", {})])
campo("OJO CON", "No confundir con F5 del navegador, que recarga toda la página y te hace perder "
      "los filtros. F3 refresca solo los datos.")
campo("DIFERENCIA", "En el Metrocar era la única forma de ver algo nuevo: si no apretabas Refresh, "
      "la pantalla se quedaba congelada con lo que había cuando la abriste.")

# ── F5 ──
ficha("F5", "Cambiar el rango de fechas de trabajo", escribe=False, permiso="D")
campo("QUÉ HACE", "Deja ver varios días juntos en la misma planilla, en vez de un día por vez.")
campo("CUÁNDO", "Cuando estás diagramando y necesitás mirar la semana completa para repartir las "
      "unidades sin superponerlas.")
campo("OJO CON", [("Mientras haya un rango puesto, las flechas ", {}),
                  ("« »", {"font": MONO, "bold": True}),
                  (" de día anterior y día siguiente quedan desactivadas: estás mirando un "
                   "período, no un día. Para volver al modo normal hay que limpiar el rango.", {})])
campo("PERMISO", [("Es la única tecla de esta familia que pide permiso. Sin la letra ", {}),
                  ("D", {"font": MONO, "bold": True}), (" no aparece ni en la ayuda.", {})])


# ═══════════════════════ 4. FAMILIA ANOTAR ═══════════════════════
h1("Familia 2 — Anotar sobre el servicio", "4")

p("Dos teclas que dejan algo escrito sobre una reserva: una nota para el turno siguiente (F2) "
  "y un recordatorio para vos mismo (F4).")

# ── F2 ──
ficha("F2", "Cargar una novedad", escribe=True)
campo("QUÉ HACE", "Escribe una anotación en el libro de guardia. Si la cargás parado sobre una "
      "fila, la novedad queda colgada de esa reserva y se ve después desde el viaje. Si la abrís "
      "desde el botón Novedad de la barra, queda como novedad suelta del día.")
campo("CUÁNDO", "Pasó algo que el turno siguiente tiene que saber y no entra en ningún campo de la "
      "reserva.")

h3("Ejemplos reales del libro")
code_block([
    "int: 8   dom: AD255RA   chof: OJEDA HUGO ORLANDO",
    "  CHOFER INFORMA QUE TIENE LA PIERNA INFLAMADA, SE LE CUBRE EL PRIMER SERVICIO,",
    "  DEJA UNIDAD EN FLORES Y VA A LABORAL",
    "",
    "int: 67  dom: AF021PY   chof: LIMOLE CARLOS FERNANDO",
    "  SERVICIO 8:10 OBELISCO-TORRE YPF, SE COMPLETO EN CATEDRAL DEJANDO PASAJEROS ABAJO",
])

h3("Paso a paso")
pasos([
    [("Parate en la fila del servicio y apretá ", {}), tecla("F2"),
     (". Se abre el libro con el formulario de alta ya desplegado.", {})],
    "El asunto viene precargado: el nombre del cliente si la novedad es de una reserva, el "
    "nombre de la empresa si es suelta. Se puede cambiar.",
    "Escribí el mensaje. Contá lo que el turno siguiente necesita saber, no lo que ya está "
    "en la reserva.",
    "Guardar. La novedad queda con tu usuario y la hora, y aparece arriba de la lista.",
])

campo("QUÉ QUEDA", "Fecha y hora, asunto, mensaje, tu usuario y —si corresponde— el número de "
      "reserva. Nada más: no hay prioridad, ni destinatario, ni teléfono, aunque el Metrocar "
      "muestre esos campos.")
campo("VOLUMEN", "Unas cinco novedades por día, mitad ligadas a una reserva y mitad sueltas.")
campo("OJO CON", [("El aviso por correo al cliente ", {"bold": True}),
                  ("no se manda desde Buslink", {"bold": True}),
                  (": ese circuito sigue saliendo del Metrocar. Cargar la novedad acá no le "
                   "avisa a nadie de afuera.", {})])

# ── F4 ──
ficha("F4", "Poner hora de aviso", escribe=True, permiso="C")
campo("QUÉ HACE", "Fija a qué hora el sistema te va a golpear la puerta por ese servicio, con un "
      "cartel y un sonido.")

h3("Los dos niveles de aviso")
p("El sistema avisa siempre, con o sin F4. Lo que hace la tecla es pisar el aviso general para "
  "un servicio puntual.")
tabla(
    ["Nivel", "De dónde sale la hora", "Alcance"],
    [
        ["Automático", "10 minutos antes de la hora de salida", "Todos los servicios del día"],
        [[("Manual — ", {}), tecla("F4")], "La hora exacta que vos elijas",
         "Solo ese servicio. Pisa al automático"],
    ],
    widths=[3.4, 6.6, 6.0], size=9.5,
)

campo("CUÁNDO", "Querés que te avise antes de lo habitual porque el servicio es delicado, el "
      "chofer es nuevo, el pasajero es un cliente sensible o la unidad viene de taller.")

h3("Paso a paso")
pasos([
    [("Parate en la fila y apretá ", {}), tecla("F4"), (".", {})],
    "El diálogo viene con la hora del aviso automático ya cargada, y arriba muestra el "
    "servicio para que confirmes que es el que querías.",
    "Elegí uno de los botones rápidos —30, 45 o 60 minutos antes— o poné una hora exacta.",
    "Debajo se lee siempre en texto a qué hora va a sonar y cuánta antelación es. Si esa "
    "hora ya pasó, te lo avisa ahí mismo.",
    "Grabar.",
])

campo("QUÉ QUEDA", [("La hora queda en la columna ", {}), ("H.Avi", {"font": MONO, "bold": True}),
                    (" de la planilla, así que se ve de un vistazo qué servicios tienen aviso "
                     "propio.", {})])
campo("CUÁNTO PIDE LA MESA", "De cada diez avisos manuales, casi diez piden entre 16 y 60 minutos. "
      "El pedido más común, lejos, es una hora antes. Por eso los botones rápidos son 30, 45 y 60.")

callout(
    "«Volver al aviso automático» no apaga el aviso",
    [("La casilla que en el Metrocar decía ", {}), ("No Avisar", {"bold": True, "font": MONO}),
     (" no deja de avisar: borra tu hora manual y el servicio vuelve al aviso general de 10 "
      "minutos antes. Por eso en Buslink se llama ", {}),
     ("Volver al aviso automático", {"bold": True}),
     (", que es lo que realmente hace. Hoy no existe forma de dejar un servicio sin ningún aviso.", {})],
)

h3("Cómo suena la alarma")
bullet("Cada minuto el sistema mira si algún servicio tiene que avisar. Si toca, aparece un "
       "cartel con la lista de los servicios que están por salir, no un número suelto.")
bullet("Suena una campana. Si el navegador todavía no te dejó reproducir sonido, alcanza con "
       "hacer un clic en cualquier parte de la página una vez.")
bullet("El cartel se cierra solo a los 20 segundos, salvo que tengas el mouse encima: "
       "mientras lo estés leyendo no se va.")
bullet([("El interruptor ", {}), ("Avisos", {"bold": True}),
        (" de la barra prende y apaga todo el motor de alarmas. Solo aparece si tenés el permiso ", {}),
        ("C", {"font": MONO, "bold": True}), (".", {})])
bullet([("Si un aviso cayó justo cuando no estabas mirando, Buslink te lo muestra igual hasta "
         "cinco minutos después, marcado como ", {}), ("atrasado", {"bold": True, "color": NARANJ}),
        (" en naranja, para que sepas que se te pasó.", {})])

callout(
    "Una corrección que Buslink trae de fábrica",
    [("Los servicios que salen pasada la medianoche —los transfers de las 00:01, más de mil por "
      "año— tienen su aviso a las 23:51 del día anterior. El Metrocar los hacía sonar 24 horas "
      "tarde, cuando el servicio ya había salido. En Buslink suenan cuando corresponde.", {})],
    fill="EAF4EC", color_titulo=VERDE, borde="1B6E3C",
)


# ═══════════════════════ 5. FAMILIA UNIDAD ═══════════════════════
h1("Familia 3 — Cambiar la unidad prevista", "5")

p("Acá están cinco de las once teclas, y son las que más se usan: más de ciento cincuenta "
  "cambios por día. Todas hacen lo mismo —mover la unidad prevista de un servicio— y se "
  "diferencian solo en a qué la mueven.")

callout(
    "Antes de seguir",
    [("Ninguna de estas teclas asigna nada. ", {"bold": True}),
     ("Cambiar el cronograma es mover el plan; sacar la unidad a la calle con su chofer es "
      "el botón ", {}), ("Asig U/P", {"bold": True}),
     (". Se puede cambiar el cronograma diez veces en el día sin que ningún vehículo se entere.", {})],
)

# ── F6 ──
ficha("F6", "Dejar sin unidad (S/C)", escribe=True, permiso="T o D")
campo("QUÉ HACE", [("Le saca la unidad prevista al servicio y lo deja en ", {}),
                   ("S/C", {"font": MONO, "bold": True}), (", sin cronograma.", {})])
campo("CUÁNDO", "Se cayó la unidad que iba a cubrirlo y todavía no sabés con qué reemplazarla. "
      "Es preferible dejarlo en S/C que dejar puesta una unidad que ya sabés que no va: así el "
      "servicio queda visible como pendiente de resolver.")
campo("OJO CON", "S/C no cancela ni suspende el servicio. El servicio sigue en pie y hay que "
      "cubrirlo; lo único que decís es que todavía no sabés con qué.")

# ── F7 ──
ficha("F7", "Pasar a la empresa", escribe=True, permiso="T o D")
campo("QUÉ HACE", [("Marca el servicio como cubierto por ", {}), ("NORTUR", {"font": MONO, "bold": True}),
                   (", sin especificar el interno.", {})])
campo("CUÁNDO", "Ya sabés que lo cubre la empresa con flota propia, pero todavía no está definido "
      "cuál unidad. Es un paso intermedio entre S/C y el interno concreto.")
campo("OJO CON", "Queda a medio camino: no se puede seguir así hasta el momento del servicio. "
      "Antes de la salida hay que bajarlo a un interno con F9.")

# ── F8 ──
ficha("F8", "Pasar a un fletero", escribe=True, permiso="T o D")
campo("QUÉ HACE", "Cambia la unidad prevista a la de un transportista contratado. Elegís primero "
      "la empresa y después el interno dentro de esa empresa; el código queda armado como TT0012.")
campo("CUÁNDO", "Ese servicio no lo va a cubrir la flota propia, sea porque no hay unidad "
      "disponible o porque el cliente o la zona lo tienen previsto así.")

h3("Paso a paso")
pasos([
    [("Parate en la fila y apretá ", {}), tecla("F8"), (".", {})],
    "Elegí el fletero de la lista. Solo aparecen los que están vigentes.",
    "Elegí el interno. La lista trae únicamente las unidades activas de ese fletero.",
    [("Si sos operador, cargá el ", {}), ("motivo", {"bold": True}),
     (": es obligatorio y queda en el historial del viaje.", {})],
    "Confirmá. Antes de grabar, el diálogo te muestra en texto cómo va a quedar el código.",
])

campo("OJO CON", [("Esto ", {}), ("no", {"bold": True, "italic": True}),
                  (" asigna chofer ni saca la unidad a la calle. La asignación real es ", {}),
                  ("Asig U/P", {"bold": True}), (".", {})])
campo("CASO PARTICULAR", "Algunos fleteros no se diagraman unidad por unidad: se les manda el "
      "servicio a la empresa entera y ellos deciden con qué lo cubren. En esos casos el "
      "cronograma queda con el nombre de la empresa —por ejemplo TEDESCHI— en vez de un código de "
      "interno. Buslink lo resuelve solo; no hay que elegir nada distinto.")

# ── F9 ──
ficha("F9", "Pasar a flota propia", escribe=True, permiso="T o D")
campo("QUÉ HACE", [("Cambia la unidad prevista a un interno propio. Queda como ", {}),
                   ("NT0049", {"font": MONO, "bold": True}), (".", {})])
campo("CUÁNDO", "Es el camino normal: la enorme mayoría de los servicios los cubre una unidad de "
      "la casa, y esta es la tecla que se usa para decir cuál.")
campo("ATAJO", [("Si ya sabés el número de interno, no hace falta ni apretar F9: tipeá el número "
                 "directamente sobre la grilla y el diálogo se abre ya posicionado. Ver ", {}),
                ("Tipeo directo", {"bold": True}), (" más abajo.", {})])

# ── Ctrl+F8 ──
ficha("Ctrl+F8", "Copiar la del diagramador", escribe=True, permiso="T o D")
campo("QUÉ HACE", [("Pisa la unidad vigente (", {}), ("U/Cb", {"font": MONO, "bold": True}),
                   (") con la que había planificado el diagramador (", {}),
                   ("U/Pr", {"font": MONO, "bold": True}), ("), sin abrir ningún diálogo.", {})])
campo("CUÁNDO", "Se probó un cambio sobre la marcha, no prosperó, y hay que volver al plan "
      "original sin tener que acordarse cuál era.")
campo("OJO CON", [("Es la única de la familia que no pide confirmación: se aplica en el acto. ", {}),
                  ("El ítem del menú se deshabilita solo", {"bold": True}),
                  (" cuando no hay nada que copiar, o sea cuando U/Pr está vacía o ya coincide con "
                   "U/Cb, así que no se puede apretar por error sobre una fila donde no aplica.", {})])

# ── tipeo ──
ficha("A-Z  0-9", "Tipeo directo", escribe=True, permiso="T o D")
campo("QUÉ HACE", "Tipear cualquier letra o número con el cursor sobre la grilla abre el cambio de "
      "unidad, ya posicionado en el interno que empieza así.")
campo("CUÁNDO", "Es el atajo de los que vienen del Metrocar y tienen los internos en la cabeza: se "
      "tipea el número directo, sin pasar por ninguna tecla de función ni por el menú.")
campo("OJO CON", [("Como cualquier tecla alfanumérica lo dispara, conviene tener presente que el "
                   "foco está en la grilla. Si querías escribir en un filtro y el diálogo se abre "
                   "solo, cancelá con ", {}), ("Esc", {"font": MONO, "bold": True}),
                  (" y hacé clic en el campo antes de tipear.", {})])
campo("DIFERENCIA", [("En el Metrocar la tecla ", {}), ("Enter", {"font": MONO, "bold": True}),
                     (" también abría el cambio de unidad. En Buslink ", {}),
                     ("Enter abre el Zoom del Viaje", {"bold": True}),
                     (", que es lo que la mesa ya tiene aprendido acá.", {})])

# ── masivo ──
h2("El cambio masivo (solo diagramador)")

p("Cuando el diagramador cambia la unidad prevista, el diálogo ofrece además aplicar el cambio a "
  "todos los servicios del día que compartan la misma unidad prevista y que todavía no tengan "
  "nadie asignado. Es lo que se usa cuando una unidad entra a taller y hay que repartir su día "
  "entero.")

bullet([("Antes de grabar, Buslink te muestra ", {}), ("exactamente a cuántas reservas alcanza", {"bold": True}),
        (" y cuáles son. Si no hay otras filas además de la actual, te lo dice con todas las letras.", {})])
bullet("Los servicios que ya tienen unidad asignada quedan afuera: esos ya están comprometidos y "
       "se cambian de a uno.")
bullet("Todo el cambio se graba junto: o entran todas las filas o no entra ninguna. No puede "
       "quedar el tablero a medio cambiar.")


# ═══════════════════════ 6. COMPARACION ═══════════════════════
h1("Las cinco teclas de unidad, comparadas", "6")

p("La misma información de arriba, ordenada para decidir rápido.")

tabla(
    ["Situación", "Tecla", "Queda como"],
    [
        ["Se cayó la unidad y no sé con qué reemplazarla", [tecla("F6")], [("S/C", {"font": MONO})]],
        ["Lo cubre la empresa pero no sé con qué interno", [tecla("F7")], [("NORTUR", {"font": MONO})]],
        ["Lo cubre el interno propio 49", [tecla("F9"), (" o tipear ", {}), ("49", {"font": MONO})],
         [("NT0049", {"font": MONO})]],
        ["Lo cubre el interno 12 de un contratado", [tecla("F8")], [("TT0012", {"font": MONO})]],
        ["Lo cubre una empresa contratada, ella decide la unidad", [tecla("F8")],
         [("TEDESCHI", {"font": MONO})]],
        ["Probé un cambio y quiero volver al plan del diagramador", [tecla("Ctrl+F8")],
         "lo que diga U/Pr"],
        ["Una unidad entra a taller y hay que repartir todo su día",
         [tecla("F9"), (" + masivo", {})], "según elijas"],
    ],
    widths=[8.4, 4.2, 3.4], size=9.5,
)

h2("Lo que comparten todas")

bullet("Actúan sobre la fila donde está el cursor.")
bullet([("Si el servicio es parte de una ", {}), ("ruta con varios tramos", {"bold": True}),
        (", el cambio se aplica a todos los tramos de la ruta, no solo al que estás mirando. "
         "Tiene sentido: es la misma unidad la que hace todo el recorrido.", {})])
bullet([("Todas ", {}), ("borran el chequeo", {"bold": True}),
        (" del servicio. Si cambió la unidad prevista, el chequeo que se había hecho sobre la "
         "anterior ya no vale y hay que rehacerlo.", {})])
bullet("Ninguna toca el vehículo, ni el odómetro, ni el chofer, ni el estado del servicio.")


# ═══════════════════════ 7. MODOS ═══════════════════════
h1("Modo diagramador y modo operador", "7")

p("Las teclas de unidad son las mismas para todos, pero no hacen lo mismo según quién las "
  "aprieta. El sistema elige el modo solo, por el permiso del usuario, y lo avisa en el título "
  "del diálogo.")

tabla(
    ["", "Modo Diagramador", "Modo Operador"],
    [
        [[("Permiso", {"bold": True})], [("D", {"font": MONO, "bold": True})],
         [("T", {"font": MONO, "bold": True})]],
        [[("Título del diálogo", {"bold": True})], "Diagrama", "Cambio de Unidad"],
        [[("Qué escribe", {"bold": True})],
         [("U/Pr y U/Cb", {"bold": True}), (" — mueve el plan y la unidad vigente juntas", {})],
         [("Solo U/Cb", {"bold": True}), (" — respeta lo que planeó el diagramador", {})]],
        [[("Motivo", {"bold": True})], "Deshabilitado. No aplica: está armando el plan.",
         [("Obligatorio", {"bold": True}), (". Sin motivo no graba.", {})]],
        [[("Historial", {"bold": True})], "No deja registro: armar el plan no es una excepción.",
         "Cada cambio queda en el historial del viaje con el motivo y tu usuario."],
        [[("Masivo", {"bold": True})], "Disponible", "No disponible"],
    ],
    widths=[3.4, 6.3, 6.3], size=9.5,
)

h2("Por qué la diferencia")

p("El diagramador arma el tablero antes de que empiece el día: todo lo que hace es «el plan», "
  "no hay nada que justificar. El operador, en cambio, cambia una unidad porque algo pasó "
  "durante el día —se rompió, llegó tarde, el chofer no vino— y ese algo hay que poder "
  "reconstruirlo después. De ahí el motivo obligatorio y el historial.")

p("También por eso el operador no toca U/Pr: la unidad programada queda como testimonio de lo "
  "que estaba previsto. Comparar U/Pr con U/Cb al final del día muestra cuánto se desvió la "
  "operación del plan.")

callout(
    "Si tenés los dos permisos",
    [("Mandan las reglas del diagramador. Un usuario con ", {}), ("D", {"font": MONO, "bold": True}),
     (" y ", {}), ("T", {"font": MONO, "bold": True}),
     (" entra siempre en modo diagramador: sin motivo, moviendo las dos columnas y con el masivo "
      "habilitado. Es importante saberlo, porque significa que esos cambios ", {}),
     ("no van a quedar en el historial del viaje", {"bold": True}), (".", {})],
)


# ═══════════════════════ 8. LO QUE NO HACEN ═══════════════════════
h1("Lo que las teclas F NO hacen", "8")

p("Media hora de búsqueda se ahorra sabiendo dónde no está lo que buscás. Estas operaciones "
  "existen, pero no cuelgan de ninguna tecla de función.")

tabla(
    ["Lo que querés hacer", "Dónde está"],
    [
        ["Asignar la unidad y el chofer reales, sacar el servicio a la calle",
         [("Botón ", {}), ("Asig U/P", {"bold": True}), (" de la barra", {})]],
        ["Cambiar la unidad de un servicio que ya está asignado",
         [("Botón ", {}), ("Reasignar", {"bold": True}), (" de la barra", {})]],
        ["Cerrar un servicio que terminó y liberar la unidad",
         [("Botón ", {}), ("Liberar", {"bold": True}), (" de la barra", {})]],
        ["Chequear que la unidad y el chofer están listos",
         [("Botón ", {}), ("Chequeo", {"bold": True}), (" de la barra", {})]],
        ["Cancelar una reserva",
         [("Zoom del Viaje ", {}), ("(Enter sobre la fila)", {"italic": True})]],
        ["Modificar los datos de la reserva: horario, recorrido, pasajeros",
         [("Zoom del Viaje ", {}), ("(Enter sobre la fila)", {"italic": True})]],
        ["Ver el historial de todo lo que le pasó a un viaje",
         "Clic derecho sobre la fila → Historial del viaje"],
        ["Ver dónde está la unidad ahora mismo",
         "Clic derecho sobre la fila → GPS"],
        ["Mandarle un correo al cliente avisando una novedad",
         [("Sigue saliendo del Metrocar. ", {"bold": True}), ("No está en Buslink.", {})]],
    ],
    widths=[8.6, 7.4], size=9.5,
)

# ═══════════════════════ 9. FAQ ═══════════════════════
h1("Preguntas frecuentes", "9")

faq = [
    ("Aprieto una tecla y no pasa nada.",
     "Casi siempre es el foco: la grilla tiene que estar activa. Hacé un clic sobre cualquier "
     "fila y probá de nuevo. Si sigue sin pasar nada, fijate en la ayuda del F1 si esa tecla "
     "aparece: puede ser que tu usuario no tenga el permiso."),
    ("Cambié el cronograma y la unidad sigue apareciendo libre en el panel de Buses.",
     "Es correcto. Cambiar el cronograma no compromete ningún vehículo: la unidad recién queda "
     "tomada cuando la asignás con Asig U/P."),
    ("¿Por qué me pide un motivo si mi compañero no?",
     "Tu compañero es diagramador y vos operador. El operador cambia la unidad porque pasó algo "
     "durante el día y ese algo queda registrado; el diagramador arma el plan y no tiene nada que "
     "justificar. Ver el capítulo 7."),
    ("Cargué la hora de aviso y no sonó nada.",
     "Tres motivos posibles: el interruptor Avisos de la barra está apagado, tu usuario no tiene "
     "el permiso C, o la hora que pusiste ya había pasado cuando grabaste. El diálogo del F4 "
     "avisa este último caso mientras cargás."),
    ("Puse el aviso y después me arrepentí. ¿Cómo lo saco?",
     "Volvés a abrir el F4 sobre esa fila y marcás «Volver al aviso automático». No queda sin "
     "aviso: vuelve al general de diez minutos antes."),
    ("Cambié la unidad en el servicio equivocado.",
     "Volvé a apretar la tecla sobre esa fila y ponele la unidad que tenía. Si era la que había "
     "planificado el diagramador, Ctrl+F8 te la devuelve sin tener que acordarte cuál era."),
    ("El servicio es una ruta con tres tramos. ¿Tengo que cambiar los tres?",
     "No. El cambio se aplica a los tres tramos automáticamente: es la misma unidad la que hace "
     "todo el recorrido."),
    ("¿Puedo dejar un servicio en NORTUR hasta la hora de salida?",
     "Mejor no. NORTUR dice «lo cubre la empresa» sin decir con qué, y sirve como paso "
     "intermedio. Antes de la salida hay que bajarlo a un interno concreto con F9."),
    ("Aprieto F5 y se me recarga toda la página.",
     "Estás fuera de la grilla. Con el cursor sobre la planilla, F5 abre el rango de fechas —si "
     "tenés permiso de diagramador—; fuera de ella, la tecla es del navegador."),
    ("¿Dónde veo las novedades que cargó el turno anterior?",
     "El botón Novedad de la barra abre el libro con las de los últimos siete días. Las que "
     "cuelgan de una reserva se ven además desde la fila, con el clic derecho."),
]
for preg, resp in faq:
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(9)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.keep_with_next = True
    run(par, preg, bold=True, size=10.5, color=AZUL)
    p(resp, space_after=2, left=0.3)


# ═══════════════════════ ANEXO A ═══════════════════════
h1("Anexo A — Qué cambia en la base con cada tecla", "A")

p("Para perfil técnico y para supervisión. Nombres reales de las columnas de la réplica.", size=10, color=GRIS)

tabla(
    ["Tecla", "Tabla", "Qué escribe", "Historial"],
    [
        [[tecla("F1")], "—", "nada", "—"],
        [[tecla("F3")], "—", "nada (solo relee)", "—"],
        [[tecla("F5")], "—", "nada (solo filtro de pantalla)", "—"],
        [[tecla("F2")], [("libro_novedad", {"font": MONO})],
         [("f_carga", {"font": MONO}), (", ", {}), ("asunto", {"font": MONO}), (", ", {}),
          ("mensaje", {"font": MONO}), (", ", {}), ("usuario_cr", {"font": MONO}), (", ", {}),
          ("id_viaje", {"font": MONO})], "la novedad es el registro"],
        [[tecla("F4")], [("viaje", {"font": MONO})], [("hs_aviso", {"font": MONO})],
         [("viaje_log", {"font": MONO}), (" motivo ", {}), ("AVISO", {"font": MONO}),
          (" — mejora sobre FoxPro", {"italic": True})]],
        [[tecla("F6"), (" a ", {}), tecla("F9"), (", ", {}), tecla("Ctrl+F8"), (", tipeo", {})],
         [("viaje", {"font": MONO})],
         [("Diagramador: ", {"bold": True}), ("cronogram2", {"font": MONO}), (" + ", {}),
          ("cronograma", {"font": MONO}), (" + ", {}), ("chequeo=0", {"font": MONO}),
          ("\nOperador: ", {"bold": True}), ("cronograma", {"font": MONO}), (" + ", {}),
          ("chequeo=0", {"font": MONO})],
         [("Solo modo operador: ", {}), ("viaje_log", {"font": MONO}), (" motivo ", {}),
          ("CBIO UNIDAD", {"font": MONO}), (", uno por tramo", {})]],
    ],
    widths=[3.4, 2.6, 6.4, 3.6], size=9,
)

h2("Detalles que importan al que lea el SQL")

bullet([("La tabla ", {}), ("viaje", {"font": MONO}), (" no tiene índice por ", {}),
        ("id_viaje", {"font": MONO}), (": la clave primaria es un artefacto de la réplica. Toda "
        "consulta tiene que anclar además en ", {}), ("f_reserva", {"font": MONO}),
        (" o termina recorriendo las 512.876 filas.", {})])
bullet([("Los nombres de columna vienen truncados a 10 caracteres desde FoxPro: ", {}),
        ("usuario_create", {"font": MONO}), (" es ", {}), ("usuario_cr", {"font": MONO}), (", ", {}),
        ("aviso_tiempo", {"font": MONO}), (" es ", {}), ("aviso_tiem", {"font": MONO}), (".", {})])
bullet([("El aviso efectivo de un servicio es ", {}),
        ("COALESCE(hs_aviso, hs_inicio − aviso_tiem)", {"font": MONO, "bold": True}),
        (". Buslink lo compara por fecha y hora completas; el Metrocar comparaba solo HH:MM, "
         "que es de donde salía el error de los servicios de medianoche.", {})])
bullet([("Todo cambio de cronograma va en una transacción única, releyendo la fila con ", {}),
        ("UPDLOCK", {"font": MONO}), (" adentro. FoxPro era efectivamente monousuario; la web no.", {})])
bullet([("Un servicio que es parte de una ruta se identifica por ", {}),
        ("id_viaje_i", {"font": MONO}), (" distinto de cero; el cambio pega a todos los tramos "
        "de ese identificador.", {})])

# ═══════════════════════ ANEXO B ═══════════════════════
h1("Anexo B — Diferencias entre Metrocar y Buslink", "B")

p("Todo lo que cambió a propósito, y por qué. Nada de esto es un error de migración.", size=10, color=GRIS)

tabla(
    ["Tecla", "Metrocar (FoxPro)", "Buslink", "Por qué"],
    [
        [[tecla("F1")], "Lista las teclas sin explicar ninguna", "Cada tecla explicada, agrupada "
         "por familia y filtrada por permiso", "La lista sola no alcanza para aprender"],
        [[tecla("F3")], "Única forma de ver algo nuevo", "La planilla se actualiza sola cada 60 s "
         "y resalta lo que cambió", "La pantalla congelada era fuente de errores"],
        [[tecla("F4")], "Precarga la hora del viaje, que es inválida: si grabás sin tocar nada, da error",
         "Precarga el aviso automático, más botones de 30/45/60 min", "Los datos dicen que el 98% "
         "de los avisos manuales piden entre 16 y 60 minutos"],
        [[tecla("F4")], "«No Avisar», que en realidad no apaga el aviso",
         "«Volver al aviso automático»", "El rótulo anterior decía algo que no pasaba"],
        [[tecla("F4")], "El aviso se pierde si no estabas mirando ese minuto exacto",
         "Se muestra hasta 5 minutos después, marcado como atrasado",
         "El operador tiene que enterarse igual"],
        [[tecla("F4")], "Cartel que dice «Inicio de 3 servicios»", "La lista de los 3 servicios, "
         "con doble clic para abrir el viaje", "Saber cuántos sin saber cuáles no sirve"],
        [[tecla("F4")], "Los servicios de después de medianoche avisan 24 h tarde",
         "Avisan cuando corresponde", "Son más de mil casos por año"],
        [[tecla("F2")], "Puede mandar correo al cliente", "Solo graba la novedad",
         "Es una acción hacia afuera; se decidió no sacarla de un sistema en prueba"],
        [[tecla("F6"), ("–", {}), tecla("F9")], "El cambio masivo confirma siempre, con un cartel "
         "que se aprieta sin leer", "Muestra a cuántas reservas alcanza antes de grabar",
         "Una confirmación que aparece siempre deja de informar"],
        [[tecla("F6"), ("–", {}), tecla("F9")], "El masivo va fila por fila, sin transacción",
         "Todo junto o nada", "Si fallaba a mitad, el tablero quedaba mezclado"],
        [[("Enter", {"font": MONO, "bold": True, "color": AZUL, "hl": HEX_CLARO})],
         "Abre el cambio de unidad", "Abre el Zoom del Viaje",
         "Es lo que la mesa ya tiene aprendido en Buslink"],
    ],
    widths=[2.0, 4.6, 4.8, 4.6], size=8.5,
)

# ═══════════════════════ ANEXO C ═══════════════════════
h1("Anexo C — Qué está activo hoy y qué espera al día D", "C")

p("Buslink todavía convive con el Metrocar. Mientras la tabla de viajes siga siendo del sistema "
  "viejo, cualquier cosa que Buslink escriba ahí sería pisada por la próxima replicación. Por eso "
  "las teclas que modifican datos están construidas y probadas, pero con la grabación frenada "
  "hasta el día del corte.")

tabla(
    ["Tecla", "Estado hoy", "Qué falta"],
    [
        [[tecla("F1")], [("Activa", {"bold": True, "color": VERDE})], "—"],
        [[tecla("F3")], [("Activa", {"bold": True, "color": VERDE})], "—"],
        [[tecla("F5")], [("Activa", {"bold": True, "color": VERDE})], "—"],
        [[tecla("F4")], [("Alarmas activas", {"bold": True, "color": VERDE}),
                         (" · grabación frenada", {"color": NARANJ})],
         "Encender la escritura el día D. Las alarmas ya funcionan con las horas que se cargan "
         "desde el Metrocar"],
        [[tecla("F2")], [("Lectura activa", {"bold": True, "color": VERDE}),
                         (" · alta frenada", {"color": NARANJ})],
         "El libro de novedades es una tabla propia y podría cortar antes que el resto"],
        [[tecla("F6"), ("–", {}), tecla("F9"), (", ", {}), tecla("Ctrl+F8")],
         [("Diálogos activos", {"bold": True, "color": VERDE}), (" · grabación frenada", {"color": NARANJ})],
         "Encender la escritura el día D"],
    ],
    widths=[3.0, 5.2, 7.8], size=9.5,
)

callout(
    "Qué significa «grabación frenada» en la pantalla",
    [("El diálogo abre normal, muestra los datos reales, valida todo y arma el cambio; el botón "
      "de grabar aparece con un candado y una leyenda que explica que se habilita cuando el "
      "módulo pase a Buslink. Se puede usar para practicar sin ningún riesgo: ", {}),
     ("no toca la base.", {"bold": True})],
    fill=HEX_CLARO, color_titulo=AZUL, borde=HEX_AZUL,
)

# ── pie ──
doc.add_paragraph()
par = doc.add_paragraph()
border(par, "top", size=8, color="D5DAE3", space=6)
par.paragraph_format.space_before = Pt(14)
run(par, "Buslink · Metrocar Nortur — Guía de teclas de la Planilla de Tráfico · agosto 2026",
    size=9, color=GRIS)
par2 = doc.add_paragraph()
run(par2, "Fuentes: relevamiento del sistema FoxPro (trafico2.scx, trafico_hs_aviso.scx, "
          "trafico_cambia_cronograma.scx, libro_novedad_abm.scx) y medición de uso real sobre la "
          "base de producción.", size=9, color=GRIS)

# ── numero de pagina en el pie ──
for sec in doc.sections:
    footer_par = sec.footer.paragraphs[0]
    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_par.add_run()
    r.font.size = Pt(8.5)
    r.font.name = BODY
    r.font.color.rgb = GRIS
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    r._r.append(fld1); r._r.append(instr); r._r.append(fld2)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "GUIA_TECLAS_TRAFICO_F1_F9.docx")
doc.save(out)
print("OK ->", out)
