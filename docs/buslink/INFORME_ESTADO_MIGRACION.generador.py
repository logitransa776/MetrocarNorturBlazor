# -*- coding: utf-8 -*-
"""
Genera INFORME_ESTADO_MIGRACION.docx — estado real de la migración Metrocar → Buslink
al 09/08/2026, qué falta, qué documentación falta y recomendaciones.

Fuentes (todas verificadas en esta sesión):
  · docs/Buslink/PLAN_MIGRACION_BUSLINK.md · ANALISIS_SISTEMA_BUSLINK.md · INFORME_AVANCE_BUSLINK.md
  · docs/PLANOFOXPRO/README.md + los 30 planos por módulo
  · .claude/skills/estado-buslink/SKILL.md (+ las 12 skills del proyecto)
  · Código: MainLayout.razor, AbmFeatureFlags.cs, AbmService.cs, ReportService.cs,
    ZoomViajeDialog.razor, appsettings.json
  · git log / git status
  · Menús FoxPro (C:\\MetroCarSys\\Menus\\*.MPR) y forms (378 .scx)
  · Consultas directas a replicaVPF (conteos y fechas de negocio por tabla)
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

AZUL   = RGBColor(0x00, 0x3A, 0xA0)
NARANJ = RGBColor(0xF9, 0x94, 0x10)
GRIS   = RGBColor(0x5A, 0x60, 0x6A)
ROJO   = RGBColor(0xB3, 0x26, 0x1E)
VERDE  = RGBColor(0x1B, 0x6E, 0x3C)
NEGRO  = RGBColor(0x1F, 0x24, 0x2C)

HEX_AZUL   = "003AA0"
HEX_NARANJ = "F99410"
HEX_CLARO  = "EEF2F9"
HEX_SUAVE  = "F7F8FA"
HEX_AVISO  = "FFF4E3"
HEX_ROJO   = "FDECEA"
HEX_VERDE  = "EAF5EE"

BODY = "Segoe UI"
MONO = "Consolas"

doc = Document()

st = doc.styles["Normal"]
st.font.name = BODY
st.font.size = Pt(10.5)
st.font.color.rgb = NEGRO
st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15

for sec in doc.sections:
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.9)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)


def shade(el, hexcolor):
    el.append(parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls("w"), hexcolor)))


def border(par, edge="bottom", size=12, color=HEX_NARANJ, space=4):
    pPr = par._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    e = OxmlElement("w:" + edge)
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(size))
    e.set(qn("w:space"), str(space))
    e.set(qn("w:color"), color)
    pBdr.append(e)


def run(par, text, bold=False, italic=False, size=10.5, color=None, font=BODY, hl=None):
    r = par.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color is not None:
        r.font.color.rgb = color
    if hl:
        shade(r._element.get_or_add_rPr(), hl)
    return r


def p(text="", size=10.5, bold=False, italic=False, color=None, space_after=6,
      space_before=0, align=None, left=0):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    if left:
        par.paragraph_format.left_indent = Cm(left)
    if align:
        par.alignment = align
    if text:
        run(par, text, bold=bold, italic=italic, size=size, color=color)
    return par


def rich(parts, space_after=6, space_before=0, left=0, size=10.5):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    if left:
        par.paragraph_format.left_indent = Cm(left)
    for txt, fmt in parts:
        f = dict(fmt)
        f.setdefault("size", size)
        run(par, txt, **f)
    return par


def h1(text, nro=None):
    par = doc.add_paragraph()
    par.paragraph_format.page_break_before = True
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.keep_with_next = True
    if nro:
        run(par, nro + "  ", bold=True, size=17, color=NARANJ)
    run(par, text, bold=True, size=17, color=AZUL)
    border(par, "bottom", size=10, color=HEX_NARANJ, space=3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def h2(text, color=AZUL):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.keep_with_next = True
    run(par, text, bold=True, size=13, color=color)
    return par


def h3(text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(9)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.keep_with_next = True
    run(par, text, bold=True, size=11, color=GRIS)
    return par


def bullet(text_parts, level=0):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(0.75 + 0.6 * level)
    if isinstance(text_parts, str):
        run(par, text_parts)
    else:
        for txt, fmt in text_parts:
            f = dict(fmt)
            f.setdefault("size", 10.5)
            run(par, txt, **f)
    return par


def pasos(items):
    for i, text_parts in enumerate(items, start=1):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.left_indent = Cm(1.05)
        par.paragraph_format.first_line_indent = Cm(-0.6)
        run(par, "{}.\t".format(i), bold=True, color=NARANJ)
        if isinstance(text_parts, str):
            run(par, text_parts)
        else:
            for txt, fmt in text_parts:
                f = dict(fmt)
                f.setdefault("size", 10.5)
                run(par, txt, **f)


def _cell_margins(cell, lr=110, tb=70):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", tb), ("left", lr), ("bottom", tb), ("right", lr)):
        e = OxmlElement("w:" + side)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def _tbl_borders(tbl, color, left_only=False, size=4):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    sides = ["left"] if left_only else ["top", "left", "bottom", "right", "insideH", "insideV"]
    todos = ["top", "left", "bottom", "right", "insideH", "insideV"]
    for s in todos:
        e = OxmlElement("w:" + s)
        if s in sides:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(24 if left_only else size))
            e.set(qn("w:color"), color)
        else:
            e.set(qn("w:val"), "none")
        e.set(qn("w:space"), "0")
        borders.append(e)
    tblPr.append(borders)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:tblHeader")
    e.set(qn("w:val"), "true")
    trPr.append(e)


def callout(titulo, texto, fill=HEX_AVISO, color_titulo=NARANJ, borde=HEX_NARANJ):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), fill)
    cell.text = ""
    par = cell.paragraphs[0]
    par.paragraph_format.space_after = Pt(2)
    run(par, titulo, bold=True, size=10.5, color=color_titulo)
    par2 = cell.add_paragraph()
    par2.paragraph_format.space_after = Pt(0)
    if isinstance(texto, str):
        run(par2, texto, size=10)
    else:
        for txt, fmt in texto:
            f = dict(fmt)
            f.setdefault("size", 10)
            run(par2, txt, **f)
    _cell_margins(cell, 140, 110)
    _tbl_borders(t, borde, left_only=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def tabla(headers, filas, widths=None, size=9.5, header_fill=HEX_AZUL, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, htxt in enumerate(headers):
        c = hdr.cells[i]
        if header_fill:
            shade(c._tc.get_or_add_tcPr(), header_fill)
        c.text = ""
        par = c.paragraphs[0]
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        run(par, htxt, bold=True, size=size, color=RGBColor(0xFF, 0xFF, 0xFF))
        _cell_margins(c, 90, 55)
    _repeat_header(hdr)
    for j, fila in enumerate(filas):
        row = t.add_row()
        for i, celda in enumerate(fila):
            c = row.cells[i]
            if zebra and j % 2 == 1:
                shade(c._tc.get_or_add_tcPr(), HEX_SUAVE)
            c.text = ""
            par = c.paragraphs[0]
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.space_before = Pt(0)
            if isinstance(celda, str):
                run(par, celda, size=size)
            else:
                for txt, fmt in celda:
                    f = dict(fmt)
                    f.setdefault("size", size)
                    run(par, txt, **f)
            _cell_margins(c, 90, 55)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    if header_fill is None:
        t._tbl.remove(t.rows[0]._tr)
    _tbl_borders(t, "D5DAE3", size=4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def code_block(lineas, size=9):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), "F4F6F9")
    cell.text = ""
    for i, ln in enumerate(lineas):
        par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        run(par, ln, font=MONO, size=size, color=RGBColor(0x33, 0x39, 0x42))
    _cell_margins(cell, 140, 90)
    _tbl_borders(t, "C9D2E0", left_only=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def mono(txt):
    return (txt, {"font": MONO, "size": 9.5, "color": AZUL})


B = {"bold": True}
N = {}
R = {"color": ROJO, "bold": True}
V = {"color": VERDE, "bold": True}
G = {"color": GRIS}

OK   = [("LISTO", {"bold": True, "color": VERDE, "size": 8.5})]
AND_ = [("ANDAMIAJE", {"bold": True, "color": NARANJ, "size": 8.5})]
PEND = [("PENDIENTE", {"bold": True, "color": ROJO, "size": 8.5})]
NO   = [("NO MIGRAR", {"bold": True, "color": GRIS, "size": 8.5})]


# ═══════════════════════════ PORTADA ═══════════════════════════
for _ in range(3):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(0)
run(par, "BUSLINK  ·  METROCAR NORTUR", bold=True, size=11, color=NARANJ)
par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(0)
run(par, "Migración FoxPro → Blazor / SQL Server", size=11, color=GRIS)

par = doc.add_paragraph()
par.paragraph_format.space_before = Pt(18)
par.paragraph_format.space_after = Pt(0)
run(par, "Estado real y", bold=True, size=32, color=AZUL)
par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(4)
run(par, "qué falta para terminar", bold=True, size=32, color=NARANJ)
border(par, "bottom", size=18, color=HEX_NARANJ, space=8)

par = doc.add_paragraph()
par.paragraph_format.space_before = Pt(14)
run(par, "Auditoría del proyecto al 9 de agosto de 2026: lo construido, lo documentado, "
         "lo que queda del Metrocar, lo que conviene NO migrar y los riesgos que hay que "
         "cerrar antes del día del corte.",
    size=12, color=GRIS)

for _ in range(2):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

tabla(
    ["", ""],
    [
        [[("Para quién", B)], "Claudio Marañon (dev) y el dueño del sistema en NORTUR"],
        [[("Método", B)], "Relevamiento de los 3 documentos maestros, los 30 planos FoxPro, las 13 skills, "
                          "el código fuente de Buslink, los menús del Metrocar y consultas directas a replicaVPF"],
        [[("Corte", B)], "9 de agosto de 2026"],
        [[("Naturaleza", B)], "Informe de auditoría — todo lo afirmado acá se verificó contra código, "
                              "documentos o base de datos; las opiniones van marcadas como tales"],
    ],
    widths=[3.4, 12.4], size=10, header_fill=None, zebra=True,
)


# ═══════════════════════════ CONTENIDO ═══════════════════════════
h1("Contenido")

contenido = [
    ("1.", "Resumen ejecutivo", "Dónde estamos parados de verdad, en una página"),
    ("2.", "Lo que fuimos haciendo", "La línea de tiempo y el inventario medido hoy"),
    ("3.", "Los documentos del proyecto", "Qué hay, para qué sirve cada uno y cuál está vencido"),
    ("4.", "Metrocar → Buslink: el mapa completo", "Módulo por módulo, qué se migró y qué no"),
    ("5.", "Lo que falta para el Día D", "El circuito viaje, fase por fase, con el estado real"),
    ("6.", "Lo que falta DESPUÉS del Día D", "Y qué conviene declarar fuera de alcance"),
    ("7.", "Las pantallas que quedan fuera de Buslink", "El cruce ítem por ítem contra el menú del Metrocar"),
    ("8.", "Documentación faltante", "Los planos que todavía no existen"),
    ("9.", "Riesgos y deudas técnicas", "Lo que encontré mirando el código y la base"),
    ("10.", "Ideas y recomendaciones", "Mi criterio, marcado como opinión"),
    ("11.", "Plan sugerido para las próximas semanas", "Orden concreto de trabajo"),
    ("A.", "Anexo — Vigencia real de cada módulo", "Medido sobre la base, no supuesto"),
    ("B.", "Anexo — Checklist de arranque inmediato", "Lo que yo haría esta semana"),
]
for nro, tit, sub in contenido:
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(5)
    par.paragraph_format.left_indent = Cm(0.2)
    run(par, nro + "  ", bold=True, size=11, color=NARANJ)
    run(par, tit, bold=True, size=11, color=AZUL)
    if sub:
        run(par, "   " + sub, size=10, color=GRIS)

callout(
    "Cómo leer este informe",
    [("Los capítulos 1, 5 y 10 son los que hay que leer sí o sí. El 4 y el 6 son el inventario "
      "para consultar. El 8 y el 9 son donde te digo lo que veo que puede doler y lo que "
      "yo haría distinto. Todo lo que dice ", N), ("verificado", B),
     (" salió de mirar el código o consultar la base en el momento de escribir esto.", N)],
    fill=HEX_CLARO, color_titulo=AZUL, borde=HEX_AZUL,
)


# ═══════════════════════ 1. RESUMEN EJECUTIVO ═══════════════════════
h1("Resumen ejecutivo", "1")

p("El proyecto está mucho más avanzado de lo que dicen sus propios documentos maestros, y "
  "al mismo tiempo más lejos del día del corte de lo que parece. Las dos cosas son ciertas "
  "y conviven por la misma razón: se construyó muchísima pantalla y muy poca infraestructura de corte.")

h2("Los cinco números que definen el momento")

tabla(
    ["Métrica", "Valor", "Comentario"],
    [
        ["Pantallas Blazor con ruta propia", [("49", B)],
         "El documento de análisis del 02/07 decía 13. Se triplicó en cinco semanas."],
        ["Ítems del menú lateral ya operativos", [("44", B)],
         "Contra 105 ítems que siguen deshabilitados (el backlog a la vista)"],
        ["Interruptores de escritura apagados", [("25", {"bold": True, "color": NARANJ})],
         "Código de alta/baja/modificación escrito y compilado, pero que nunca tocó la base"],
        ["ABMs con escritura REAL en producción", [("1", {"bold": True, "color": ROJO})],
         "Usuarios y Permisos. Es el único dato que Buslink escribe hoy."],
        ["Ítems de la Fase 0 todavía abiertos", [("5 de 8", {"bold": True, "color": ROJO})],
         "Y son justamente los que habilitan el corte, no los que agregan pantallas"],
    ],
    widths=[5.4, 2.2, 8.2], size=9.5,
)

h2("El diagnóstico en tres frases")

pasos([
    [("La lectura está terminada.", B), (" Todo el sistema operativo de NORTUR se puede consultar "
      "hoy desde Buslink: tráfico, flota, choferes, facturación con su motor de tarifas validado, "
      "combustible, reservas e informes. En varios puntos Buslink ya es mejor que el Metrocar "
      "(cross-filter, exportaciones, tablero de ocupación, centro de ayuda, avisos por hora).", N)],
    [("La escritura está construida pero no encendida.", B), (" Hay 25 flags en ", N), mono("false"),
     (". Es capital real —el código existe, sigue los planos y compila— pero es capital ", N),
     ("no probado contra la base", B), (": un método que nunca se ejecutó no está validado, por "
      "más fiel que sea al plano.", N)],
    [("Lo que falta no es pantalla: es el corte.", B), (" Faltan el interruptor de sincronización, "
      "el bloqueo del FoxPro, el mapeo campo a campo de las 12 tablas del circuito, el motor "
      "unificado de escritura y el ensayo general. Nada de eso se resuelve construyendo otra "
      "pantalla más.", N)],
])

h2("Los cinco bloqueantes reales del Día D")

tabla(
    ["#", "Bloqueante", "Estado hoy"],
    [
        ["1", [("El interruptor de sincronización DBF→SQL no está documentado", B),
               ("\nNadie escribió cómo se apaga la réplica tabla por tabla, quién la opera, cuánto "
                "tarda en propagar, ni qué hace con filas que existen en SQL y no en DBF.", N)],
         [("Fase 0 ítem 3 · sin empezar", {"color": ROJO, "bold": True})]],
        ["2", [("El bloqueo del FoxPro no está probado", B),
               ("\nNo hay mecanismo verificado para dejar el Metrocar en solo consulta. Es el paso 5 "
                "del runbook del día D y hoy es una hipótesis.", N)],
         [("Fase 0 ítem 4 · sin empezar", {"color": ROJO, "bold": True})]],
        ["3", [("El mapeo campo a campo de las 12 tablas no existe", B),
               ("\nIncluye cómo se asigna ", N), mono("_sync_id"), (" en un INSERT hecho por Blazor. "
                "El propio plan lo llama show-stopper si se descubre el día D.", N)],
         [("Fase 0 ítem 5 · sin empezar", {"color": ROJO, "bold": True})]],
        ["4", [("El motor compartido de escritura del viaje nunca se construyó", B),
               ("\nNo existe ", N), mono("ViajeAbmService.cs"), (". Las primitivas se fueron "
                "copiando dentro de ", N), mono("AbmService"), (" pantalla por pantalla: el INSERT "
                "en ", N), mono("viaje_log"), (" ya está escrito tres veces en tres lugares distintos.", N)],
         [("Fase 2 · desviada", {"color": ROJO, "bold": True})]],
        ["5", [("Faltan tablas por replicar en el servidor nuevo", B),
               ("\n", N), mono("viaje_log_chofer"), (" (75.001 filas) no está replicada en ningún lado — "
                "verificado hoy: tampoco existe en el servidor local. Y ", N), mono("cabecera"),
               (", ", N), mono("chofer_franco"), (", ", N), mono("chofer_viatico"),
               (" estaban en el viejo y no en el nuevo.", N)],
         [("Fase 0 ítem 8 · abierto", {"color": ROJO, "bold": True})]],
    ],
    widths=[0.8, 10.0, 5.0], size=9.5,
)

callout(
    "Lo más importante que encontré",
    [("El último commit del repositorio es del ", N), ("21 de julio", B), (". Desde entonces hay ",N),
     ("44 archivos modificados (+9.987 / −3.406 líneas) y 31 archivos nuevos sin versionar", B),
     (", que incluyen módulos enteros: el cambio de cronograma (F6-F9), el menú del panel Buses, "
      "el centro de ayuda, el tablero de ocupación de flota, la auditoría de accesos y el control "
      "de acceso por Internet. Son casi tres semanas de trabajo sin respaldo en GitHub. "
      "Es lo primero que hay que resolver, antes que cualquier otra cosa de este informe.", N)],
    fill=HEX_ROJO, color_titulo=ROJO, borde="B3261E",
)


# ═══════════════════════ 2. LO QUE FUIMOS HACIENDO ═══════════════════════
h1("Lo que fuimos haciendo", "2")

p("Esta es la reconstrucción de la historia del proyecto a partir de los documentos, las skills "
  "y el historial de git. Sirve para dos cosas: para ver el ritmo real (que es alto) y para "
  "entender por qué el orden de las fases terminó siendo distinto al del plan.")

h2("Línea de tiempo")

tabla(
    ["Cuándo", "Qué se entregó", "Qué significó"],
    [
        ["Jun 2026", "Arquitectura Blazor Server + MudBlazor + ApexCharts + tema NORTUR + drawer propio; "
                     "login contra la tabla usuario con el flujo exacto del FoxPro",
                     "La base sobre la que se apoyó todo lo demás"],
        ["15/06/2026", "Choferes y Vehículos-Flota en solo lectura (lista + ficha multi-pestaña)",
                       "Se estrenó el patrón «lista + ficha» que después se repitió 20 veces"],
        ["16/06/2026", "Reglas de performance: pooling de conexiones, warmup, Virtualize en grillas grandes",
                       "El sistema pasó de sentirse lento a sentirse instantáneo"],
        ["18-22/06/2026", "Módulo Facturación en lectura + motor de tarifas migrado y validado "
                          "contra 8.656 viajes históricos (99,4% exacto)",
                          "La pieza técnicamente más difícil del sistema, y ya está resuelta"],
        ["01/07/2026", "ABM de Usuarios y Permisos: el PRIMER ABM con escritura real",
                       "Se probó de punta a punta que la estrategia de cambiar de dueño una tabla funciona"],
        ["02/07/2026", "Plan de migración Buslink aprobado + biblioteca FoxPro reorganizada + "
                       "especificación de escritura del despacho",
                       "El proyecto pasó de «migrar informes» a «migrar la operación»"],
        ["03-07/07/2026", "Informes analíticos con cross-filter estilo Power BI; Reservas, Tráfico, "
                          "Combustible y Vehículos completados con andamiaje de escritura",
                          "Se vaciaron de placeholders cuatro menús enteros"],
        ["15-30/07/2026", "Rediseño de Reservas por cliente con las correcciones de la clienta; "
                          "banda horaria con alcance corregido; comparación entre períodos",
                          "Primera devolución real de un usuario final incorporada"],
        ["31/07-04/08/2026", "Tablero de ocupación de flota; las 9 teclas de función de la planilla "
                             "(F1-F9 + Ctrl+F8); menú completo del panel Buses; centro de ayuda; "
                             "auditoría de accesos; control de acceso por Internet",
                             "La planilla de tráfico dejó de ser una grilla y pasó a ser una consola de despacho"],
    ],
    widths=[2.5, 7.5, 5.8], size=9.5,
)

h2("El inventario, medido hoy")

p("Lo siguiente no viene de ningún documento: lo conté sobre el código en esta sesión.", size=10, color=GRIS)

tabla(
    ["Pieza", "Cantidad", "Detalle"],
    [
        ["Páginas con ruta", "49", "Components/Pages/*.razor"],
        ["Diálogos y componentes compartidos", "68", "Components/Shared/*.razor"],
        ["Métodos de lectura", "121", "ReportService.cs — 8.721 líneas"],
        ["Métodos de escritura", "66", "AbmService.cs — 4.780 líneas, 43 transacciones"],
        ["Exportaciones a Excel", "—", "ExcelExportService.cs — 2.393 líneas"],
        ["Planos de lógica FoxPro", "30", "docs/PLANOFOXPRO/ en 7 carpetas de módulo"],
        ["Skills de conocimiento", "13", ".claude/skills/ — el «cómo trabajar» del proyecto"],
        ["Tests automatizados", "4 archivos", "Playwright: smoke + clientes + un filtro de tráfico"],
    ],
    widths=[6.0, 2.4, 7.4], size=9.5,
)

callout(
    "Por qué el orden de las fases se dio vuelta",
    "El plan decía: cerrar incógnitas (Fase 0) → catálogos (1) → motor (2) → Tráfico (3) → Reservas (4). "
    "Lo que pasó en la práctica fue construir pantallas de Fase 3 y 4 con andamiaje mientras la Fase 0 "
    "quedaba abierta. No fue un error: cada pantalla nueva daba valor visible y consolidaba el patrón. "
    "Pero tiene un costo que hoy hay que pagar: hay 25 escrituras codificadas que nunca se ejecutaron, "
    "y ninguna de ellas se puede probar de verdad hasta que exista el entorno de corte.",
    fill=HEX_CLARO, color_titulo=AZUL, borde=HEX_AZUL,
)


# ═══════════════════════ 3. LOS DOCUMENTOS ═══════════════════════
h1("Los documentos del proyecto", "3")

p("Me pediste saber qué documentos recorrí. Son estos, y de paso te marco cuáles siguen "
  "siendo fuente de verdad y cuáles ya quedaron desactualizados.")

h2("Documentos maestros")

tabla(
    ["Documento", "Qué es", "Vigencia"],
    [
        [[("CLAUDE.md", {"font": MONO, "size": 9})], "Las instrucciones vivas del proyecto: stack, conexión, "
         "arquitectura, modelo de datos, reglas de negocio, estado por módulo",
         [("Vigente", V), (" — es lo mejor mantenido del repo", N)]],
        [[("docs/Buslink/PLAN_MIGRACION_BUSLINK.md", {"font": MONO, "size": 9})],
         "El roadmap: 8 fases, las 12 tablas del corte, runbook del día D, top 8 de riesgos, "
         "18 criterios de «listo»",
         [("Vigente en la estrategia", V), (", desactualizado en el estado (dice Fase 0 en curso "
          "con 1 de 7 ítems; hoy hay Fase 3 y 4 construidas con andamiaje)", N)]],
        [[("docs/Buslink/ANALISIS_SISTEMA_BUSLINK.md", {"font": MONO, "size": 9})],
         "Análisis del sistema para el cliente: arquitectura, inventario, seguridad, glosario",
         [("VENCIDO", R), (" — corte 02/07. Dice 13 páginas y 11 migradas; hoy son 49 y 44. "
          "Regenerarlo con este informe como insumo.", N)]],
        [[("docs/Buslink/INFORME_AVANCE_BUSLINK.md", {"font": MONO, "size": 9})],
         "El plan convertido en pasos operativos para tildar semana a semana",
         [("VENCIDO", R), (" — mismo corte 02/07, mismo problema", N)]],
        [[("docs/PLANOFOXPRO/README.md", {"font": MONO, "size": 9})],
         "Índice maestro de los 30 planos, con el estado de migración de cada uno",
         [("Vigente", V), (" — es el mapa más útil del proyecto", N)]],
        [[(".claude/skills/estado-buslink/SKILL.md", {"font": MONO, "size": 9})],
         "El detalle módulo por módulo de qué está migrado, con trampas y decisiones",
         [("Vigente", V), (" — 593 líneas, es la memoria técnica real. Ojo: todavía no está "
          "versionado en git.", N)]],
    ],
    widths=[5.2, 5.4, 5.2], size=9.5,
)

h2("La biblioteca de planos FoxPro (30 documentos)")

p("Cada plano es la lógica real de un formulario del Metrocar, extraída del binario ", size=10)
rich([mono(".scx"), (" con un lector propio, no reconstruida de memoria. Incluye el SQL exacto, "
      "las validaciones y —clave— los bugs heredados que hay que NO copiar.", N)], size=10)

tabla(
    ["Carpeta", "Documentos", "Cubre"],
    [
        ["trafico/", "14", "Zoom, filtros, toolbar de escritura, cronograma, F2/F4, menú Buses, "
                           "historial, GPS, cabeceras, francos, viáticos, voucher/guardia/contactos"],
        ["reservas/", "5", "Alta manual, plantillas, importa Excel, informes de banda horaria y por cliente"],
        ["catalogos/", "8", "Los ABMs de la Fase 1: motivos, feriados, destinos, operadores, clientes, guías, grupos"],
        ["vehiculos-choferes/", "8", "Choferes, odómetros, siniestros, fleteros, tipos, agenda, los 2 informes de flota"],
        ["facturacion/", "1", "El módulo entero en un solo documento largo"],
        ["combustible/", "2", "Las dos eras del módulo + el mapeo del menú"],
        ["sistema/", "1", "Permisos, niveles, flujo de login"],
    ],
    widths=[3.6, 2.0, 10.2], size=9.5,
)

h2("Documentación técnica de apoyo")

bullet([mono("docs/performance/PERFORMANCE_GRILLAS_Y_CONEXION.md"), (" — por qué las grillas eran lentas y las reglas que salieron de ahí", N)])
bullet([mono("docs/performance/PENDIENTE_GRILLA_TRAFICO_BLANQUEO.md"), (" — un problema abierto documentado con mediciones y con los intentos fallidos. Es buena práctica: evita re-intentar lo que ya falló.", N)])
bullet([mono("docs/Buslink/ACTIVAR_ABM_VEHICULOS_CHOFERES.md"), (" — el checklist de encendido de un ABM", N)])
bullet([mono("docs/sql/usuario_sesion.sql"), (", ", N), mono("usuarios_logs.sql"), (", ", N), mono("docs/sql-indices/"), (" — scripts que hay que correr en el servidor nuevo antes del corte", N)])
bullet([mono("docs/Buslink/GUIA_TECLAS_TRAFICO_F1_F9.docx"), (" — manual de usuario de las teclas, ya en Word", N)])
bullet("Seis documentos .docx más para el cliente (combustible, facturación, seguridad, performance, testing, publicación IIS)")


# ═══════════════════════ 4. EL MAPA COMPLETO ═══════════════════════
h1("Metrocar → Buslink: el mapa completo", "4")

p("El menú principal del Metrocar tiene 271 entradas y 378 formularios detrás. El menú lateral "
  "de Buslink lo replica entero: lo migrado es un link, lo pendiente aparece deshabilitado. "
  "Esta es la foto módulo por módulo.")

h2("Leyenda")
rich([("LISTO", {"bold": True, "color": VERDE}), (" = funciona en Buslink.   ", N),
      ("ANDAMIAJE", {"bold": True, "color": NARANJ}), (" = la pantalla existe y la escritura está "
       "codificada, pero apagada por un interruptor.   ", N),
      ("PENDIENTE", {"bold": True, "color": ROJO}), (" = no se tocó.   ", N),
      ("NO MIGRAR", {"bold": True, "color": GRIS}), (" = mi recomendación de descartarlo (ver cap. 6).", N)],
     size=10)

h2("Reservas")
tabla(
    ["Ítem del Metrocar", "Estado", "Nota"],
    [
        ["Reservas Especiales (alta manual)", AND_, "Puerta de alta al circuito viaje — se enciende el día D"],
        ["Reservas por Plantillas (armado masivo)", AND_, "Con preview dry-run, que el FoxPro no tiene"],
        ["Mantenimiento de Plantillas", AND_, ""],
        ["Crear Plantillas", PEND, "No pedido en las entregas anteriores"],
        ["Importa Reservas desde Excel", PEND, "Fase 4 puerta 3 — candidata explícita a descarte"],
        ["Operadores · Grupos · Destinos", AND_, "Grupos corta el día D (su baja cancela viajes en cascada)"],
        ["Clientes (desde el menú Reservas)", OK, "Reusa la pantalla de Facturación, en lectura"],
        ["Informes: fecha/servicio, banda horaria, por cliente", OK, "Con cross-filter, drill-down y Excel multi-hoja"],
        ["Bandas Horarios (catálogo)", PEND, "6 filas, casi estático"],
    ],
    widths=[6.6, 2.0, 7.2], size=9.5,
)

h2("Tráfico")
tabla(
    ["Ítem del Metrocar", "Estado", "Nota"],
    [
        ["Planilla de Tráfico (operación del día)", OK, "Lectura completa: 25 columnas, panel Buses, 8 filtros, "
         "auto-refresh, menú contextual, ocupación de flota"],
        ["Zoom del Viaje · Historial", OK, "Lectura. La edición es la última entrega de la Fase 3"],
        ["Teclas F1 · F3 · F5 (ver)", OK, "Activas"],
        ["Tecla F4 (alarma por hora)", AND_, "El motor de avisos YA está activo; solo la escritura espera"],
        ["Teclas F6-F9 · Ctrl+F8 (cambio de cronograma)", AND_, "La operación más frecuente del circuito: 154 por día"],
        ["Tecla F2 (libro de guardia)", AND_, "Podría cortar antes del día D: la tabla es autocontenida"],
        ["Toolbar de despacho (chequeo, asignar, reasignar, finalizar, cancelar, reactivar)", PEND,
         "Es el corazón de la Fase 3 — hoy no hay ni un botón"],
        ["Menú del panel Buses (16 + 4 ítems)", AND_, "Logoneo bloqueado además por tabla sin replicar"],
        ["Cabeceras · Francos (3 pantallas) · Viáticos (3)", AND_, "Las tablas no están en el servidor nuevo"],
        ["Voucher · Guardia · Contactos · Rubros · Lista de pasajeros", AND_, ""],
        ["Adicionales Stock (Mantenimiento e Ingreso)", PEND, [("Módulo VIVO y sin documentar — ver cap. 6", {"bold": True, "color": ROJO})]],
        ["Libro de Novedades (menú) · Envío de correos · Parámetros de correo", PEND, "Tabla viva: 48.353 filas"],
        ["Controles sobre estados de reservas · Web Aeropuertos 2000", PEND, ""],
    ],
    widths=[6.6, 2.0, 7.2], size=9.5,
)

h2("Vehículos y Choferes")
tabla(
    ["Ítem del Metrocar", "Estado", "Nota"],
    [
        ["Choferes · Vehículos-Flota", OK, "Lista + ficha multi-pestaña, con vencimientos resaltados"],
        ["Odómetros · Siniestros · Agenda de Vencimientos", OK, "Lectura + informe"],
        ["Fleteros · Tipo de Vehículos", AND_, "Fleteros es catálogo compartido con Facturación"],
        ["Informes de flota (viajes por chofer, km unidades)", OK, "Corrigen dos bugs del original"],
        ["Apercibimientos + Motivos", NO, "Tabla con 0 filas: nunca se usó"],
        ["Capacitaciones (consulta y armado)", NO, "Último registro: 2013"],
    ],
    widths=[6.6, 2.0, 7.2], size=9.5,
)

h2("Facturación")
tabla(
    ["Ítem del Metrocar", "Estado", "Nota"],
    [
        ["Resumen de Liquidaciones", OK, "Maestro-detalle + comprobante"],
        ["Liquidación a Clientes", OK, [("Valorizada en vivo por el motor de tarifas. Falta el «Graba».", B)]],
        ["Liquidaciones estimadas", OK, [("Construida pero SIN link en el menú lateral — está huérfana", {"color": ROJO})]],
        ["ABM Clientes", OK, "Lectura"],
        ["Clientes Tarifas · Descuentos · Empresa Facturación", PEND, ""],
        ["Tarifario de Venta (4 pantallas)", PEND, "Tarifas vivas, con vigencias cargadas hasta 10/2026"],
        ["Tarifario de Choferes (4 pantallas)", PEND, "3.751 filas"],
        ["Adicionales · Rubro Adicionales (4 pantallas)", PEND, "Se leen desde el motor, el ABM sigue en FoxPro"],
        ["Liquidación a Fleteros", NO, "12 liquidaciones en total, la última de 2023"],
        ["Liquidación a Choferes (5 pantallas)", PEND, "Adelantos hasta 02/2025 — confirmar con el dueño si sigue viva"],
    ],
    widths=[6.6, 2.0, 7.2], size=9.5,
)

h2("Combustible · Taller · ABM del sistema · Utilitarios")
tabla(
    ["Ítem del Metrocar", "Estado", "Nota"],
    [
        ["Combustible — los 10 ítems del menú", [("LISTO", {"bold": True, "color": VERDE, "size": 8.5}),
         (" / ", {"size": 8.5}), ("ANDAMIAJE", {"bold": True, "color": NARANJ, "size": 8.5})],
         "Informes en lectura + conciliación con andamiaje. El menú no tiene placeholders."],
        ["Taller — los 16 ítems (OT, stock, chequeo, catálogos)", NO,
         [("Última orden de trabajo: 05/08/2019. Stock y depósitos con 0 filas.", B)]],
        ["ABM del sistema — 14 catálogos", PEND, "Feriados, motivos, servicios, zonas, guías, dueños, "
         "permisos, IATA, nacionalidades, profesiones, cronogramas de servicio"],
        ["ABM del sistema — Parámetros (4 pantallas)", PEND,
         [("La tabla parametro cambia de dueño el día D y su pantalla no está migrada ni documentada", {"color": ROJO})]],
        ["ABM del sistema — Facturación (monedas, cotizaciones, bancos, impuestos, comprobantes)", PEND, ""],
        ["Utilitarios — Scheduler, agenda, backup, chat, conectados, servicio XML", NO,
         "Herramientas de sistema del FoxPro; el equivalente moderno ya existe o no aplica"],
        ["Utilitarios — Reparaciones (normalizadores, indexadores, compactar)", NO,
         "Son utilidades de mantenimiento de los DBF. Sin sentido en SQL Server."],
        ["Utilitarios — los 3 informes (por cliente, por chofer, km unidades)", OK, "Ya migrados y mejorados"],
        ["Accesos — Tablero de Control", PEND, "Permiso X; solo lo tiene el supervisor"],
        ["Accesos — Cambio de password", PEND, "Hoy se cambia desde el ABM de Usuarios"],
    ],
    widths=[6.6, 2.0, 7.2], size=9.5,
)


# ═══════════════════════ 5. LO QUE FALTA PARA EL DÍA D ═══════════════════════
h1("Lo que falta para el Día D", "5")

p("El día D es el momento en que 12 tablas cambian de dueño y el FoxPro queda de consulta. "
  "Esto es lo que falta, fase por fase, con el estado verificado hoy —no el que dice el plan.")

h2("Fase 0 — Cerrar las incógnitas del corte")

tabla(
    ["#", "Entregable", "Estado verificado"],
    [
        ["1", "Especificación de escritura del despacho (toolbar de tráfico)",
         [("HECHO", V), (" — TRAFICO2_TOOLBAR.md + matriz ESCRITURA_CIRCUITO.md", N)]],
        ["2", "Decisión sobre la integración GPS",
         [("Documentado, falta la firma del dueño", {"color": NARANJ, "bold": True}),
          (". Hallazgo: hoy es un no-op, los dos interruptores están apagados en producción.", N)]],
        ["3", "Documento del interruptor de sincronización",
         [("SIN EMPEZAR", R), (" — verificado: no existe el archivo. Es la palanca del corte Y del rollback.", N)]],
        ["4", "Documento del bloqueo del FoxPro",
         [("SIN EMPEZAR", R), (" — probar en copia, no el día D", N)]],
        ["5", "Mapeo campo a campo de las 12 tablas + resolución de _sync_id",
         [("SIN EMPEZAR", R), (" — el plan lo marca como show-stopper", N)]],
        ["6", "Re-plantear el índice por id_viaje al cliente",
         [("Sin evidencia de resolución", {"color": NARANJ, "bold": True}),
          (". Mientras tanto rige la regla: todo WHERE de escritura lleva f_reserva.", N)]],
        ["7", "Regla del permiso F (ocultar importes)",
         [("PARCIAL", {"color": NARANJ, "bold": True}), (" — aplicado en las 4 pantallas de Facturación "
          "y en el alta de reservas, pero ", N), ("NO en el Zoom del Viaje", B),
          (": las líneas 193 y 198 muestran «Importe a Liq.» e «Importe a Pagar» sin control. "
           "Es justo la pantalla que abre el operador de tráfico, que es quien no debería verlos.", N)]],
        ["8", "Pedir la réplica de las tablas faltantes",
         [("ABIERTO", R), (" — ", N), mono("viaje_log_chofer"), (" no existe en ninguno de los dos "
          "servidores (verificado hoy sobre replicaVPF). Además hay que re-verificar ", N),
          mono("cabecera"), (", ", N), mono("chofer_franco"), (" y ", N), mono("chofer_viatico"),
          (" en el servidor nuevo.", N)]],
    ],
    widths=[0.8, 5.6, 9.4], size=9.5,
)

h2("Fase 1 — Catálogos con cutover temprano")

p("El plan pedía 5 ABMs que pueden cambiar de dueño antes del día D, para achicar el alcance del corte. "
  "Estado real: de los 5, hay 2 construidos con andamiaje (Destinos y Operadores) y 3 sin empezar "
  "(motivos de cancelación, feriados y el maestro de Clientes en escritura). "
  "Del grupo que corta el día D, están Grupos y Plantillas; faltan Guías y el ABM de Francos "
  "puntual. Ninguno hizo cutover todavía.", size=10)

callout(
    "Un pendiente operativo que no es de software",
    "Hay CERO feriados de 2026 cargados en el sistema (15 filas, todas de años anteriores). "
    "Mientras tanto, el armado de plantillas del FoxPro genera viajes en los feriados como si fueran "
    "días comunes. Esto se arregla cargándolos en el FoxPro hoy mismo: no hace falta esperar ninguna migración.",
    fill=HEX_AVISO,
)

h2("Fase 2 — El motor de escritura del circuito")

rich([("Acá está la desviación más importante del plan. ", B),
      ("No existe ", N), mono("Services/ViajeAbmService.cs"), (". Lo que hay es ", N),
      mono("AbmService.cs"), (" con 4.780 líneas y 66 métodos, donde las primitivas del circuito se "
      "fueron escribiendo por pantalla:", N)])

bullet([("El INSERT en ", N), mono("viaje_log"), (" está escrito ", N), ("tres veces", B),
        (" en tres lugares distintos del archivo (líneas 2544, 2656 y 3770).", N)])
bullet([("El contador atómico de ", N), mono("parametro"), (" sí está unificado en un helper "
        "(", N), mono("SiguienteParametroAsync"), ("), aunque usa ", N), mono("UPDATE + SELECT"),
        (" en vez del ", N), mono("OUTPUT inserted"), (" que pedía el plan. Dentro de la transacción "
        "es correcto —el lock exclusivo se sostiene hasta el commit—, pero conviene unificarlo.", N)])
bullet([("Hay 43 transacciones, o sea que la mejora sobre el FoxPro (que no usa ninguna) "
        "sí se está aplicando.", N)])

p("Mi lectura: el motor existe de hecho, pero disperso. Antes de arrancar la Fase 3 —que agrega "
  "diez operaciones más sobre las mismas tablas— hay que consolidarlo. Si no, cada operación nueva "
  "vuelve a copiar la bitácora y las validaciones, y el día que haya que cambiar una regla habrá "
  "que cambiarla en diez lugares. Es media semana de trabajo que ahorra semanas después.")

h2("Fase 3 — Tráfico en escritura (la fase que da el valor)")

p("Es la que hace que NORTUR pueda cargar los internos desde Buslink. Diez operaciones, en este orden:", size=10)

tabla(
    ["#", "Operación", "Estado"],
    [
        ["1", "Chequeo — el «hola mundo» de la escritura", PEND],
        ["2", "Asignar unidad y chofer — el corazón del despacho", PEND],
        ["3", "Liberar (volver a Sin Asignar)", PEND],
        ["4", "Otra unidad / Reasignar con motivo", PEND],
        ["5", "Finalizar (el botón «Libe», que en realidad cierra el viaje)", PEND],
        ["6", "Cancelar con motivo + cascada de grupo", PEND],
        ["7", "Reactivar", PEND],
        ["8", "Francos", AND_],
        ["9", "Zoom del Viaje en edición (~35 campos + diff de auditoría)", PEND],
        ["10", "Duplicar viaje y valor de servicio", PEND],
        ["+", "Cambio de cronograma (F6-F9) — no estaba en el plan original", AND_],
        ["+", "Alarma por hora (F4) y libro de guardia (F2)", AND_],
    ],
    widths=[0.8, 10.2, 4.8], size=9.5,
)

p("O sea: de las diez operaciones del plan, cero están construidas. Lo que sí se construyó "
  "(cronograma, F4, F2) fueron operaciones que el plan no había identificado y que resultaron "
  "ser las más frecuentes de la operación real. Eso es bueno —se descubrió trabajando— pero no "
  "reemplaza a la toolbar.", size=10)

h2("Fases 4, 5 y 6")

tabla(
    ["Fase", "Estado", "Qué falta concretamente"],
    [
        ["4 — Reservas (las 3 puertas de alta)",
         [("2 de 3 con andamiaje", {"color": NARANJ, "bold": True})],
         "Alta manual y plantillas están codificadas. Falta la importación desde Excel (28 columnas, "
         "3 etapas de validación) — es la candidata declarada a descarte. Falta también migrar el "
         "«deshacer lote», que es el botón de emergencia del primer día."],
        ["5 — Facturación: el Graba",
         [("Sin empezar", R)],
         "El motor de cálculo ya está y validado. Falta el grabado transaccional, el Revertir "
         "corregido, los servicios 2º y 3º, las rutas, el ajuste global manual y el test de cuadre "
         "contra las últimas 3 liquidaciones reales."],
        ["6 — Ensayo general",
         [("Sin empezar", R)],
         "Feature flag global de escritura, operación sombra de 3 a 5 días con backup fresco, "
         "test de gemelos, ensayo del rollback, capacitación de los 4 usuarios reales, "
         "runbook impreso."],
    ],
    widths=[4.2, 2.6, 9.0], size=9.5,
)

callout(
    "Estimación honesta",
    [("Si se ataca en el orden correcto y sin interrupciones, veo entre ", N), ("10 y 14 semanas", B),
     (" hasta el día D: 2 para cerrar la Fase 0 y consolidar el motor, 4 para la toolbar de Tráfico, "
      "2 para completar Reservas y el Graba, 2 para el ensayo general, y un colchón de 2 a 4 semanas "
      "para lo que siempre aparece. Las estimaciones del plan original (4 meses) siguen siendo "
      "razonables, con la diferencia de que hoy la mitad de las pantallas ya están construidas "
      "y la otra mitad del esfuerzo se fue a pantallas que el plan no contemplaba.", N)],
    fill=HEX_CLARO, color_titulo=AZUL, borde=HEX_AZUL,
)


# ═══════════════════════ 6. DESPUÉS DEL DÍA D ═══════════════════════
h1("Lo que falta DESPUÉS del Día D", "6")

p("Acá está, para mí, el aporte más útil de este informe. El backlog visible dice «105 ítems "
  "pendientes», y eso asusta y desenfoca. Fui a la base a preguntar cuáles de esos módulos "
  "todavía se usan. La respuesta cambia bastante el panorama.")

h2("Módulos MUERTOS — mi recomendación es declararlos fuera de alcance por escrito")

tabla(
    ["Módulo", "Evidencia en la base", "Formularios que ahorra"],
    [
        ["Taller completo (OT, stock, depósitos, chequeo de unidades, catálogos)",
         [("6.582 órdenes históricas, pero la última entrada al taller es del ", N),
          ("5 de agosto de 2019", B), (". ", N), mono("taller_stock"), (" y ", N),
          mono("taller_deposito"), (" tienen 0 filas: nunca se usaron.", N)],
         "24 formularios · 16 ítems de menú"],
        ["Cuenta corriente",
         [("Las 5 tablas (", N), mono("ctacte"), (", ", N), mono("ctacte_detalle"), (", ", N),
          mono("ctacte_pago"), ("…) tienen ", N), ("0 filas", B), (". Se programó y nunca se usó.", N)],
         "17 formularios"],
        ["Capacitaciones de choferes",
         [("417 registros, el último de ", N), ("junio de 2013", B), (".", N)],
         "2 formularios"],
        ["Apercibimientos y sanciones",
         [(mono("chofer_sancion")[0], {"font": MONO, "size": 9}), (" y ", N),
          (mono("chofer_sancion_motivo")[0], {"font": MONO, "size": 9}), (" con ", N),
          ("0 filas", B), (".", N)],
         "2 formularios"],
        ["Liquidación a Fleteros",
         [("12 liquidaciones tipo PROVEEDOR en todo el histórico, la última de ", N),
          ("diciembre de 2023", B), (".", N)],
         "El plan ya lo marcaba como «barato y casi sin uso»"],
        ["Peajes",
         [("6 formularios en el FoxPro y ", N), ("ninguna tabla correspondiente", B),
          (" en la réplica: es código muerto.", N)],
         "6 formularios"],
        ["Utilitarios de mantenimiento de DBF (indexar, compactar, normalizar)",
         "No tienen sentido contra SQL Server.",
         "~10 ítems de menú"],
    ],
    widths=[4.0, 8.0, 3.8], size=9.5,
)

callout(
    "Cuánto recorta esto",
    [("Entre 60 y 70 formularios del Metrocar y aproximadamente ", N), ("45 de los 105 ítems "
      "pendientes del menú", B), (". El proyecto pasa de «infinito» a «finito». Mi sugerencia "
      "concreta: llevarle esta tabla al dueño y pedirle que firme el descarte. Y en el menú de "
      "Buslink, cambiar el rótulo de esos ítems de «próximamente» a «no se migra — consultar en "
      "Metrocar», para no inflar la expectativa del cliente.", N)],
    fill=HEX_VERDE, color_titulo=VERDE, borde="1B6E3C",
)

h2("Módulos VIVOS que todavía no se tocaron")

tabla(
    ["Módulo", "Evidencia de que está vivo", "Prioridad sugerida"],
    [
        ["Adicionales Stock (menú Tráfico)",
         [("11.321 filas, ", N), ("4.576 desde 2025 y la última del 7 de julio de 2026", B),
          (". Es el módulo vivo más grande que nadie miró: no está migrado, no está documentado "
           "y no tiene plano.", N)],
         [("ALTA", R)]],
        ["Libro de Novedades (menú completo + correos)",
         [("48.353 filas, la última del 10 de julio de 2026. El alta por F2 ya está con andamiaje, "
           "pero el menú de mantenimiento, el envío de correos y sus parámetros no.", N)],
         [("MEDIA", {"bold": True, "color": NARANJ})]],
        ["Tarifarios de venta y de choferes",
         [(mono("lista_precio")[0], {"font": MONO, "size": 9}), (" con 2.791 filas y vigencias "
           "cargadas hasta ", N), ("octubre de 2026", B), ("; ", N),
          (mono("lista_precio_chofer")[0], {"font": MONO, "size": 9}), (" con 3.751. Buslink los "
           "LEE (el motor de tarifas depende de ellos) pero su ABM sigue en FoxPro.", N)],
         [("ALTA post-corte", R)]],
        ["Pantallas de Parámetros del sistema",
         [("La tabla ", N), mono("parametro"), (" es una de las 12 que cambian de dueño el día D, "
           "y su pantalla de mantenimiento no está migrada ni documentada. Después del corte, "
           "cambiar un parámetro no va a tener dónde hacerse.", N)],
         [("ALTA — es del día D", R)]],
        ["Catálogos del ABM del sistema (servicios, zonas, guías, dueños, permisos, IATA…)",
         [("Se usan todos los días desde las pantallas migradas, pero su mantenimiento sigue "
           "en FoxPro.", N)],
         [("MEDIA", {"bold": True, "color": NARANJ})]],
        ["Liquidación a Choferes",
         [("Adelantos hasta el 28/02/2025 y tarifario de choferes vivo. No está claro si sigue "
           "operando: ", N), ("hay que preguntarle al dueño", B), (".", N)],
         [("A CONFIRMAR", {"bold": True, "color": NARANJ})]],
        ["Chequeo de unidades (módulo Auditoría)",
         [("3.644 controles cargados y 71 ítems de chequeo definidos. No pude fechar el último "
           "por falta de columna de fecha de negocio: ", N), ("verificar con el dueño", B), (".", N)],
         [("A CONFIRMAR", {"bold": True, "color": NARANJ})]],
    ],
    widths=[3.8, 8.6, 3.4], size=9.5,
)

callout(
    "El más importante de esta lista",
    [("Las pantallas de Parámetros.", B), (" El día D la tabla ", N), mono("parametro"),
     (" pasa a ser propiedad de SQL, y el FoxPro queda bloqueado. Si nadie migró la pantalla que "
      "la edita, a partir de ese día no hay forma de cambiar un umbral de aviso, un contador o "
      "una configuración sin entrar a la base a mano. Es un ítem chico que hoy nadie tiene en "
      "el radar y que debería entrar en la Fase 1.", N)],
    fill=HEX_ROJO, color_titulo=ROJO, borde="B3261E",
)


# ═══════════════════════ 7. PANTALLAS FUERA DE BUSLINK ═══════════════════════
h1("Las pantallas que quedan fuera de Buslink", "7")

p("Este capítulo responde una pregunta concreta: ¿qué pantallas existen en el Metrocar y no "
  "existen en Buslink? Para contestarla no alcanzaba con mirar el menú lateral: fui al archivo "
  "de menú del FoxPro y extraje el formulario exacto que abre cada ítem, y recién ahí lo crucé "
  "contra las 44 rutas activas de Buslink.")

h2("El número, con la letra chica")

tabla(
    ["", "Cantidad", "Lectura"],
    [
        ["Ítems del menú Metrocar sin equivalente en Buslink", [("~102", B)],
         "El número que asusta"],
        ["De esos, módulos muertos o utilidades propias del FoxPro", [("~62", {"color": GRIS, "bold": True})],
         "No se migran: no tienen a quién servir"],
        ["Pantallas vivas que sí hay que migrar", [("~40", {"color": NARANJ, "bold": True})],
         "El trabajo real que queda"],
        ["De esas, el núcleo crítico", [("~20", {"color": ROJO, "bold": True})],
         "Bloquean el día D o son operación diaria"],
    ],
    widths=[7.4, 2.4, 6.0], size=9.5,
)

p("Dicho de otra manera: el menú del Metrocar tiene 271 entradas y el backlog visible de Buslink "
  "muestra 105 ítems deshabilitados, pero de todo eso hay unas veinte pantallas que de verdad "
  "importan. Ese es el trabajo, no el resto.")

h2("Grupo 1 — Críticas: bloquean el Día D o son de uso diario")

tabla(
    ["Módulo", "Pantalla del Metrocar", "Formulario", "Por qué es crítica"],
    [
        ["ABM del sistema", [("Parámetros Generales", B)], [mono("parametro")],
         [("La tabla ", N), mono("parametro"), (" cambia de dueño el día D. Sin esta pantalla, "
          "después del corte NO hay dónde editarla: contadores, umbrales de aviso, cliente "
          "interno y flags del GPS quedan sin interfaz.", N)]],
        ["ABM del sistema", "Parámetros Empresa", [mono("parametro_empresa")], "Ídem"],
        ["ABM del sistema", "Parámetros Pantalla Tráfico", [mono("parametro_trafico")], "Ídem"],
        ["ABM del sistema", "Parámetros SQL Server para GPS", [mono("parametro_sql_server")],
         "Ídem — y es donde viven los dos interruptores del GPS que hoy están apagados"],
        ["Tráfico", [("Adicionales Stock → Mantenimiento", B)], [mono("adicional_stock")],
         [("11.321 filas, última carga el ", N), ("07/07/2026", B), (". Es carga diaria: agua, "
          "hielo y propinas por unidad y cliente. Ni migrada ni documentada.", N)]],
        ["Tráfico", [("Adicionales Stock → Ingreso", B)], [mono("adicional_stock_abm")],
         "La pantalla de alta del mismo módulo"],
        ["Utilitarios", [("Elimina Lotes de Carga", B)], [mono("trafico_elimina_lote")],
         "El plan lo pide migrado ANTES del día D: es el botón de emergencia si el armado de "
         "plantillas genera viajes de más. Está enterrado entre las utilidades de reparación de DBF."],
        ["Sistema", [("Cambio de password", B)], [mono("usuario_cambio_password")],
         "Hoy en Buslink un usuario común no puede cambiarse su propia clave: solo el supervisor "
         "desde el ABM. Es una pantalla chica con impacto directo en la operación."],
    ],
    widths=[2.6, 4.4, 3.2, 5.6], size=9,
)

h2("Grupo 2 — Vivas, alta prioridad después del corte")

tabla(
    ["Módulo", "Pantallas", "Formularios", "Evidencia de que está viva"],
    [
        ["Tráfico", [("Libro de Novedades", B), (" (3): el libro · envío de correos · parámetros "
          "de correo", N)],
         [mono("libro_novedad"), (", ", N), mono("_envia_correo"), (", ", N), mono("_parametro")],
         "48.353 filas, última del 10/07/2026. El alta por F2 ya tiene andamiaje; el resto del "
         "módulo no está"],
        ["Facturación", [("Tarifario de Venta", B), (" (4): altas y copias · mantenimiento de "
          "precios · definición de lista · listadores", N)],
         [mono("lista_precio_cliente"), (", ", N), mono("_mantenimiento"), (", ", N),
          mono("lista_precio_modelo"), (", ", N), mono("_imprimir")],
         [("2.791 filas con vigencias cargadas ", N), ("hasta el 01/10/2026", B), (". Buslink las "
          "LEE (el motor de tarifas depende de ellas); el mantenimiento sigue en FoxPro.", N)]],
        ["Facturación", [("Adicionales", B), (" (6): tarifarios de venta (2) · tarifarios de pago "
          "(2) · adicionales · rubros", N)],
         [mono("adicional_lista_precio*"), (", ", N), mono("adicional_lista_pago*"), (", ", N),
          mono("adicional"), (", ", N), mono("adicional_rubro")],
         "27 adicionales, 329 precios de pago, 95 de venta. La solapa Adicionales de Liquidación "
         "los usa todos los días"],
        ["Reservas", [("Crear Plantillas", B)], [mono("reserva_plantilla_crear")],
         "El mantenimiento de plantillas ya está migrado; falta el alta de una plantilla nueva"],
        ["Sistema", [("Tablero de Control", B)], [mono("tablero"), (" + ", N), mono("tablero_zoom")],
         "Permiso X, solo el supervisor"],
    ],
    widths=[2.2, 4.6, 4.2, 4.8], size=9,
)

h2("Grupo 3 — Los catálogos del ABM del sistema (11 + 6)")

p("Son pantallas simples de código y nombre. Hoy Buslink los lee en combos y grillas, pero "
  "editarlos obliga a abrir el Metrocar.", size=10)

tabla(
    ["Catálogo", "Formulario", "Filas", "Nota"],
    [
        ["Servicios", [mono("servicio")], "62", "El catálogo central del negocio"],
        ["Cronogramas de Servicio", [mono("cronograma")], "97", ""],
        ["Guías", [mono("guia")], "1.135", "Tabla del día D — la escribe el alta de reservas"],
        ["IATA", [mono("iata")], "106", "Códigos de aeropuerto"],
        ["Feriados", [mono("feriado")], "15", [("Cero de 2026 cargados", {"color": ROJO, "bold": True})]],
        ["Motivos de Cancelación", [mono("viaje_motivo_cancela")], "6", "Lo necesita Cancelar (Fase 3)"],
        ["Motivos de Cambio de Cronograma", [mono("viaje_motivo_cambio")], "11", "Lo necesita Reasignar (Fase 3)"],
        ["Motivos de Llegadas Tardes", [mono("viaje_motivo_tarde")], "17", ""],
        ["Zonas · Dueños · Permisos", [mono("zona"), (" · ", N), mono("dueno"), (" · ", N), mono("permiso")],
         "6 · 2 · 14", "Zonas las escribe el cierre del viaje"],
        ["Configuración de Facturación (6)",
         [mono("empresa"), (", ", N), mono("moneda_tipo"), (", ", N), mono("moneda_cotizacion"),
          (", ", N), mono("ctacte_banco"), (", ", N), mono("ctacte_impuesto"), (", ", N),
          mono("ctacte_tipo_comprobante")],
         "32 cotizaciones", "Empresas para facturar, monedas, cotizaciones, bancos, impuestos y "
         "tipos de comprobante"],
    ],
    widths=[4.0, 4.6, 2.0, 5.2], size=9,
)

h2("Grupo 4 — A confirmar con el dueño antes de decidir")

tabla(
    ["Módulo", "Pantallas", "La duda concreta"],
    [
        ["Facturación", [("Tarifario de Choferes", B), (" (4) + ", N), ("Liquidación a Choferes", B),
          (" (5): genera · parámetros · ingreso, mantenimiento y motivo de adelantos", N)],
         [(mono("lista_precio_chofer")[0], {"font": MONO, "size": 9}), (" tiene 3.751 filas (vivo), "
          "pero ", N), (mono("chofer_adelanto")[0], {"font": MONO, "size": 9}),
          (" se detuvo el ", N), ("28/02/2025", B), (". ¿Se sigue liquidando a los choferes desde "
           "el sistema o pasó a otro circuito?", N)]],
        ["Utilitarios", "Km entre Localidades · Conectados al sistema",
         "La tabla de localidades existe. «Conectados» ya está semi-cubierto por la Auditoría de "
         "accesos y la tabla de sesiones que se construyeron en agosto"],
        ["Reservas", "Importa Reservas desde Excel",
         "Es la puerta 3 de la Fase 4 y el descarte declarado del plan. Hay que decidirlo ANTES "
         "del día D, no ese día"],
    ],
    widths=[2.2, 6.4, 7.2], size=9,
)

h2("Grupo 5 — Lo que recomiendo no migrar (~62 ítems)")

tabla(
    ["Módulo", "Ítems", "Evidencia dura"],
    [
        ["Taller completo", "16", "Última orden de trabajo: 05/08/2019. Stock y depósitos con 0 filas"],
        ["Chequeo de unidades (está dentro de Taller)", "4",
         [("Verificado en esta auditoría: la última de las 61 auditorías es del ", N),
          ("09/08/2019", B), (" — muerto igual que el resto del módulo", N)]],
        ["Apercibimientos", "2", "0 filas: nunca se usó"],
        ["Capacitaciones + Cursos-Descripción", "3", "Última: junio de 2013"],
        ["Liquidación a Fleteros", "1", "12 liquidaciones históricas, la última de 2023"],
        ["Nacionalidades · Profesiones", "2", "1 fila y 0 filas"],
        ["Utilitarios de sistema", "8", "Scheduler, agenda, calendario, calculadora, editor de "
         "imagen, backup (×2) y editor de log"],
        ["Reparaciones de DBF", "8", "Indexar, compactar, normalizar: no aplican a SQL Server "
         "(la excepción es «Elimina Lotes», que va en el grupo 1)"],
        ["Chat · Configuración regional · Log de errores · Servicio XML", "5",
         "El servicio XML es la vía del GPS, ya verificada como apagada"],
    ],
    widths=[6.0, 1.6, 8.2], size=9,
)

h2("Las tres cosas que este cruce puso sobre la mesa")

pasos([
    [("Las 4 pantallas de Parámetros son un agujero del Día D, no un anillo siguiente.", B),
     (" El día que se bloquea el FoxPro, la tabla ", N), mono("parametro"),
     (" queda sin interfaz. Y ahí viven los contadores del circuito, los umbrales de aviso, el "
      "cliente interno y los flags del GPS. Esto tiene que entrar en la Fase 1.", N)],
    [("Adicionales Stock es la mejor relación esfuerzo/valor que queda pendiente.", B),
     (" No toca el circuito ", N), mono("viaje"), (", no tiene riesgo, es una pantalla mediana, "
      "y hay alguien cargándola todos los días. Nadie la tenía en el radar porque no aparece en "
      "ninguno de los documentos del proyecto.", N)],
    [("El «Elimina Lotes de Carga» está escondido donde nadie lo va a buscar.", B),
     (" El plan lo nombra como el botón de emergencia del primer día, pero en el menú del Metrocar "
      "está enterrado entre las utilidades de reparación de archivos DBF. Es exactamente el tipo "
      "de ítem que se pasa por alto justo cuando más se necesita.", N)],
])

callout(
    "Y una definición de alcance que conviene tomar ahora",
    [("Los tarifarios —10 pantallas entre venta y adicionales— son, después del circuito ", N),
     mono("viaje"), (", la razón principal por la que el FoxPro no se va a poder apagar del todo "
      "el día D. Vale la pena ponerlos explícitamente como el anillo siguiente al corte, con fecha "
      "asignada, en vez de dejarlos en la bolsa genérica de «lo que falta».", N)],
    fill=HEX_CLARO, color_titulo=AZUL, borde=HEX_AZUL,
)


# ═══════════════════════ 8. DOCUMENTACIÓN FALTANTE ═══════════════════════
h1("Documentación faltante", "8")

p("La metodología del proyecto es «extraer antes de construir»: primero el plano del formulario "
  "FoxPro, después el código. Estos son los planos que faltan, ordenados por urgencia.")

h2("Bloqueantes del Día D (no son planos de pantallas, y por eso se postergaron)")

tabla(
    ["Documento que falta", "Por qué importa"],
    [
        ["El interruptor de sincronización DBF→SQL",
         "Es la palanca del corte y del rollback. Hay que documentar cómo se apaga tabla por tabla, "
         "quién lo opera, cuánto tarda en propagar y —lo más importante— qué hace la sincronización "
         "con filas que existen en SQL y no en los DBF. Sin esto, el día D es a ciegas."],
        ["El mecanismo de bloqueo del FoxPro",
         "Cómo dejar el Metrocar en solo consulta, probado en una copia. El runbook lo asume resuelto."],
        ["Mapeo campo a campo de las 12 tablas del circuito",
         "Contra INFORMATION_SCHEMA, incluyendo identidades, defaults y —crítico— cómo se asigna "
         "_sync_id en un INSERT hecho por Blazor. El plan lo llama show-stopper."],
        ["Plano de las 4 pantallas de Parámetros",
         "parametro cambia de dueño el día D y su mantenimiento no está migrado ni documentado."],
    ],
    widths=[5.4, 10.4], size=9.5,
)

h2("Planos de pantallas vivas que faltan")

bullet([mono("adicional_stock.scx"), (" + ", N), mono("adicional_stock_abm.scx"),
        (" — el módulo vivo sin documentar (11.321 filas, activo en julio 2026)", B)])
bullet([mono("libro_novedad.scx"), (", ", N), mono("_view"), (", ", N), mono("_envia_correo"), (", ", N),
        mono("_parametro"), (" — hoy solo está documentado el alta por F2", N)])
bullet([("Los 16 formularios de tarifarios (", N), mono("lista_precio*"), (") — hoy están cubiertos "
        "a nivel de módulo pero sin plano por pantalla, y son la próxima cosa que hay que migrar "
        "después del corte", N)])
bullet([mono("chofer_franco_modifica.scx"), (" y ", N), mono("chofer_franco_auditoria.scx"),
        (" — ya listados como pendientes", N)])
bullet([mono("trafico_liberar_hora_adicional.scx"), (" — el subdiálogo de horas extra del cierre, "
        "necesario para la operación 5 de la Fase 3", N)])
bullet([mono("viaje_motivo_tarde*.scx"), (" — 17 motivos cargados", N)])
bullet([("Catálogos del ABM del sistema: ", N), mono("servicio"), (", ", N), mono("zona"), (", ", N),
        mono("iata"), (", ", N), mono("dueno"), (", ", N), mono("permiso"), (", ", N),
        mono("cronograma"), (", ", N), mono("nacionalidad"), (", ", N), mono("profesion")])
bullet([("Facturación: monedas, cotizaciones, bancos, impuestos, tipos de comprobante, "
        "clientes-tarifas, clientes-descuentos", N)])
bullet([mono("tablero.scx"), (" — el Tablero de Control (permiso X)", N)])
bullet([mono("auditoria*.scx"), (" — el chequeo de unidades, si se confirma que sigue en uso", N)])

h2("Documentos que hay que actualizar, no crear")

bullet([mono("docs/Buslink/ANALISIS_SISTEMA_BUSLINK.md"), (" y su versión .docx — corte 02/07, "
        "quedaron a menos de la mitad de la realidad", N)])
bullet([mono("docs/Buslink/INFORME_AVANCE_BUSLINK.md"), (" — dice «Fase 0 en curso, 1 de 7 ítems»", N)])
bullet([mono("docs/Buslink/PLAN_MIGRACION_BUSLINK.md"), (" — la estrategia sigue siendo válida; hay que "
        "reflejar que Fases 3 y 4 tienen andamiaje construido y que la Fase 2 se desvió", N)])
bullet([("El índice ", N), mono("docs/PLANOFOXPRO/README.md"), (" — no incluye todavía los planos "
        "nuevos de agosto (menú Buses, cronograma, F2, F4)", N)])


# ═══════════════════════ 9. RIESGOS Y DEUDAS ═══════════════════════
h1("Riesgos y deudas técnicas", "9")

p("Esto es lo que encontré mirando el código y la base con ojo de auditoría. Ordenado por lo "
  "que más me preocuparía si el proyecto fuera mío.")

h2("Riesgo 1 — Tres semanas de trabajo sin respaldo")
rich([("El último commit es del 21 de julio. Hoy hay 44 archivos modificados y 31 archivos nuevos "
       "sin versionar, incluyendo servicios enteros (", N), mono("OcupacionFlota.cs"), (", ", N),
      mono("AccesoRedService.cs"), (", ", N), mono("SesionCircuitoTracker.cs"), ("), 18 diálogos "
      "nuevos y 5 planos FoxPro. Un disco que falla hoy borra tres semanas. ", N),
      ("Esto se arregla en diez minutos y es lo primero de la lista.", B)])

h2("Riesgo 2 — 25 escrituras codificadas que nunca se ejecutaron")
p("Hay 4.780 líneas de lógica de escritura escrita contra los planos y nunca ejecutada contra una "
  "base real. Los planos son buenos, pero un INSERT de 35 campos con desnormalizados calculados "
  "no se valida leyendo: se valida corriéndolo. Cuanto más crece esa masa sin ejecutarse, más "
  "grande es la sorpresa el día que se encienden todos juntos.")
rich([("Mi recomendación: ", B), ("levantar un entorno de pruebas con un backup fresco de producción "
      "y encender los flags ahí, uno por uno, con el protocolo de dos señales. No hace falta esperar "
      "al día D para probar el código: hace falta esperar al día D para ", N), ("usarlo en producción", B),
      (". Son dos cosas distintas y hoy están mezcladas.", N)])

h2("Riesgo 3 — Credenciales en el repositorio")
rich([("El ", N), mono("appsettings.json"), (" versionado tiene usuario ", N), mono("sa"),
      (" y contraseña en texto plano, y lo mismo aparece en ", N), mono("CLAUDE.md"),
      (". Además, las contraseñas de los usuarios se guardan en texto plano en la tabla ", N),
      mono("usuario"), (" (herencia del FoxPro, campo de 15 caracteres). Mientras el sistema viva "
      "en la LAN es un riesgo acotado; el día que se active el acceso por Internet —que ya está "
      "codificado y esperando— pasa a ser un problema serio.", N)])
bullet("Corto plazo: sacar la cadena de conexión del repo (variables de entorno o User Secrets) y usar un login SQL con permisos acotados, no sa.")
bullet("Antes de exponer a Internet: hashear las contraseñas de usuario (con migración transparente en el primer login).")

h2("Riesgo 4 — Inconsistencia en a qué servidor apunta la aplicación")
rich([("El ", N), mono("CLAUDE.md"), (" dice que el servidor nuevo es ", N), mono("172.25.69.217"),
      ("; el ", N), mono("appsettings.json"), (" tiene comentada la IP ", N), mono("172.25.80.234"),
      ("; y la conexión activa hoy es la del servidor local ", N), mono("DESKTOP-CV6LF0O\\SQLEXPRESS"),
      (". Tres verdades distintas para la misma cosa. Antes del corte esto tiene que quedar en un "
       "solo lugar y sin ambigüedad.", N)])

h2("Riesgo 5 — La réplica local está desfasada")
p("Los datos más recientes de la réplica local son de alrededor del 10 de julio de 2026. Hoy es "
  "9 de agosto: hay un mes de operación que no está. Para la operación sombra de la Fase 6 hace "
  "falta un backup fresco de producción, y conviene pedirlo con tiempo.")

h2("Otras deudas técnicas")

tabla(
    ["Deuda", "Impacto", "Sugerencia"],
    [
        [[mono("Command Timeout=0"), (" en la cadena de conexión", N)],
         "Una consulta mal formada se cuelga para siempre y se lleva puesto el circuito SignalR del usuario",
         "Poner un timeout razonable (30–60 s) y manejar la excepción con un mensaje claro"],
        ["SQL Server 2012 fuera de soporte",
         "Sin parches de seguridad desde julio de 2022. Además obliga a escribir SQL antiguo.",
         "Plantearle al cliente la actualización como parte del proyecto, no después"],
        ["Cobertura de tests baja",
         "4 archivos de Playwright y cero tests unitarios del motor de tarifas o de la escritura. "
         "Para transacciones sobre el circuito viaje es poco.",
         "Tests de integración con transacción y rollback automático: se prueba contra la base real "
         "sin ensuciarla. Es la red de seguridad que falta para la Fase 3."],
        ["Clases muy grandes",
         [(mono("ReportService.cs")[0], {"font": MONO, "size": 9}), (" con 8.721 líneas y ", N),
          (mono("AbmService.cs")[0], {"font": MONO, "size": 9}), (" con 4.780. Todavía manejable, "
           "pero la Fase 3 le suma bastante.", N)],
         "Partirlas por módulo (clases parciales o un servicio por dominio) antes de empezar Tráfico"],
        ["Blanqueo de la grilla de Tráfico al scrollear rápido",
         "Problema conocido, documentado, con dos intentos de solución fallidos",
         "Está bien dejarlo abierto: molesta pero no bloquea. No volver a intentarlo sin leer el documento."],
        ["Pantalla huérfana",
         "«Liquidaciones estimadas» existe y funciona pero no tiene link en el menú lateral",
         "Agregar el ítem al menú de Facturación — es un renglón"],
        ["Sin integración continua ni backup automatizado del servidor nuevo",
         "El día D, ese servidor pasa a ser el único original de la operación",
         "Backup diario verificado con restore de prueba, desde antes del corte"],
    ],
    widths=[4.0, 6.6, 5.2], size=9.5,
)


# ═══════════════════════ 10. IDEAS ═══════════════════════
h1("Ideas y recomendaciones", "10")

p("Todo lo de este capítulo es opinión mía, marcada como tal. Son cosas que no están en el plan "
  "y que creo que valen la pena.")

h2("1. Separar «probar la escritura» de «usar la escritura»")
p("Es la idea más valiosa que se me ocurre. Hoy los 25 flags están apagados porque encenderlos "
  "en producción sería peligroso. Pero eso mezcló dos decisiones distintas: si el código funciona, "
  "y si el negocio ya opera con él. Con un entorno de pruebas con backup fresco, los flags se pueden "
  "encender ahí y validar todo con datos reales, sin ningún riesgo. Cuando llegue el día D, el "
  "código llegaría probado en vez de estrenado.")

h2("2. Un tablero de salud del corte, dentro de Buslink")
p("El plan tiene ocho consultas de monitoreo para la primera semana después del corte. En vez de "
  "que sean scripts sueltos que alguien corre a mano, convertirlas en una pantalla del sistema "
  "(visible solo para el supervisor): bitácora sin huérfanos, consistencia entre viaje y vehículo, "
  "contadores monótonos, desnormalizados sanos, detección de escritura fantasma de la sincronización. "
  "Con semáforos. El día D uno mira esa pantalla en vez de abrir el SQL Management Studio.")

h2("3. Aprovechar lo que ya se construyó para el propio corte")
bullet([("La ", N), ("auditoría de accesos", B), (" y la tabla de sesiones ya permiten saber quién "
        "está conectado. El día D, antes de cortar, sirve para verificar que no quedó nadie operando.", N)])
bullet([("El ", N), ("centro de ayuda", B), (" ya publica manuales en PDF. La capacitación previa al "
        "día D puede entregarse ahí en vez de por mail: queda dentro del sistema y siempre a mano.", N)])
bullet([("El ", N), ("libro de novedades", B), (" (F2) es el lugar natural para que el despacho "
        "registre cualquier cosa rara durante la primera semana post-corte.", N)])

h2("4. Un banner que diga quién manda")
p("Mientras dure la transición, una barra fija arriba: «Buslink en modo consulta — la operación se "
  "carga en Metrocar» o «Buslink OPERATIVO — no cargar en Metrocar», según el estado del flag global. "
  "Los errores más caros de una migración no son de código: son dos personas cargando en dos sistemas "
  "distintos porque nadie les dijo cuál estaba vivo. Y del lado del Metrocar, cambiarle el título de "
  "la ventana a «SOLO CONSULTA» el día del corte.")

h2("5. Cerrar el alcance por escrito")
p("Con la evidencia de los capítulos 6 y 7 en la mano, pedirle al dueño una firma sobre tres cosas: qué "
  "módulos se declaran muertos y no se migran, si la importación desde Excel entra o no en el "
  "alcance del día 1, y qué se hace con la integración GPS. Son las tres decisiones que hoy están "
  "abiertas y que cada tanto vuelven a discutirse. Firmadas, dejan de volver.")

h2("6. Cambiar el rótulo del backlog visible")
p("Que el menú muestre 105 ítems «próximamente» le dice al cliente que faltan 105 cosas. Si unos 45 "
  "son módulos muertos, conviene distinguirlos: «no se migra — consultar en Metrocar». La misma "
  "pantalla, contando una historia más cercana a la verdad.")

h2("7. Empezar por Adicionales Stock cuando haya un hueco")
p("Es una pantalla de tamaño mediano, viva, de bajo riesgo (no toca el circuito viaje), sin "
  "documentar y que nadie tiene en el radar. Es el candidato ideal para una semana en la que haya "
  "que esperar una definición del cliente sobre otra cosa.")

h2("8. Sobre el motor de tarifas")
p("Está validado al 99,4% contra 8.656 viajes. Ese 0,6% que no coincide son unos 50 casos. "
  "Antes del Graba conviene mirarlos uno por uno y clasificarlos: si son bugs del FoxPro que Buslink "
  "corrige, hay que documentarlo y avisarle al cliente que los números van a cambiar (poco, pero "
  "van a cambiar). Si son casos que Buslink calcula mal, hay que arreglarlos antes de grabar. "
  "Hoy están en la misma bolsa y son dos cosas muy distintas.")


# ═══════════════════════ 11. PLAN SUGERIDO ═══════════════════════
h1("Plan sugerido para las próximas semanas", "11")

p("Este es el orden en que yo lo haría, con el criterio de que cada semana cierre algo verificable "
  "y que lo que desbloquea vaya antes que lo que agrega.")

h2("Semana 1 — Higiene y bloqueantes")
pasos([
    [("Commitear y subir todo.", B), (" Tres semanas de trabajo sin respaldo. Antes que nada.", N)],
    [("Sacar las credenciales del repositorio", B), (" y unificar a qué servidor apunta la aplicación.", N)],
    [("Cerrar el permiso F en el Zoom del Viaje", B), (" — es la deuda de seguridad más concreta y "
      "es media hora de trabajo (líneas 193 y 198 de ", N), mono("ZoomViajeDialog.razor"), (").", N)],
    [("Agregar el link faltante de «Liquidaciones estimadas» al menú.", N)],
    [("Pedirle al cliente", B), (": el backup fresco de producción, la réplica de ", N),
     mono("viaje_log_chofer"), (", la verificación de las 5 tablas en el servidor nuevo, y la "
     "carga de los feriados de 2026 en el FoxPro.", N)],
])

h2("Semana 2 — Las incógnitas del corte")
pasos([
    [("Documentar el interruptor de sincronización", B), (" — con el operador de la réplica, en vivo, "
      "y probando qué pasa con una fila escrita solo en SQL.", N)],
    [("Documentar y probar el bloqueo del FoxPro", B), (" en una copia.", N)],
    [("Mapeo campo a campo de las 12 tablas", B), (", con la resolución de ", N), mono("_sync_id"),
     (" verificada contra el catálogo del sistema, no supuesta.", N)],
    [("Llevarle al dueño los capítulos 6 y 7 de este informe", N), (" y cerrar el alcance por escrito.", N)],
])

h2("Semana 3 — Consolidar el motor")
pasos([
    [("Unificar las primitivas de escritura del circuito", B), (": la bitácora escrita una sola vez, "
      "los contadores, las transiciones de estado, la firma obligatoria con fecha de reserva.", N)],
    [("Montar el entorno de pruebas con backup fresco", B), (" y estrenar el protocolo de dos señales "
      "encendiendo ahí los flags ya construidos, empezando por el cambio de cronograma (que es la "
      "operación más frecuente y la de menor riesgo).", N)],
    [("Escribir los primeros tests de integración con rollback", N), (" — la red de seguridad de la Fase 3.", N)],
])

h2("Semanas 4 a 7 — Tráfico en escritura")
p("Las diez operaciones en el orden del plan: chequeo primero (valida la tubería completa con "
  "riesgo mínimo), después asignar y liberar juntos (se prueban mutuamente), después reasignar, "
  "después el cierre del viaje, después cancelar y reactivar, y el Zoom en edición al final. "
  "Una operación por entrega, cada una verificada con dos señales.")

h2("Semanas 8 y 9 — Cerrar Reservas y Facturación")
p("El «deshacer lote» de plantillas —que es el botón de emergencia del primer día— y el Graba "
  "transaccional con su test de cuadre contra las últimas tres liquidaciones reales. Decidir con "
  "el dueño si la importación desde Excel entra o queda para después.")

h2("Semanas 10 a 12 — Ensayo general")
p("Feature flag global, operación sombra de tres a cinco días con diff diario, test de gemelos, "
  "ensayo del rollback completo, capacitación de los cuatro usuarios reales con su matriz de "
  "permisos, y el runbook impreso. Recién cuando esto da limpio se elige la fecha del corte.")

callout(
    "El criterio para elegir la fecha",
    "No antes de que los 18 criterios de «listo» del plan estén tildados, y sobre el día de menor "
    "volumen histórico de viajes, con la noche anterior como ventana. Y con el dev disponible toda "
    "esa semana, no solo ese día.",
    fill=HEX_CLARO, color_titulo=AZUL, borde=HEX_AZUL,
)


# ═══════════════════════ ANEXO A ═══════════════════════
h1("Anexo A — Vigencia real de cada módulo", "A")

p("Consultas hechas sobre replicaVPF el 09/08/2026. Las fechas son de negocio (no de la "
  "sincronización), que es lo único que dice si un módulo se usa de verdad.")

tabla(
    ["Tabla", "Filas", "Última actividad real", "Lectura"],
    [
        [[mono("viaje")], "533.483", "Reservas cargadas hasta 10/05/2027", [("Núcleo", V)]],
        [[mono("viaje_log")], "4.513.301", "Bitácora del circuito", [("Núcleo", V)]],
        [[mono("viaje_adicional")], "160.498", "—", [("Núcleo", V)]],
        [[mono("liquidacion")], "4.245", "Cliente: 07/07/2026 · Proveedor: 21/12/2023", [("Vivo / fletero muerto", {"color": NARANJ})]],
        [[mono("liquidacion_detalle")], "828.960", "—", [("Vivo", V)]],
        [[mono("vehiculo_sobre")], "110.239", "Cargas de combustible activas", [("Vivo", V)]],
        [[mono("libro_novedad")], "48.353", "10/07/2026", [("Vivo, sin migrar", {"color": ROJO})]],
        [[mono("chofer_franco")], "72.079", "Francos activos", [("Vivo", V)]],
        [[mono("adicional_stock")], "11.321", "07/07/2026", [("Vivo, sin migrar ni documentar", {"color": ROJO, "bold": True})]],
        [[mono("cliente_grupo")], "11.574", "—", [("Vivo", V)]],
        [[mono("vehiculo_km")], "10.638", "Odómetros", [("Vivo", V)]],
        [[mono("taller_service")], "6.582", "05/08/2019", [("MUERTO", {"color": GRIS, "bold": True})]],
        [[mono("taller_service_item")], "7.979", "—", [("MUERTO", {"color": GRIS, "bold": True})]],
        [[mono("taller_stock")], "0", "Nunca se usó", [("MUERTO", {"color": GRIS, "bold": True})]],
        [[mono("taller_deposito")], "0", "Nunca se usó", [("MUERTO", {"color": GRIS, "bold": True})]],
        [[mono("ctacte")], "0", "Nunca se usó (las 5 tablas)", [("MUERTO", {"color": GRIS, "bold": True})]],
        [[mono("chofer_curso")], "417", "19/06/2013", [("MUERTO", {"color": GRIS, "bold": True})]],
        [[mono("chofer_sancion")], "0", "Nunca se usó", [("MUERTO", {"color": GRIS, "bold": True})]],
        [[mono("chofer_adelanto")], "375", "28/02/2025", [("Dudoso — preguntar", {"color": NARANJ})]],
        [[mono("lista_precio")], "2.791", "Vigencias hasta 01/10/2026", [("Vivo, ABM en FoxPro", {"color": ROJO})]],
        [[mono("lista_precio_chofer")], "3.751", "—", [("Vivo, ABM en FoxPro", {"color": ROJO})]],
        [[mono("auditoria_control")], "3.644", "Sin fecha de negocio — verificar", [("A confirmar", {"color": NARANJ})]],
        [[mono("chofer_viatico")], "0", "Nunca se usó", [("Migrado igual (andamiaje)", G)]],
        [[mono("vehiculo_chofer")], "0", "Tabla vacía en la réplica", [("Conocido", G)]],
        [[mono("viaje_log_chofer")], "—", "No existe en SQL (75.001 filas en el DBF)", [("BLOQUEANTE", {"color": ROJO, "bold": True})]],
    ],
    widths=[4.2, 2.2, 5.6, 3.8], size=9,
)


# ═══════════════════════ ANEXO B ═══════════════════════
h1("Anexo B — Checklist de arranque inmediato", "B")

p("Lo que yo haría esta semana, sin esperar a nadie.")

h2("Hoy mismo (una hora de trabajo)")
bullet([("Commit y push de todo lo pendiente", B), (" — 44 modificados + 31 nuevos", N)])
bullet([("Agregar ", N), mono("appsettings.json"), (" con credenciales al ", N), mono(".gitignore"),
        (" y mover la cadena de conexión a variables de entorno", N)])
bullet([("Sumar el link de «Liquidaciones estimadas» al menú de Facturación", N)])

h2("Esta semana")
bullet([("Gate del permiso F en el Zoom del Viaje", B), (" (líneas 193 y 198)", N)])
bullet([("Revisar todas las grillas con importes y aplicarles el mismo criterio", N)])
bullet([("Actualizar ", N), mono("docs/PLANOFOXPRO/README.md"), (" con los planos de agosto", N)])

h2("Pedidos al cliente — mandarlos juntos en un solo mail")
pasos([
    "Backup fresco de la base de producción, para el entorno de pruebas y la operación sombra.",
    "Replicar la tabla viaje_log_chofer (75.001 filas) — hoy bloquea el logoneo de conductores.",
    "Verificar que cabecera, chofer_franco, chofer_viatico y sus dos catálogos existan en el servidor nuevo.",
    "Cargar en el Metrocar los feriados de 2026 que faltan (hay cero cargados y el armado de plantillas los ignora).",
    "Una reunión de 30 minutos con quien opera la sincronización DBF→SQL, para documentar el interruptor.",
    "Decisión firmada sobre la integración GPS (la evidencia dice que está apagada hace años).",
    "Confirmar si la liquidación a choferes y el chequeo de unidades siguen en uso.",
    "El índice por id_viaje en la tabla viaje, con el argumento nuevo: ahora va a haber modificaciones, no solo consultas.",
])

h2("Preguntas abiertas para el dueño")
bullet("¿Firmamos que Taller, Cuenta Corriente, Capacitaciones, Apercibimientos y Liquidación a Fleteros no se migran?")
bullet("¿La importación de reservas desde Excel entra en el alcance del día 1 o queda para después?")
bullet("¿Quién es el usuario responsable de cada módulo el día del corte, para la capacitación?")
bullet("¿Hay presupuesto y ventana para actualizar SQL Server 2012, que está sin soporte desde 2022?")


# ── pie ──
doc.add_paragraph()
par = doc.add_paragraph()
border(par, "top", size=8, color="D5DAE3", space=6)
par.paragraph_format.space_before = Pt(14)
run(par, "Buslink · Metrocar Nortur — Informe de estado de la migración · 9 de agosto de 2026",
    size=9, color=GRIS)
par2 = doc.add_paragraph()
run(par2, "Fuentes verificadas: documentos maestros del proyecto, 30 planos FoxPro, 13 skills, "
          "código fuente de Buslink (49 páginas · 68 componentes · ReportService y AbmService), "
          "historial de git, menús y formularios del Metrocar (C:\\MetroCarSys) y consultas directas "
          "a la base replicaVPF.", size=9, color=GRIS)

for sec in doc.sections:
    footer_par = sec.footer.paragraphs[0]
    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_par.add_run()
    r.font.size = Pt(8.5)
    r.font.name = BODY
    r.font.color.rgb = GRIS
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    r._r.append(fld1); r._r.append(instr); r._r.append(fld2)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "INFORME_ESTADO_MIGRACION.docx")
doc.save(out)
print("OK ->", out)
