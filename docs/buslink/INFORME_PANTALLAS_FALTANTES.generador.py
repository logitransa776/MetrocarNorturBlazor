# -*- coding: utf-8 -*-
"""
Genera INFORME_PANTALLAS_FALTANTES.docx — el cruce ítem por ítem del menú del Metrocar
(FoxPro) contra las rutas activas de Buslink: qué pantallas quedan afuera, de qué módulo,
si el módulo está vivo o muerto, y el backlog priorizado.

Método: se parseó C:\\MetroCarSys\\Menus\\MENU_PRINCIPAL.MPR extrayendo el formulario .scx
exacto que abre cada entrada del menú (271 entradas), y se cruzó contra las 44 rutas del
drawer de Buslink + las 49 páginas .razor. La vigencia de cada módulo se midió con
consultas directas a replicaVPF usando fechas de NEGOCIO (no metadata de la réplica).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

AZUL   = RGBColor(0x00, 0x3A, 0xA0)
NARANJ = RGBColor(0xF9, 0x94, 0x10)
GRIS   = RGBColor(0x5A, 0x60, 0x6A)
ROJO   = RGBColor(0xB3, 0x26, 0x1E)
VERDE  = RGBColor(0x1B, 0x6E, 0x3C)
NEGRO  = RGBColor(0x1F, 0x24, 0x2C)

HEX_AZUL   = "003AA0"
HEX_NARANJ = "F99410"
HEX_CLARO  = "EEF2F9"
HEX_SUAVE  = "F7F8FA"
HEX_AVISO  = "FFF4E3"
HEX_ROJO   = "FDECEA"
HEX_VERDE  = "EAF5EE"
HEX_MUERTO = "F0F0F0"

BODY = "Segoe UI"
MONO = "Consolas"

doc = Document()

st = doc.styles["Normal"]
st.font.name = BODY
st.font.size = Pt(10.5)
st.font.color.rgb = NEGRO
st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.15

for sec in doc.sections:
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.9)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)


def shade(el, hexcolor):
    el.append(parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls("w"), hexcolor)))


def border(par, edge="bottom", size=12, color=HEX_NARANJ, space=4):
    pPr = par._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    e = OxmlElement("w:" + edge)
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(size))
    e.set(qn("w:space"), str(space))
    e.set(qn("w:color"), color)
    pBdr.append(e)


def run(par, text, bold=False, italic=False, size=10.5, color=None, font=BODY, hl=None):
    r = par.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color is not None:
        r.font.color.rgb = color
    if hl:
        shade(r._element.get_or_add_rPr(), hl)
    return r


def p(text="", size=10.5, bold=False, italic=False, color=None, space_after=6,
      space_before=0, align=None, left=0):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    if left:
        par.paragraph_format.left_indent = Cm(left)
    if align:
        par.alignment = align
    if text:
        run(par, text, bold=bold, italic=italic, size=size, color=color)
    return par


def rich(parts, space_after=6, space_before=0, left=0, size=10.5):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    if left:
        par.paragraph_format.left_indent = Cm(left)
    for txt, fmt in parts:
        f = dict(fmt)
        f.setdefault("size", size)
        run(par, txt, **f)
    return par


def h1(text, nro=None):
    par = doc.add_paragraph()
    par.paragraph_format.page_break_before = True
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.keep_with_next = True
    if nro:
        run(par, nro + "  ", bold=True, size=17, color=NARANJ)
    run(par, text, bold=True, size=17, color=AZUL)
    border(par, "bottom", size=10, color=HEX_NARANJ, space=3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def h2(text, color=AZUL):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.keep_with_next = True
    run(par, text, bold=True, size=13, color=color)
    return par


def h3(text):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(9)
    par.paragraph_format.space_after = Pt(2)
    par.paragraph_format.keep_with_next = True
    run(par, text, bold=True, size=11, color=GRIS)
    return par


def bullet(text_parts, level=0):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(0.75 + 0.6 * level)
    if isinstance(text_parts, str):
        run(par, text_parts)
    else:
        for txt, fmt in text_parts:
            f = dict(fmt)
            f.setdefault("size", 10.5)
            run(par, txt, **f)
    return par


def pasos(items):
    for i, text_parts in enumerate(items, start=1):
        par = doc.add_paragraph()
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.left_indent = Cm(1.05)
        par.paragraph_format.first_line_indent = Cm(-0.6)
        run(par, "{}.\t".format(i), bold=True, color=NARANJ)
        if isinstance(text_parts, str):
            run(par, text_parts)
        else:
            for txt, fmt in text_parts:
                f = dict(fmt)
                f.setdefault("size", 10.5)
                run(par, txt, **f)


def _cell_margins(cell, lr=110, tb=70):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", tb), ("left", lr), ("bottom", tb), ("right", lr)):
        e = OxmlElement("w:" + side)
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)


def _tbl_borders(tbl, color, left_only=False, size=4):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    sides = ["left"] if left_only else ["top", "left", "bottom", "right", "insideH", "insideV"]
    todos = ["top", "left", "bottom", "right", "insideH", "insideV"]
    for s in todos:
        e = OxmlElement("w:" + s)
        if s in sides:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(24 if left_only else size))
            e.set(qn("w:color"), color)
        else:
            e.set(qn("w:val"), "none")
        e.set(qn("w:space"), "0")
        borders.append(e)
    tblPr.append(borders)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:tblHeader")
    e.set(qn("w:val"), "true")
    trPr.append(e)


def callout(titulo, texto, fill=HEX_AVISO, color_titulo=NARANJ, borde=HEX_NARANJ):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), fill)
    cell.text = ""
    par = cell.paragraphs[0]
    par.paragraph_format.space_after = Pt(2)
    run(par, titulo, bold=True, size=10.5, color=color_titulo)
    par2 = cell.add_paragraph()
    par2.paragraph_format.space_after = Pt(0)
    if isinstance(texto, str):
        run(par2, texto, size=10)
    else:
        for txt, fmt in texto:
            f = dict(fmt)
            f.setdefault("size", 10)
            run(par2, txt, **f)
    _cell_margins(cell, 140, 110)
    _tbl_borders(t, borde, left_only=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def tabla(headers, filas, widths=None, size=9.5, header_fill=HEX_AZUL, zebra=True, fills=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0]
    for i, htxt in enumerate(headers):
        c = hdr.cells[i]
        if header_fill:
            shade(c._tc.get_or_add_tcPr(), header_fill)
        c.text = ""
        par = c.paragraphs[0]
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        run(par, htxt, bold=True, size=size, color=RGBColor(0xFF, 0xFF, 0xFF))
        _cell_margins(c, 90, 55)
    _repeat_header(hdr)
    for j, fila in enumerate(filas):
        row = t.add_row()
        rowfill = fills[j] if fills and j < len(fills) and fills[j] else None
        for i, celda in enumerate(fila):
            c = row.cells[i]
            if rowfill:
                shade(c._tc.get_or_add_tcPr(), rowfill)
            elif zebra and j % 2 == 1:
                shade(c._tc.get_or_add_tcPr(), HEX_SUAVE)
            c.text = ""
            par = c.paragraphs[0]
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.space_before = Pt(0)
            if isinstance(celda, str):
                run(par, celda, size=size)
            else:
                for txt, fmt in celda:
                    f = dict(fmt)
                    f.setdefault("size", size)
                    run(par, txt, **f)
            _cell_margins(c, 90, 55)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    if header_fill is None:
        t._tbl.remove(t.rows[0]._tr)
    _tbl_borders(t, "D5DAE3", size=4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def ficha(titulo, modulo, form, prioridad, color=ROJO, fill=HEX_ROJO):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade(cell._tc.get_or_add_tcPr(), fill)
    cell.text = ""
    par = cell.paragraphs[0]
    par.paragraph_format.space_after = Pt(1)
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.keep_with_next = True
    run(par, titulo, bold=True, size=13, color=NEGRO)
    run(par, "     " + prioridad + " ", bold=True, size=8,
        color=RGBColor(0xFF, 0xFF, 0xFF), hl=("B3261E" if color == ROJO else "C9741A"))
    par2 = cell.add_paragraph()
    par2.paragraph_format.space_after = Pt(0)
    par2.paragraph_format.keep_with_next = True
    run(par2, modulo, size=9, color=GRIS)
    run(par2, "     " + form, size=9, font=MONO, color=AZUL)
    _cell_margins(cell, 140, 100)
    _tbl_borders(t, ("B3261E" if color == ROJO else HEX_NARANJ), left_only=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def campo(etiqueta, texto):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(3)
    par.paragraph_format.left_indent = Cm(0.3)
    run(par, etiqueta + "   ", bold=True, size=9, color=NARANJ)
    if isinstance(texto, str):
        run(par, texto, size=10)
    else:
        for txt, fmt in texto:
            f = dict(fmt)
            f.setdefault("size", 10)
            run(par, txt, **f)
    return par


def mono(txt):
    return (txt, {"font": MONO, "size": 9.5, "color": AZUL})


def mono8(txt):
    return (txt, {"font": MONO, "size": 8.5, "color": AZUL})


B = {"bold": True}
N = {}
G = {"color": GRIS}

# marcas de estado para el anexo
M_OK  = [("MIGRADO", {"bold": True, "color": VERDE, "size": 8})]
M_AND = [("ANDAMIAJE", {"bold": True, "color": NARANJ, "size": 8})]
M_FAL = [("FALTA", {"bold": True, "color": ROJO, "size": 8})]
M_NO  = [("NO MIGRAR", {"bold": True, "color": GRIS, "size": 8})]


# ═══════════════════════════ PORTADA ═══════════════════════════
for _ in range(3):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(0)
run(par, "BUSLINK  ·  METROCAR NORTUR", bold=True, size=11, color=NARANJ)
par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(0)
run(par, "Auditoría de cobertura funcional", size=11, color=GRIS)

par = doc.add_paragraph()
par.paragraph_format.space_before = Pt(18)
par.paragraph_format.space_after = Pt(0)
run(par, "Qué pantallas del", bold=True, size=30, color=AZUL)
par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(0)
run(par, "Metrocar todavía", bold=True, size=30, color=AZUL)
par = doc.add_paragraph()
par.paragraph_format.space_after = Pt(4)
run(par, "no están en Buslink", bold=True, size=30, color=NARANJ)
border(par, "bottom", size=18, color=HEX_NARANJ, space=8)

par = doc.add_paragraph()
par.paragraph_format.space_before = Pt(14)
run(par, "El cruce ítem por ítem del menú del sistema viejo contra el sistema nuevo: qué falta, "
         "de qué módulo, si ese módulo todavía se usa, y en qué orden conviene atacarlo.",
    size=12, color=GRIS)

for _ in range(2):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

tabla(
    ["", ""],
    [
        [[("Pregunta que responde", B)], "¿Qué pantallas existen en Metrocar y no existen en Buslink?"],
        [[("Método", B)], "Se extrajo el formulario .scx exacto detrás de cada una de las 271 entradas "
                          "del menú del FoxPro y se cruzó contra las 44 rutas activas de Buslink. "
                          "La vigencia de cada módulo se midió con consultas a la base usando fechas "
                          "de negocio."],
        [[("Corte", B)], "9 de agosto de 2026"],
        [[("Documento hermano", B)], "«Estado real y qué falta para terminar» — el informe general "
                                     "del proyecto. Este documento profundiza su capítulo 7."],
    ],
    widths=[3.6, 12.4], size=10, header_fill=None, zebra=True,
)


# ═══════════════════════════ CONTENIDO ═══════════════════════════
h1("Contenido")

contenido = [
    ("1.", "El resultado en una página", "Los cuatro números que hay que retener"),
    ("2.", "Cómo se hizo este cruce", "Por qué no alcanzaba con mirar el menú lateral"),
    ("3.", "Grupo 1 — Las críticas", "8 pantallas que bloquean el corte o son de uso diario"),
    ("4.", "Grupo 2 — Vivas, alta prioridad", "15 pantallas del anillo siguiente"),
    ("5.", "Grupo 3 — Los catálogos", "17 pantallas chicas del ABM del sistema"),
    ("6.", "Grupo 4 — A confirmar con el dueño", "9 pantallas con la vigencia en duda"),
    ("7.", "Grupo 5 — Lo que no conviene migrar", "62 ítems, con la evidencia de cada uno"),
    ("8.", "Backlog priorizado", "El orden de trabajo, con estimación"),
    ("A.", "Anexo — El menú completo del Metrocar", "Los 9 módulos, ítem por ítem, con su estado"),
]
for nro, tit, sub in contenido:
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(5)
    par.paragraph_format.left_indent = Cm(0.2)
    run(par, nro + "  ", bold=True, size=11, color=NARANJ)
    run(par, tit, bold=True, size=11, color=AZUL)
    if sub:
        run(par, "   " + sub, size=10, color=GRIS)

callout(
    "Para qué sirve cada parte",
    [("Si tenés cinco minutos, leé el capítulo 1 y el 8. Si vas a una reunión con el dueño, "
      "llevá el capítulo 7 (lo que no se migra, con la evidencia) y el 8 (el orden y el costo). "
      "El anexo A es el ", N), ("checklist imprimible", B), (": el menú entero del Metrocar con "
      "una marca de estado por ítem, para tildar a medida que se avanza.", N)],
    fill=HEX_CLARO, color_titulo=AZUL, borde=HEX_AZUL,
)


# ═══════════════════════ 1. RESULTADO ═══════════════════════
h1("El resultado en una página", "1")

p("El menú del Metrocar tiene 271 entradas. El menú lateral de Buslink lo replica entero y "
  "muestra 105 ítems deshabilitados, lo que da la sensación de que falta muchísimo. "
  "El cruce real dice otra cosa.")

tabla(
    ["", "Cantidad", "Qué significa"],
    [
        ["Ítems del menú del Metrocar sin equivalente en Buslink", [("~102", {"bold": True, "size": 12})],
         "El número que asusta y desenfoca"],
        ["De esos: módulos muertos o utilidades propias del FoxPro", [("~62", {"bold": True, "size": 12, "color": GRIS})],
         "No tienen a quién servir. Se descartan con evidencia, no por intuición."],
        ["Pantallas vivas que sí hay que migrar", [("~40", {"bold": True, "size": 12, "color": NARANJ})],
         "El trabajo real que queda por delante"],
        ["De esas: el núcleo crítico", [("~20", {"bold": True, "size": 12, "color": ROJO})],
         "Bloquean el Día D o son operación de todos los días"],
    ],
    widths=[7.6, 2.4, 6.0], size=10,
)

h2("Cómo se reparte lo que falta, por módulo")

tabla(
    ["Módulo del Metrocar", "Ítems", "Migrado o con andamiaje", "Falta", "De lo que falta, ¿vale la pena?"],
    [
        ["Sistema (Accesos)", "4", "1", "3", [("2 sí", {"color": ROJO}), (" · 1 a confirmar", G)]],
        ["Reservas", "11", "8", "3", [("3 sí", {"color": ROJO})]],
        ["Tráfico", "20", "13", "7", [("6 sí", {"color": ROJO}), (" · 1 trivial", G)]],
        ["Vehículos y Choferes", "11", "7", "4", [("0 — los 4 son módulos muertos", {"color": GRIS})]],
        ["Facturación", "29", "6", "23", [("17 sí", {"color": ROJO}), (" · 5 a confirmar · 1 no", G)]],
        ["Taller", "16", "0", "16", [("0 — módulo muerto desde 2019", {"color": GRIS})]],
        ["Combustible", "10", "10", "0", [("—", G)]],
        ["ABM del sistema", "24", "0", "24", [("21 sí", {"color": ROJO}), (" · 3 no", G)]],
        ["Utilitarios", "27", "3", "24", [("1 sí (crítico)", {"color": ROJO}), (" · 2 a confirmar · 21 no", G)]],
    ],
    widths=[4.4, 1.4, 3.4, 1.4, 5.4], size=9.5,
)

callout(
    "La conclusión de una línea",
    [("Buslink ya cubre los tres módulos que sostienen la operación diaria —Tráfico, Combustible "
      "y Vehículos— casi por completo. Lo que falta se concentra en ", N),
     ("Facturación (los tarifarios) y en el ABM del sistema (los catálogos y los parámetros)", B),
     (", más un puñado de pantallas sueltas de Tráfico que nadie tenía en el radar.", N)],
    fill=HEX_VERDE, color_titulo=VERDE, borde="1B6E3C",
)


# ═══════════════════════ 2. MÉTODO ═══════════════════════
h1("Cómo se hizo este cruce", "2")

p("Vale explicarlo porque el método cambia el resultado.")

h2("Por qué no alcanzaba con mirar el menú lateral de Buslink")

p("El drawer de Buslink replica el menú del Metrocar, y los ítems no migrados aparecen "
  "deshabilitados. Es un buen backlog visual, pero tiene tres problemas para una auditoría:")

bullet("Un ítem puede estar deshabilitado y sin embargo tener la pantalla construida en otro lado "
       "(pasó con «Liquidaciones estimadas», que existe y funciona pero quedó sin link).")
bullet("Un ítem puede estar habilitado pero apuntar a una pantalla que reutiliza otra "
       "(«Estaciones» de Combustible y «Contactos» de Tráfico son el mismo formulario del FoxPro).")
bullet("El drawer no dice si el módulo que falta se sigue usando. Un pendiente de un módulo "
       "muerto no es un pendiente: es ruido.")

h2("Los tres pasos")

pasos([
    [("Extraer el menú real del FoxPro.", B), (" Se parseó ", N), mono("MENU_PRINCIPAL.MPR"),
     (" leyendo las definiciones de barra y sus acciones, para obtener ", N),
     ("el formulario .scx exacto que abre cada entrada", B), (". Son 271 entradas, de las cuales "
      "unas 150 son ítems reales (el resto son separadores).", N)],
    [("Cruzar contra Buslink.", B), (" Las 44 rutas activas del drawer y las 49 páginas ", N),
     mono(".razor"), (" del proyecto, más los 25 interruptores de escritura de ", N),
     mono("AbmFeatureFlags"), (" para distinguir «migrado» de «migrado con la escritura apagada».", N)],
    [("Medir la vigencia de cada módulo en la base.", B), (" Consultas directas a ", N),
     mono("replicaVPF"), (" usando ", N), ("fechas de negocio", B), (" —la fecha de la orden de "
      "trabajo, de la carga, del adelanto— y nunca ", N), mono("_updated_at"),
     (", que es metadata de la réplica y muestra a todas las tablas como «recientes» aunque el "
      "módulo esté muerto hace años. Este paso es el que separa los 102 pendientes de los 40 reales.", N)],
])

callout(
    "La trampa que este método evita",
    [("Si uno mira ", N), mono("_updated_at"), (", el módulo Taller parece vivo: la última "
      "actualización figura en 2026. Pero esa fecha es de cuándo la réplica copió la fila, no de "
      "cuándo alguien la cargó. Mirando ", N), mono("f_i_taller"), (" —la fecha real de ingreso al "
      "taller— la última orden es de ", N), ("agosto de 2019", B), (". Son siete años de diferencia "
      "y cambian por completo la decisión.", N)],
    fill=HEX_AVISO,
)


# ═══════════════════════ 3. GRUPO 1 ═══════════════════════
h1("Grupo 1 — Las críticas", "3")

p("Ocho pantallas. O bloquean el Día D, o alguien las está usando todos los días. "
  "Son las que yo pondría en el plan antes que cualquier otra cosa de este documento.")

h2("3.1 · Las cuatro pantallas de Parámetros")

ficha("Parámetros Generales · Empresa · Pantalla Tráfico · SQL Server GPS",
      "ABM del sistema → Parámetros",
      "parametro · parametro_empresa · parametro_trafico · parametro_sql_server",
      "BLOQUEA EL DÍA D")

campo("QUÉ SON", "Las cuatro pantallas que editan la configuración global del sistema: contadores "
                 "del circuito, umbrales de aviso de vencimientos, cliente interno de prueba, "
                 "parámetros visuales de la planilla de tráfico y los interruptores de la "
                 "integración GPS.")
campo("POR QUÉ ES CRÍTICO",
      [("La tabla ", N), mono("parametro"), (" es una de las 12 que cambian de dueño el Día D. "
       "Ese día el FoxPro queda bloqueado. Si estas pantallas no están migradas, ", N),
       ("a partir de ese momento no hay ninguna interfaz para editar la configuración del sistema", B),
       (": habría que entrar a la base con un cliente SQL y hacer un UPDATE a mano. Y no es "
        "configuración de adorno: ahí viven los contadores que asignan el número de viaje y de "
        "lote.", N)])
campo("ESTADO", [("No migrado y ", N), ("sin plano documentado", B), (". No aparece en ninguno de "
      "los 30 documentos de la biblioteca FoxPro.", N)])
campo("ESFUERZO ESTIMADO", "3 a 4 días para las cuatro, incluida la extracción del plano. "
                           "Son formularios de campos sueltos, sin grilla ni lógica de negocio compleja.")
campo("RECOMENDACIÓN", [("Entra en la Fase 1 del plan, junto con los catálogos. No es un anillo "
       "siguiente: es una condición del corte.", B)])

h2("3.2 · Adicionales Stock")

ficha("Adicionales Stock — Mantenimiento e Ingreso",
      "Tráfico → Adicionales Stock",
      "adicional_stock · adicional_stock_abm",
      "VIVO · SIN DOCUMENTAR")

campo("QUÉ ES", "El registro del stock de adicionales que se carga en cada unidad: agua, hielo, "
                "propinas. Se anota por dominio del vehículo, cliente y fecha.")
campo("EVIDENCIA DE QUE ESTÁ VIVO",
      [("11.321 filas en total, ", N), ("4.576 de ellas desde 2025", B), (", y la última carga es "
       "del ", N), ("7 de julio de 2026", B), (". Las últimas filas son de GATE1 TRAVEL, con "
       "cantidades de 10 a 13 unidades por vehículo. Alguien lo carga todos los días.", N)])
campo("POR QUÉ NADIE LO VIO",
      "No aparece en el plan de migración, ni en el análisis del sistema, ni en la biblioteca de "
      "planos FoxPro. En el menú lateral de Buslink figura como una sección «próximamente» sin "
      "detalle. Salió a la luz recién al cruzar el menú del FoxPro contra la base.")
campo("ESFUERZO ESTIMADO", "3 a 5 días: extraer el plano, la grilla con filtros, el diálogo de "
                           "alta y la exportación. Sin riesgo: no toca el circuito del viaje.")
campo("RECOMENDACIÓN", [("Es la mejor relación esfuerzo/valor que queda pendiente en todo el "
       "proyecto.", B), (" Candidata ideal para una semana en la que haya que esperar una "
       "definición del cliente sobre otra cosa.", N)])

h2("3.3 · Elimina Lotes de Carga")

ficha("Elimina Lotes de Carga",
      "Utilitarios → Reparaciones y Utilitarios",
      "trafico_elimina_lote",
      "BOTÓN DE EMERGENCIA DEL DÍA 1")

campo("QUÉ HACE", "Borra de una sola vez todos los viajes generados por una corrida del armado de "
                  "plantillas, identificados por su número de lote.")
campo("POR QUÉ ES CRÍTICO",
      [("El propio plan de migración lo nombra como el botón de emergencia del primer día: si el "
       "armado de plantillas genera cientos de viajes mal, esto es lo que permite deshacerlo de "
       "forma quirúrgica. Está listado dentro del riesgo número 6 del plan.", N)])
campo("LA TRAMPA",
      [("En el menú del Metrocar no está en Tráfico ni en Reservas: está enterrado en ", N),
       ("Utilitarios → Reparaciones", B), (", entre las herramientas de indexar y compactar "
        "archivos DBF. Es exactamente el tipo de ítem que se descarta en bloque al revisar un "
        "menú por encima, justo cuando más se lo necesita.", N)])
campo("ESFUERZO ESTIMADO", "1 a 2 días. Es una consulta por lote, una previsualización y un borrado "
                           "transaccional.")

h2("3.4 · Cambio de password")

ficha("Cambio de password",
      "Sistema (Accesos) → Cambio de password",
      "usuario_cambio_password",
      "OPERACIÓN DIARIA")

campo("QUÉ HACE", "Permite que cada usuario cambie su propia contraseña.")
campo("EL PROBLEMA HOY",
      [("En Buslink un usuario común ", N), ("no puede cambiarse la clave", B), (": la única forma "
       "es que el supervisor entre al ABM de Usuarios y se la cambie. En el Metrocar cada uno lo "
       "hace desde su propio menú. Es una regresión funcional respecto del sistema viejo, y se "
       "vuelve más visible el día que Buslink sea el único sistema.", N)])
campo("ESFUERZO ESTIMADO", "Medio día. La infraestructura de escritura sobre la tabla de usuarios "
                           "ya existe y funciona.")
campo("NOTA DE SEGURIDAD",
      [("Es el momento natural para resolver también la deuda de las contraseñas en texto plano: "
       "si se implementa el cambio de clave con hash y se migra a cada usuario de forma "
       "transparente en su primer cambio, la deuda se paga sola sin un proyecto aparte.", N)])


# ═══════════════════════ 4. GRUPO 2 ═══════════════════════
h1("Grupo 2 — Vivas, alta prioridad", "4")

p("Quince pantallas de módulos que se usan, pero que no bloquean el corte. Son el primer anillo "
  "después del Día D — con la excepción de los tarifarios, que conviene mirar antes.")

h2("4.1 · Libro de Novedades (3 pantallas)")

tabla(
    ["Pantalla", "Formulario", "Qué hace"],
    [
        ["Libro de Novedades", [mono("libro_novedad")], "La grilla de consulta y mantenimiento de "
         "las novedades cargadas: filtros, búsqueda, edición"],
        ["Envío de correos", [mono("libro_novedad_envia_correo")], "Manda la novedad por mail al "
         "cliente. Esta parte NO se migró con el F2 y sigue viviendo en FoxPro"],
        ["Correos Electrónicos Parámetros", [mono("libro_novedad_parametro")], "La configuración de "
         "destinatarios y plantillas de ese envío (10 filas)"],
    ],
    widths=[4.2, 4.4, 7.4], size=9.5,
)
rich([("Evidencia: ", B), ("48.353 filas, la última del 10 de julio de 2026", N), (". El alta rápida "
      "desde la planilla (tecla F2) ya está construida con andamiaje, pero el módulo completo —ver, "
      "buscar, editar y sobre todo mandar el correo— sigue siendo del Metrocar.", N)], size=10)

h2("4.2 · Tarifario de Venta (4 pantallas)")

tabla(
    ["Pantalla", "Formulario"],
    [
        ["Altas y Copias de Lista de Precios", [mono("lista_precio_cliente")]],
        ["Mantenimiento de Precios", [mono("lista_precio_cliente_mantenimiento")]],
        ["Definición de Lista de Precio", [mono("lista_precio_modelo")]],
        ["Listadores", [mono("lista_precio_tarifario_imprimir")]],
    ],
    widths=[8.0, 8.0], size=9.5,
)
rich([("Evidencia: ", B), ("2.791 filas con vigencias cargadas hasta el 1 de octubre de 2026", N),
      (". Buslink las ", N), ("lee", B), (" —el motor de tarifas validado al 99,4% depende "
      "enteramente de ellas— pero el mantenimiento sigue en el Metrocar. ", N),
      ("Es la razón principal por la que el FoxPro no se va a poder apagar del todo el Día D.", B)], size=10)

h2("4.3 · Adicionales y sus tarifarios (6 pantallas)")

tabla(
    ["Pantalla", "Formulario", "Datos"],
    [
        ["Tarifarios Ventas → Alta y Copia", [mono("adicional_lista_precio")], "95 precios de venta"],
        ["Tarifarios Ventas → Mantenimiento", [mono("adicional_lista_precio_mantenimiento")], ""],
        ["Tarifarios Pagos → Alta y Copia", [mono("adicional_lista_pago")], "329 precios de pago"],
        ["Tarifarios Pagos → Mantenimiento", [mono("adicional_lista_pago_mantenimiento")], ""],
        ["Adicionales", [mono("adicional")], "27 adicionales (agua, hielo, propinas…)"],
        ["Rubro Adicionales", [mono("adicional_rubro")], "8 rubros"],
    ],
    widths=[5.4, 6.4, 4.2], size=9.5,
)
p("La solapa Adicionales de Liquidación a Clientes los valoriza todos los días contra estos "
  "tarifarios. Están vivos y son parte del mismo paquete que el tarifario de venta.", size=10)

h2("4.4 · Crear Plantillas y Tablero de Control")

tabla(
    ["Pantalla", "Módulo", "Formulario", "Nota"],
    [
        ["Crear Plantillas", "Reservas → Plantillas", [mono("reserva_plantilla_crear")],
         "El mantenimiento de plantillas ya está migrado con andamiaje; falta el alta de una "
         "plantilla nueva. Hoy hay 9 plantillas con 574 filas."],
        ["Tablero de Control", "Sistema (Accesos)", [mono("tablero"), (" + ", N), mono("tablero_zoom")],
         "Permiso X, exclusivo del supervisor. Es la vista gerencial del sistema viejo."],
    ],
    widths=[3.2, 3.6, 4.6, 4.6], size=9.5,
)


# ═══════════════════════ 5. GRUPO 3 ═══════════════════════
h1("Grupo 3 — Los catálogos del ABM del sistema", "5")

p("Diecisiete pantallas chicas: código, nombre y poco más. Buslink las lee en combos y grillas, "
  "pero editarlas obliga a abrir el Metrocar. Individualmente son triviales; juntas son "
  "unas tres semanas de trabajo y la última atadura con el sistema viejo para la operación diaria.")

h2("5.1 · Catálogos del negocio (11)")

tabla(
    ["Catálogo", "Formulario", "Filas", "Por qué importa"],
    [
        ["Servicios", [mono("servicio")], "62", "El catálogo central: todo viaje tiene un servicio"],
        ["Cronogramas de Servicio", [mono("cronograma")], "97", "Los cronogramas que asignan las teclas F6-F9"],
        ["Guías", [mono("guia")], "1.135", [("Tabla del Día D", B), (" — la escribe el alta de reservas", N)]],
        ["IATA", [mono("iata")], "106", "Códigos de aeropuerto para los transfers"],
        ["Feriados", [mono("feriado")], "15", [("Cero cargados para 2026", {"color": ROJO, "bold": True}),
         (" — el armado de plantillas los ignora", N)]],
        ["Motivos de Cancelación", [mono("viaje_motivo_cancela")], "6",
         "Lo necesita la operación «Cancelar» de la Fase 3"],
        ["Motivos de Cambio de Cronograma", [mono("viaje_motivo_cambio")], "11",
         [("Lo necesita «Reasignar». Ojo: su modifica está roto en el fuente del FoxPro", N)]],
        ["Motivos de Llegadas Tardes", [mono("viaje_motivo_tarde")], "17", ""],
        ["Zonas", [mono("zona")], "6", "Las escribe el cierre del viaje (la unidad cambia de zona)"],
        ["Dueños", [mono("dueno")], "2", "Titulares de los vehículos"],
        ["Permisos", [mono("permiso")], "14", "Permisos de circulación de la flota"],
    ],
    widths=[4.6, 4.4, 1.6, 5.4], size=9.5,
)

h2("5.2 · Configuración de Facturación (6)")

tabla(
    ["Catálogo", "Formulario", "Nota"],
    [
        ["Empresas para Facturar", [mono("empresa")], "2 empresas cargadas"],
        ["Tipo de Monedas", [mono("moneda_tipo")], "3 monedas"],
        ["Cotizaciones", [mono("moneda_cotizacion")], "32 cotizaciones — lo usa el cálculo de liquidación"],
        ["Bancos", [mono("ctacte_banco")], "Formulario compartido con Cuenta Corriente (módulo muerto)"],
        ["Impuestos sobre Ventas", [mono("ctacte_impuesto")], "Lo usa el IVA de la liquidación"],
        ["Tipos de Comprobantes", [mono("ctacte_tipo_comprobante")], "Ídem"],
    ],
    widths=[4.4, 5.0, 6.6], size=9.5,
)

callout(
    "Un detalle que conviene notar",
    [("Tres de estos seis formularios pertenecen al módulo de Cuenta Corriente, que está muerto "
      "(0 filas en sus cinco tablas). Pero los catálogos que editan —bancos, impuestos y tipos de "
      "comprobante— sí los usa la liquidación a clientes, que está viva. O sea: el módulo se "
      "descarta, pero estas tres pantallitas hay que rescatarlas.", N)],
    fill=HEX_AVISO,
)


# ═══════════════════════ 6. GRUPO 4 ═══════════════════════
h1("Grupo 4 — A confirmar con el dueño", "6")

p("Nueve pantallas donde la base da señales contradictorias. No las clasifico yo: "
  "son preguntas para el dueño del sistema.")

h2("6.1 · Liquidación a Choferes y su tarifario (9 pantallas)")

tabla(
    ["Pantalla", "Formulario"],
    [
        ["Tarifario de Choferes → Altas y Copia", [mono("lista_precio_chofer")]],
        ["Tarifario de Choferes → Mantenimiento", [mono("lista_precio_chofer_mantenimiento")]],
        ["Tarifario de Choferes → Definición", [mono("lista_precio_modelo_chofer")]],
        ["Tarifario de Choferes → Listadores", [mono("lista_precio_chofer_imprimir")]],
        ["Liquidación a Choferes → Genera Liquidación", [mono("liquidacion_chofer_por_hora")]],
        ["Liquidación a Choferes → Parámetros", [mono("liquidacion_chofer_por_hora_parametro")]],
        ["Liquidación a Choferes → Ingreso de Adelantos", [mono("chofer_adelanto_abm")]],
        ["Liquidación a Choferes → Mantenimiento de Adelantos", [mono("chofer_adelanto")]],
        ["Liquidación a Choferes → Motivo de Adelantos", [mono("chofer_adelanto_motivo")]],
    ],
    widths=[9.0, 7.0], size=9.5,
)

campo("LA SEÑAL A FAVOR", [("El tarifario de choferes tiene ", N), ("3.751 filas", B),
      (" y 25 modelos definidos: es un tarifario real y cargado.", N)])
campo("LA SEÑAL EN CONTRA", [("Los adelantos a choferes se detuvieron el ", N),
      ("28 de febrero de 2025", B), (": 375 registros y ninguno después de esa fecha.", N)])
campo("LA PREGUNTA", [("¿Se sigue liquidando a los choferes desde el sistema, o eso pasó a "
      "liquidación de sueldos por otro circuito? Según la respuesta, son nueve pantallas de "
      "trabajo o nueve ítems que se descartan.", B)])

h2("6.2 · Dos utilitarios sueltos")

tabla(
    ["Pantalla", "Formulario", "La duda"],
    [
        ["Km entre Localidades", [mono("km_localidad")], "La tabla de localidades existe (34 filas). "
         "¿Se usa para calcular distancias en las reservas o quedó de otra época?"],
        ["Conectados al sistema", [mono("login_conectado")], "Ya está semi-cubierto: la auditoría de "
         "accesos y la tabla de sesiones construidas en agosto muestran quién está conectado. "
         "¿Alcanza o falta algo del original?"],
    ],
    widths=[3.6, 4.0, 8.4], size=9.5,
)

h2("6.3 · La decisión pendiente de Reservas")

rich([("La ", N), ("importación de reservas desde Excel", B), (" (", N), mono("importa_excel_viaje"),
      (", 28 columnas y 3 etapas de validación) es la tercera puerta de alta de la Fase 4 y el "
       "descarte declarado del plan si el cronograma aprieta. ", N),
      ("Esta decisión hay que tomarla antes del Día D, no ese día", B), (": si se descarta, el "
       "workaround del día 1 es cargar a mano o por plantilla, y eso hay que avisarlo con "
       "anticipación a quien hoy usa el Excel.", N)])


# ═══════════════════════ 7. GRUPO 5 ═══════════════════════
h1("Grupo 5 — Lo que no conviene migrar", "7")

p("Sesenta y dos ítems. Cada uno con la evidencia que respalda el descarte. Este es el capítulo "
  "para llevar a la reunión con el dueño y pedir una firma.")

h2("7.1 · Módulos muertos, medidos en la base")

tabla(
    ["Módulo", "Ítems", "Formularios", "Evidencia"],
    [
        [[("Taller", B), ("\nMantenimiento · Agenda de vencimiento · OT ingreso · OT egreso · "
          "Stock (4) · Artículos (3) · Ítem del service", N)], "12",
         [mono8("taller_*"), (" (20 forms)", {"size": 8.5})],
         [("Última orden de trabajo: ", N), ("5 de agosto de 2019", B), (". ", N),
          mono8("taller_stock"), (" y ", {"size": 9}), mono8("taller_deposito"),
          (" con 0 filas: el stock nunca se usó.", N)]],
        [[("Chequeo de unidades", B), ("\nConsulta · Ingreso · Impresión · Parámetros de ítem", N)], "4",
         [mono8("auditoria*")],
         [("Está dentro del menú Taller. Verificado: la última de las 61 auditorías es del ", N),
          ("9 de agosto de 2019", B), (".", N)]],
        [[("Cuenta Corriente", B), ("\nNo tiene menú propio, pero aporta 3 catálogos", N)], "—",
         [mono8("ctacte_*"), (" (17 forms)", {"size": 8.5})],
         [("Las 5 tablas con ", N), ("0 filas", B), (". Se programó y nunca se usó. Rescatar solo "
          "bancos, impuestos y tipos de comprobante (ver capítulo 5).", N)]],
        [[("Apercibimientos", B), ("\nApercibimientos · Motivos", N)], "2",
         [mono8("chofer_sancion"), (" · ", {"size": 9}), mono8("chofer_sancion_motivo")],
         [("Ambas tablas con ", N), ("0 filas", B), (".", N)]],
        [[("Capacitaciones", B), ("\nConsulta · Armado · Cursos-Descripción", N)], "3",
         [mono8("chofer_curso_*")],
         [("417 registros, el último de ", N), ("junio de 2013", B), (".", N)]],
        [[("Liquidación a Fleteros", B)], "1", [mono8("liquidacion_fletero_nueva")],
         [("12 liquidaciones tipo PROVEEDOR en todo el histórico, la última de ", N),
          ("diciembre de 2023", B), (".", N)]],
        [[("Nacionalidades · Profesiones", B)], "2",
         [mono8("nacionalidad"), (" · ", {"size": 9}), mono8("profesion")],
         "1 fila y 0 filas respectivamente."],
    ],
    widths=[4.2, 1.2, 3.6, 7.0], size=9,
)

h2("7.2 · Utilitarios propios del FoxPro (21 ítems)")

tabla(
    ["Grupo", "Ítems", "Por qué no aplica"],
    [
        ["Herramientas de escritorio", "Scheduler · Agenda · Calendario · Calculadora · Editor de imagen",
         "Son accesorios que el FoxPro incorporaba porque en su época no había alternativa. "
         "Hoy los resuelve el sistema operativo o el navegador."],
        ["Backup", "Backup · Backup Minimiza · Editor de log de backup",
         "Copiaban archivos DBF. El resguardo de SQL Server se hace con el motor de la base, no "
         "desde la aplicación."],
        ["Reparaciones de archivos", "Normaliza km · Cronograma en plantillas · Viajes · "
         "Indexa base · Pasa Excel a viajes · Compacta base · Destinos · Guías",
         [("Reparan corrupción e índices de archivos DBF, un problema que SQL Server no tiene. ", N),
          ("La excepción es «Elimina Lotes de Carga», que sí hay que migrar (capítulo 3).", B)]],
        ["Comunicación interna", "Chat - Conversar · Chat - Server",
         "Chat punto a punto entre puestos de la red local."],
        ["Otros", "Editor de log de errores · Configuración regional · Servicio XML",
         [("El servicio XML (", N), mono8("viaje_intranet"), (") es la vía del GPS, ya verificada "
          "como apagada en producción.", N)]],
    ],
    widths=[3.4, 6.0, 6.6], size=9,
)

callout(
    "Lo que se gana firmando este descarte",
    [("Se sacan del alcance ", N), ("62 ítems de menú y unos 65 formularios del FoxPro", B),
     (". El backlog visible pasa de 105 pendientes a poco más de 40, y el proyecto deja de "
      "parecer infinito. Sugerencia concreta: en el menú de Buslink, cambiar el rótulo de estos "
      "ítems de «próximamente» a ", N), ("«no se migra — consultar en Metrocar»", B),
     (". La misma pantalla, contando una historia mucho más cercana a la verdad.", N)],
    fill=HEX_VERDE, color_titulo=VERDE, borde="1B6E3C",
)


# ═══════════════════════ 8. BACKLOG ═══════════════════════
h1("Backlog priorizado", "8")

p("El orden en que yo lo atacaría, con estimación de esfuerzo. Las estimaciones son mías y "
  "asumen un desarrollador asistido, con el patrón de pantallas ya consolidado del proyecto.")

tabla(
    ["Orden", "Qué", "Ítems", "Días", "Cuándo"],
    [
        ["1", [("Las 4 pantallas de Parámetros", B), ("\nCondición del corte: sin esto, después del "
          "Día D no hay dónde editar la configuración", N)], "4", "3-4",
         [("Antes del Día D", {"color": ROJO, "bold": True})]],
        ["2", [("Elimina Lotes de Carga", B), ("\nEl botón de emergencia del primer día", N)], "1", "1-2",
         [("Antes del Día D", {"color": ROJO, "bold": True})]],
        ["3", [("Los 3 catálogos que usa la Fase 3", B), ("\nMotivos de cancelación, motivos de "
          "cambio de cronograma y feriados", N)], "3", "2-3",
         [("Antes del Día D", {"color": ROJO, "bold": True})]],
        ["4", [("Cambio de password", B), ("\nCon hash, resolviendo de paso la deuda de seguridad", N)],
         "1", "0,5-1", [("Antes del Día D", {"color": ROJO, "bold": True})]],
        ["5", [("Adicionales Stock", B), ("\nMódulo vivo, sin riesgo, sin documentar", N)], "2", "3-5",
         [("Cuando haya un hueco", {"color": NARANJ, "bold": True})]],
        ["6", [("Tarifarios: venta + adicionales", B), ("\nLa atadura principal con el FoxPro "
          "después del corte", N)], "10", "10-15",
         [("Primer anillo post-corte", {"color": NARANJ, "bold": True})]],
        ["7", [("Libro de Novedades completo", B), ("\nIncluido el envío de correos", N)], "3", "4-6",
         [("Primer anillo post-corte", {"color": NARANJ, "bold": True})]],
        ["8", [("Los 8 catálogos restantes del negocio", B), ("\nServicios, cronogramas, guías, "
          "IATA, zonas, dueños, permisos, motivos de llegada tarde", N)], "8", "6-8",
         [("Segundo anillo", G)]],
        ["9", [("Los 6 catálogos de configuración de Facturación", B)], "6", "4-5", [("Segundo anillo", G)]],
        ["10", [("Crear Plantillas · Tablero de Control · Bandas horarias · Controles de estado · "
          "Clientes Tarifas/Descuentos/Empresa", B)], "7", "6-8", [("Segundo anillo", G)]],
        ["11", [("Liquidación a Choferes + su tarifario", B), ("\nSolo si el dueño confirma que "
          "sigue en uso", N)], "9", "8-12", [("A confirmar", {"color": NARANJ})]],
        ["12", [("Importa Reservas desde Excel", B), ("\nDecisión de alcance pendiente", N)], "1", "5-8",
         [("A decidir", {"color": NARANJ})]],
    ],
    widths=[1.2, 7.4, 1.2, 1.4, 4.8], size=9,
)

h2("Los totales")

tabla(
    ["Tramo", "Ítems", "Esfuerzo estimado"],
    [
        [[("Antes del Día D (órdenes 1 a 4)", B)], "9", [("7 a 10 días", B)]],
        ["Primer anillo post-corte (5 a 7)", "15", "17 a 26 días"],
        ["Segundo anillo (8 a 10)", "21", "16 a 21 días"],
        ["Condicionales (11 y 12)", "10", "13 a 20 días"],
        [[("TOTAL de lo que vale la pena migrar", B)], [("55", B)], [("53 a 77 días", B)]],
    ],
    widths=[7.6, 2.4, 6.0], size=9.5,
)

callout(
    "Cómo leer estos totales",
    [("Los 53 a 77 días son ", N), ("todo lo que queda del Metrocar que vale la pena migrar", B),
     (", y son independientes del trabajo del circuito ", N), ("viaje", {"italic": True}),
     (" (la toolbar de Tráfico, el Graba de Facturación y el ensayo general), que es el camino "
      "crítico hacia el Día D. Lo importante de esta tabla es la primera fila: ", N),
     ("9 ítems y unos 7 a 10 días son condición del corte", B), (" y hoy no están en ningún plan.", N)],
    fill=HEX_CLARO, color_titulo=AZUL, borde=HEX_AZUL,
)


# ═══════════════════════ ANEXO A ═══════════════════════
h1("Anexo A — El menú completo del Metrocar", "A")

p("El menú entero, ítem por ítem, con el formulario del FoxPro y el estado en Buslink. "
  "Pensado para imprimir y tildar.", size=10)

rich([("MIGRADO", {"bold": True, "color": VERDE, "size": 9}), (" funciona en Buslink   ·   ", {"size": 9}),
      ("ANDAMIAJE", {"bold": True, "color": NARANJ, "size": 9}),
      (" pantalla lista, escritura apagada hasta el Día D   ·   ", {"size": 9}),
      ("FALTA", {"bold": True, "color": ROJO, "size": 9}), (" hay que migrarlo   ·   ", {"size": 9}),
      ("NO MIGRAR", {"bold": True, "color": GRIS, "size": 9}), (" descarte recomendado", {"size": 9})],
     size=9)

W = [4.7, 4.5, 2.1, 4.7]
HDR = ["Ítem del menú", "Formulario FoxPro", "Estado", "Nota"]


def menu_tabla(filas):
    fills = [HEX_MUERTO if f[2] is M_NO else None for f in filas]
    tabla(HDR, filas, widths=W, size=8.5, fills=fills, zebra=False)


h2("Sistema (Accesos)")
menu_tabla([
    ["Accesos", [mono8("usuario")], M_OK, "/usuarios-abm — escritura real"],
    ["Cambio de password", [mono8("usuario_cambio_password")], M_FAL, "El usuario no puede cambiar su clave"],
    ["Cambiar Empresa", "—", M_FAL, "2 empresas cargadas — confirmar si se usa"],
    ["Tablero de Control", [mono8("tablero")], M_FAL, "Permiso X"],
])

h2("Reservas")
menu_tabla([
    ["Reservas Especiales", [mono8("reserva_transportacion_con_adicional")], M_AND, "Puerta de alta — Día D"],
    ["Reservas por Plantillas", [mono8("reserva_plantilla_armar")], M_AND, "Con preview dry-run"],
    ["Clientes", [mono8("cliente")], M_OK, "Lectura"],
    ["Operadores", [mono8("cliente_operador")], M_AND, ""],
    ["Grupos", [mono8("cliente_grupo")], M_AND, "Cascada sobre viaje — Día D"],
    ["Destinos", [mono8("destino")], M_AND, ""],
    ["Plantillas → Crear Plantillas", [mono8("reserva_plantilla_crear")], M_FAL, ""],
    ["Plantillas → Mantenimiento", [mono8("reserva_plantilla_mantenimiento")], M_AND, ""],
    ["Importa Reservas desde Excel", [mono8("importa_excel_viaje")], M_FAL, "Descarte posible"],
    ["Informes → Fecha y banda horaria", [mono8("trafico_resumen_horario")], M_OK, "+2 informes propios"],
    ["Informes → Bandas Horarios", [mono8("trafico_resumen_horario_banda")], M_FAL, "Catálogo de 6 bandas"],
])

h2("Tráfico")
menu_tabla([
    ["Operación de Tráfico", [mono8("trafico2")], M_OK, "La pantalla central"],
    ["Web de Aeropuertos 2000", "— (link externo)", M_FAL, "Es un hipervínculo"],
    ["Cabeceras - Recorridos", [mono8("cabecera_recorrido")], M_AND, ""],
    ["Francos → Ingreso", [mono8("chofer_franco_abm")], M_AND, "Alta masiva"],
    ["Francos → Mantenimiento", [mono8("chofer_franco")], M_AND, ""],
    ["Francos → Auditoría", [mono8("chofer_franco_auditoria")], M_AND, "Informe matriz"],
    ["Viáticos → Viáticos", [mono8("chofer_viatico")], M_AND, "Tabla vacía"],
    ["Viáticos → Motivo", [mono8("chofer_viatico_motivo")], M_AND, "Vacía"],
    ["Viáticos → Forma Liquidación", [mono8("chofer_viatico_liquida")], M_AND, "Vacía"],
    ["Voucher Recepción", [mono8("trafico_voucher")], M_AND, "Escribe viaje — Día D"],
    ["Guardia", [mono8("trafico_guardia")], M_AND, ""],
    ["Contactos y Proveedores → Contactos", [mono8("estacion")], M_AND, "Compartido con Combustible"],
    ["Contactos y Proveedores → Rubros", [mono8("estacion_rubro")], M_AND, ""],
    ["Lista de pasajeros", [mono8("trafico_pasajero_planilla")], M_OK, ""],
    ["Controles sobre estados de reservas", [mono8("viaje_estado_query")], M_FAL, "Es un informe"],
    ["Adicionales Stock → Mantenimiento", [mono8("adicional_stock")], M_FAL, [("VIVO · sin documentar", {"bold": True, "color": ROJO, "size": 8.5})]],
    ["Adicionales Stock → Ingreso", [mono8("adicional_stock_abm")], M_FAL, [("VIVO · sin documentar", {"bold": True, "color": ROJO, "size": 8.5})]],
    ["Libro de Novedades → Libro", [mono8("libro_novedad")], M_FAL, "48.353 filas"],
    ["Libro de Novedades → Envío de correos", [mono8("libro_novedad_envia_correo")], M_FAL, ""],
    ["Libro de Novedades → Parámetros correo", [mono8("libro_novedad_parametro")], M_FAL, ""],
])

h2("Vehículos y Choferes")
menu_tabla([
    ["Choferes", [mono8("chofer")], M_OK, "Ficha de 5 pestañas"],
    ["Apercibimientos → Apercibimientos", [mono8("chofer_sancion")], M_NO, "0 filas"],
    ["Apercibimientos → Motivos", [mono8("chofer_sancion_motivo")], M_NO, "0 filas"],
    ["Capacitaciones → Consulta", [mono8("chofer_curso_consulta")], M_NO, "Última: 2013"],
    ["Capacitaciones → Armado", [mono8("chofer_curso_arma")], M_NO, "Última: 2013"],
    ["Odómetros", [mono8("vehiculo_km")], M_OK, ""],
    ["Siniestros", [mono8("siniestro")], M_OK, "5 solapas"],
    ["Vehículos - Flota", [mono8("vehiculo")], M_OK, "Ficha de 6 pestañas"],
    ["Agenda de Vencimientos", [mono8("agenda_vencimiento")], M_OK, ""],
    ["Fleteros", [mono8("fletero")], M_AND, "Compartido con Facturación"],
    ["Tipo de Vehículos", [mono8("vehiculo_tipo")], M_AND, ""],
])

h2("Facturación")
menu_tabla([
    ["Clientes → ABM Clientes", [mono8("cliente")], M_OK, "Lectura"],
    ["Clientes → Clientes Tarifas", [mono8("cliente_tarifa")], M_FAL, "4 filas"],
    ["Clientes → Clientes Descuentos", [mono8("cliente_descuento")], M_FAL, "0 filas"],
    ["Clientes → Empresa Facturación", [mono8("cliente_cambia_empresa_fc")], M_FAL, ""],
    ["Fleteros", [mono8("fletero")], M_AND, "Mismo form que V. y Choferes"],
    ["Grupos", [mono8("cliente_grupo")], M_AND, "Mismo form que Reservas"],
    ["Tarifario de Venta → Altas y Copias", [mono8("lista_precio_cliente")], M_FAL, "2.791 filas vigentes"],
    ["Tarifario de Venta → Mantenimiento", [mono8("lista_precio_cliente_mantenimiento")], M_FAL, ""],
    ["Tarifario de Venta → Definición", [mono8("lista_precio_modelo")], M_FAL, "49 modelos"],
    ["Tarifario de Venta → Listadores", [mono8("lista_precio_tarifario_imprimir")], M_FAL, ""],
    ["Tarifario de Choferes → Altas y Copia", [mono8("lista_precio_chofer")], M_FAL, "3.751 filas"],
    ["Tarifario de Choferes → Mantenimiento", [mono8("lista_precio_chofer_mantenimiento")], M_FAL, ""],
    ["Tarifario de Choferes → Definición", [mono8("lista_precio_modelo_chofer")], M_FAL, "25 modelos"],
    ["Tarifario de Choferes → Listadores", [mono8("lista_precio_chofer_imprimir")], M_FAL, ""],
    ["Adicionales → Tarif. Ventas: alta", [mono8("adicional_lista_precio")], M_FAL, "95 precios"],
    ["Adicionales → Tarif. Ventas: mant.", [mono8("adicional_lista_precio_mantenimiento")], M_FAL, ""],
    ["Adicionales → Tarif. Pagos: alta", [mono8("adicional_lista_pago")], M_FAL, "329 precios"],
    ["Adicionales → Tarif. Pagos: mant.", [mono8("adicional_lista_pago_mantenimiento")], M_FAL, ""],
    ["Adicionales → Adicionales", [mono8("adicional")], M_FAL, "27 adicionales"],
    ["Adicionales → Rubro Adicionales", [mono8("adicional_rubro")], M_FAL, "8 rubros"],
    ["Resumen de Liquidaciones", [mono8("liquidacion_cliente")], M_OK, "Maestro-detalle"],
    ["Liquidación a Clientes", [mono8("facturacion_cliente_nueva")], M_OK, "Falta el Graba (Fase 5)"],
    ["Liquidación a Fleteros", [mono8("liquidacion_fletero_nueva")], M_NO, "12 liq., última 2023"],
    ["Liq. a Choferes → Genera", [mono8("liquidacion_chofer_por_hora")], M_FAL, "A confirmar"],
    ["Liq. a Choferes → Parámetros", [mono8("liquidacion_chofer_por_hora_parametro")], M_FAL, "A confirmar"],
    ["Liq. a Choferes → Ingreso Adelantos", [mono8("chofer_adelanto_abm")], M_FAL, "Última: 02/2025"],
    ["Liq. a Choferes → Mant. Adelantos", [mono8("chofer_adelanto")], M_FAL, "A confirmar"],
    ["Liq. a Choferes → Motivo Adelantos", [mono8("chofer_adelanto_motivo")], M_FAL, "A confirmar"],
    ["Liquidaciones estimadas", [mono8("facturacion_cliente_estimada")], M_OK, [("Sin link en el menú", {"color": ROJO, "size": 8.5})]],
])

h2("Taller — módulo muerto desde 2019")
menu_tabla([
    ["Mantenimiento", [mono8("taller_servicio")], M_NO, "Última OT: 05/08/2019"],
    ["Agenda de Vencimiento", [mono8("taller_servicio_aviso")], M_NO, ""],
    ["Órdenes de trabajo: Ingreso", [mono8("taller_orden_trabajo")], M_NO, ""],
    ["Órdenes de trabajo: Egreso", [mono8("taller_orden_trabajo_final")], M_NO, ""],
    ["Stock → Consulta", [mono8("taller_stock")], M_NO, "0 filas"],
    ["Stock → Ingresos / Egresos", [mono8("taller_stock_movimiento_ingreso_egreso")], M_NO, "0 filas"],
    ["Stock → Movimiento entre depósitos", [mono8("taller_stock_movimiento_deposito")], M_NO, "0 filas"],
    ["Stock → Depósitos", [mono8("taller_deposito")], M_NO, "0 filas"],
    ["Chequeo de unidades → Consulta", [mono8("auditoria")], M_NO, "Última: 09/08/2019"],
    ["Chequeo de unidades → Ingreso", [mono8("auditoria_abm")], M_NO, ""],
    ["Chequeo de unidades → Impresión", [mono8("auditoria_imprime")], M_NO, ""],
    ["Chequeo de unidades → Parámetros ítem", [mono8("auditoria_item_abm")], M_NO, ""],
    ["Artículos → Artículos", [mono8("taller_articulo")], M_NO, "30 filas"],
    ["Artículos → Tipo Vehículos", [mono8("taller_articulo_veh")], M_NO, "9 filas"],
    ["Artículos → Grupo Motor", [mono8("taller_articulo_motor")], M_NO, "4 filas"],
    ["Ítem del service", [mono8("taller_servicio_item")], M_NO, "21 filas"],
])

h2("Combustible — módulo completo")
menu_tabla([
    ["Promedio de Consumos", [mono8("vehiculo_combustible_consumo")], M_OK, "Informe mejorado"],
    ["ABM y Conciliación cargas", [mono8("vehiculo_combustible_mant_sobre_lote")], M_AND, ""],
    ["Saldos de Estaciones", [mono8("vehiculo_estacion_saldo")], M_OK, "Sin uso desde 2017"],
    ["Depósitos → Carga", [mono8("vehiculo_estacion_saldo_carga")], M_AND, ""],
    ["Depósitos → Mantenimiento", [mono8("vehiculo_estacion_saldo_mant")], M_AND, ""],
    ["Estaciones", [mono8("estacion")], M_AND, "Reusa Contactos de Tráfico"],
    ["Rubro de Consumos", [mono8("estacion_rubro")], M_AND, "Reusa Rubros de Tráfico"],
    ["Artículos por Rubro de Consumo", [mono8("estacion_rubro_articulo")], M_AND, ""],
    ["(propio de Buslink) Control de cargas", "—", M_OK, "No existe en FoxPro"],
    ["(propio de Buslink) Consumo Mensual", "—", M_OK, "No existe en FoxPro"],
])

h2("ABM del sistema — ninguno migrado")
menu_tabla([
    ["Cronogramas de Servicio", [mono8("cronograma")], M_FAL, "97 filas"],
    ["Servicios", [mono8("servicio")], M_FAL, "62 filas"],
    ["IATA", [mono8("iata")], M_FAL, "106 filas"],
    ["Cursos - Descripción", [mono8("chofer_curso_parametro")], M_NO, "Muerto con Capacitaciones"],
    ["Guías", [mono8("guia")], M_FAL, "1.135 — tabla del Día D"],
    ["Dueños", [mono8("dueno")], M_FAL, "2 filas"],
    ["Permisos", [mono8("permiso")], M_FAL, "14 filas"],
    ["Zonas", [mono8("zona")], M_FAL, "6 — las escribe el cierre del viaje"],
    ["Nacionalidades", [mono8("nacionalidad")], M_NO, "1 fila"],
    ["Profesiones", [mono8("profesion")], M_NO, "0 filas"],
    ["Feriados", [mono8("feriado")], M_FAL, [("0 cargados de 2026", {"color": ROJO, "bold": True, "size": 8.5})]],
    ["Motivos de Cancelaciones", [mono8("viaje_motivo_cancela")], M_FAL, "Lo pide la Fase 3"],
    ["Motivos de Cambio de Cronogramas", [mono8("viaje_motivo_cambio")], M_FAL, "Lo pide Reasignar"],
    ["Motivos de Llegadas Tardes", [mono8("viaje_motivo_tarde")], M_FAL, "17 filas"],
    ["Parámetros → Generales", [mono8("parametro")], M_FAL, [("Bloquea el Día D", {"color": ROJO, "bold": True, "size": 8.5})]],
    ["Parámetros → Empresa", [mono8("parametro_empresa")], M_FAL, [("Bloquea el Día D", {"color": ROJO, "bold": True, "size": 8.5})]],
    ["Parámetros → Pantalla Tráfico", [mono8("parametro_trafico")], M_FAL, [("Bloquea el Día D", {"color": ROJO, "bold": True, "size": 8.5})]],
    ["Parámetros → SQL Server para GPS", [mono8("parametro_sql_server")], M_FAL, [("Bloquea el Día D", {"color": ROJO, "bold": True, "size": 8.5})]],
    ["Facturación → Empresas para Facturar", [mono8("empresa")], M_FAL, "2 empresas"],
    ["Facturación → Tipo de Monedas", [mono8("moneda_tipo")], M_FAL, "3 monedas"],
    ["Facturación → Cotizaciones", [mono8("moneda_cotizacion")], M_FAL, "32 cotizaciones"],
    ["Facturación → Bancos", [mono8("ctacte_banco")], M_FAL, "Rescatar de cta. cte."],
    ["Facturación → Impuestos sobre Ventas", [mono8("ctacte_impuesto")], M_FAL, "Lo usa el IVA"],
    ["Facturación → Tipos de Comprobantes", [mono8("ctacte_tipo_comprobante")], M_FAL, ""],
])

h2("Utilitarios")
menu_tabla([
    ["Scheduler", [mono8("scheduler")], M_NO, "Accesorio de escritorio"],
    ["Agenda", [mono8("agenda")], M_NO, "Accesorio"],
    ["Calendario", [mono8("calendario")], M_NO, "Accesorio"],
    ["Calculadora", [mono8("calculadora")], M_NO, "Accesorio"],
    ["Editor de Imagen", [mono8("vfppaint")], M_NO, "Accesorio"],
    ["Backup / Backup Minimiza / Editor Log", [mono8("backup")], M_NO, "Copiaba DBF"],
    ["Reparaciones → Normaliza Km turismo", [mono8("reserva_normaliza_km")], M_NO, ""],
    ["Reparaciones → Elimina Lotes de Carga", [mono8("trafico_elimina_lote")], M_FAL, [("CRÍTICO — botón de emergencia", {"color": ROJO, "bold": True, "size": 8.5})]],
    ["Reparaciones → Cronograma en Plantillas", [mono8("reserva_plantilla_repara")], M_NO, ""],
    ["Reparaciones → Viajes", [mono8("trafico_repara")], M_NO, "Repara DBF"],
    ["Reparaciones → Indexa Base de Tráfico", [mono8("trafico_indexa")], M_NO, "No aplica a SQL"],
    ["Reparaciones → Pasa Excel a Viajes", [mono8("pasa_excel_a_viaje")], M_NO, ""],
    ["Reparaciones → Compacta Base de Datos", [mono8("compacta_base")], M_NO, "No aplica a SQL"],
    ["Reparaciones → Destinos / Guías", [mono8("destino_repara"), (" · ", {"size": 8.5}), mono8("guia_repara")], M_NO, ""],
    ["Editor Log Errores", "—", M_NO, "Log del FoxPro"],
    ["Reservas por Cliente", [mono8("viaje_analisis")], M_OK, "Migrado y mejorado"],
    ["Viajes por Choferes", [mono8("viaje_analisis_chofer")], M_OK, "Migrado y mejorado"],
    ["Km Unidades Vs Servicios", [mono8("viaje_analisis_km")], M_OK, "Bug del % corregido"],
    ["Conectados al sistema", [mono8("login_conectado")], M_FAL, "Semi-cubierto por Auditoría de accesos"],
    ["Chat - Conversar / Chat - Server", [mono8("frm_remoto"), (" · ", {"size": 8.5}), mono8("frm_host")], M_NO, "Chat de LAN"],
    ["Servicio XML", [mono8("viaje_intranet")], M_NO, "Vía del GPS, apagada"],
    ["Configuración Regional", "—", M_NO, "Del sistema operativo"],
    ["Km entre Localidades", [mono8("km_localidad")], M_FAL, "Confirmar si se usa"],
])


# ── pie ──
doc.add_paragraph()
par = doc.add_paragraph()
border(par, "top", size=8, color="D5DAE3", space=6)
par.paragraph_format.space_before = Pt(14)
run(par, "Buslink · Metrocar Nortur — Qué pantallas del Metrocar todavía no están en Buslink · "
         "9 de agosto de 2026", size=9, color=GRIS)
par2 = doc.add_paragraph()
run(par2, "Método: parseo de MENU_PRINCIPAL.MPR (271 entradas) para obtener el formulario .scx de "
          "cada ítem, cruce contra las 44 rutas activas y las 49 páginas de Buslink, y medición de "
          "vigencia por módulo con consultas directas a replicaVPF usando fechas de negocio.",
    size=9, color=GRIS)

for sec in doc.sections:
    footer_par = sec.footer.paragraphs[0]
    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_par.add_run()
    r.font.size = Pt(8.5)
    r.font.name = BODY
    r.font.color.rgb = GRIS
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    r._r.append(fld1); r._r.append(instr); r._r.append(fld2)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "INFORME_PANTALLAS_FALTANTES.docx")
doc.save(out)
print("OK ->", out)
